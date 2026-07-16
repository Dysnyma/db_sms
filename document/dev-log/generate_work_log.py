"""
生成每日工作总结 .docx
用法：python document/dev-log/generate_work_log.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUT = os.path.join(BASE, "document", "dev-log", "每日工作总结.docx")

doc = Document()

# ── 样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.35

for sec in doc.sections:
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

def sf(run, name):
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rPr.append(rf)
    rf.set(qn('w:eastAsia'), name)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; sf(r, '黑体')
        if level == 1: r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        elif level == 2: r.font.size = Pt(14); r.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
        elif level == 3: r.font.size = Pt(12)

def P(text, bold=False, sz=11, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(sz)
    run.bold = bold
    sf(run, '宋体')
    return p

def BULLET(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    sf(run, '宋体')

def TAG(text):
    """返回带背景色的 tag 元素"""
    p = doc.add_paragraph()
    run = p.add_run(f'  {text}  ')
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sf(run, '微软雅黑')
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '2E86C1')
    shading.set(qn('w:val'), 'clear')
    pPr = p._element.get_or_add_pPr()
    pPr.append(shading)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p

def TABLE(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)

# ══ 封面 ══
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('数据库系统课程设计\n每日工作总结')
r.font.size = Pt(28); r.bold = True; r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
sf(r, '黑体')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('学生姓名：蔡坤灿    专    业：计算机科学与技术')
r.font.size = Pt(14)
sf(r, '宋体')

for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026年7月')
r.font.size = Pt(14)
sf(r, '宋体')

doc.add_page_break()

# ══ 总览表 ══
H('工作总览', level=1)
TABLE(['日期', '主要工作', '产出物'],
      [['07/07', '数据库设计 + Python CLI 骨架', '8张表/3视图/5触发器/9SP/2函数 + CLI 应用'],
       ['07/08', 'CLI 功能完整实现', '6个CRUD+分页器+三角色+排序规则修复'],
       ['07/09', 'Streamlit 双界面搭建', '22个Web页面+管理页重构+备份恢复'],
       ['07/13', '目录规范化 + 代码质量修复', 'Pylint/Flake8清零+目录结构规范'],
       ['07/14', '功能增强 + 图表 + 校验', 'Plotly图表/Pydantic校验/自动班级名/专业管理'],
       ['07/15', '大数据量支持', '200万数据生成器+LOAD DATA优化+性能调优'],
       ['07/16', '文档收尾 + 答辩准备', 'PPT+说明书重构+代码附录+源码打包']])

doc.add_page_break()

# ════════════════════════════════════════
# 7月7日
# ════════════════════════════════════════
H('第一天：2026年7月7日 — 数据库设计 + CLI 应用骨架', level=1)
TAG('选题切换')
P('原定题目一"进销存管理系统"已完成 17 张表、8 个视图、7 个触发器和 15 个存储过程的设计与开发。经综合考虑后决定切换至题目二"学生成绩管理系统"，原进销存代码保留至 archive/ 目录作为学习参考。新数据库命名为 db_sms（Student Management System）。')

TAG('概念结构设计')
P('确定实体关系模型的核心决策：学生与班级采用 1:N 关系，在 student 表中添加 class_id 外键，不建中间表，避免过度设计。学生与课程之间为 M:N 关系，通过 enrollment 中间表实现，该中间表同时承担成绩存储职责。课程与教师之间为 M:N 关系，采用 teacher_course（讲授资格）和 course_offering（实际排课）两层设计，实现逻辑分离。')

TAG('逻辑结构设计')
P('完成 8 张表的完整 DDL 设计：class（班级）、student（学生）、course（课程）、teacher（教师）、course_offering（排课）、enrollment（选课）、teacher_course（讲授关系）、grade_point_rule（绩点规则）。全部表采用 InnoDB 引擎、utf8mb4 字符集、utf8mb4_unicode_ci 排序规则、INT 自增主键，且全部支持逻辑删除（is_deleted 字段）。')

TAG('数据库对象开发')
P('完成 3 个视图（v_student_message、v_course_plan、v_enrollment）、5 个触发器（选课前名额检查、选课后人数+1、退选后人数-1、录成绩重算学籍分和 GPA）、9 个存储过程（sp_show_courses、sp_enroll、sp_unenroll、sp_grade_input、sp_student_roster、sp_class_grade_report、sp_student_semester_avg、sp_teacher_info、sp_teacher_list）以及 2 个存储函数（fn_get_student_id、fn_get_teacher_id）的编写。')

TAG('AI 辅助与修正')
P('AI 辅助生成了大部分 DDL 和存储过程的框架代码，但在以下方面需要人工修正：course_offering 表的 effective_time 字段与 enroll_start_time 语义重复，已删除；sp_enroll 未能检查同课程不同教师的排课重复，新增了第 4 步 JOIN 校验；sp_show_courses 同样存在该问题，新增了 course_id NOT IN 子查询；MySQL 8.0 默认 utf8mb4_0900_ai_ci 排序规则与数据库 utf8mb4_unicode_ci 冲突（1267 错误），通过在 pymysql 连接参数中加 collation 并在存储过程和函数中显式加 COLLATE 双重修复。')

TAG('Python CLI 开发')
P('创建了完整的命令行交互系统：main.py 入口统一管理，auth.py 处理登录认证，核心工具函数（render_menu、show_table、Paginator、confirm、hr、cls）封装至 utils.py。实现了数据驱动的菜单渲染和通用的分页器组件。角色模块按学生（student_tui.py）、教师（teacher_tui.py）、教务（admin.py）分类，tester.py 作为测试员跳板可切换任意角色身份。')

# ════════════════════════════════════════
# 7月8日
# ════════════════════════════════════════
doc.add_page_break()
H('第二天：2026年7月8日 — CLI 功能完整实现', level=1)

TAG('CRUD 完整实现')
P('完成班级、课程、教师、学生、排课、选课六大实体的完整增删改查功能。采用逐字段交互修改模式，回车保留原值；新增时自动校验必填字段和外键约束。教师管理的职称和电话字段设计为可空，学生新增时列出可选班级列表。排课新增时自动筛选有资格讲授该课程的教师列表。教务端实现强制退选功能。')

TAG('分页器推广')
P('将 Paginator 分页器应用到全部管理菜单、班级学生名单、教师列表、我的成绩、可选课程列表等十余处展示场景，解决了数据量增大后的展示问题。分页器支持上一页/下一页翻页操作，每页显示 10 条记录。')

TAG('教师端增强')
P('成绩录入流程优化为两步分页：先选择排课（显示已选人数/上限），再进入循环录入模式，录入完成后自动刷新列表。支持 CSV 文件批量导入成绩，通过循环调用 sp_grade_input 存储过程实现逐行反馈。')

TAG('Streamlit 起步')
P('开始构建 Web 界面。app.py 采用 session_state 管理用户登录状态，建立 st.navigation 原生多页导航。完成学生端 5 个页面（可选课程、选课、退选、我的成绩、学期均分）和教师端 3 个页面（录入成绩、批量 CSV 上传、查看课程学生）的开发。')

TAG('代码质量')
P('tester.py 从 175 行自建子菜单精简为 83 行跳板，直接调用真实 menu 函数消除代码重复。统一修复 import 路径问题，函数添加 paged 参数兼容 CLI 和 Streamlit 双模式调用。')

# ════════════════════════════════════════
# 7月9日
# ════════════════════════════════════════
doc.add_page_break()
H('第三天：2026年7月9日 — Streamlit 双界面搭建', level=1)

TAG('教务端 14 页面')
P('完成教务管理员的全部 14 个 Streamlit 页面：查询统计类 6 个（数据概览、班级学生名单、班级成绩统计、班级成绩明细、教师信息、教师列表），数据管理类 6 个（班级、课程、教师、学生、排课、选课 CRUD），系统工具类 2 个（备份、恢复）。')

TAG('管理页面重构')
P('经历了多轮架构迭代：st.tabs 同时渲染所有 tab 导致组件互相干扰，改为 st.radio 单模式渲染；st.form + st.form_submit_button 在动态 key 下预填值不更新且保存无效，全部改为 st.button + 动态 key + session_state 消息传递，最终方案稳定。')

TAG('备份恢复加固')
P('备份命令补全 --routines --triggers --add-drop-table 参数，确保存储过程、触发器和 DROP TABLE 一同备份。恢复改用 SET FOREIGN_KEY_CHECKS=0 包裹，避免外键冲突导致的恢复失败，不再使用 DROP DATABASE 方式。')

TAG('导航系统')
P('侧边栏从 st.radio 升级为 st.navigation 原生多页导航。22 个页面按角色分三组：学生（5 页）、教师（3 页）、教务（14 页），退出登录后自动回到登录页。')

# ════════════════════════════════════════
# 7月13日
# ════════════════════════════════════════
doc.add_page_break()
H('第四天：2026年7月13日 — 目录规范化 + 代码质量修复', level=1)

TAG('目录结构规范')
P('按照标准化项目结构对目录进行重组：规范化目录命名（SQL/ → sql、document/上交/ → submission/），清理过时的日志副本，添加 README.md 项目说明。')

TAG('代码质量修复')
P('对项目进行全面的代码质量检查与修复：修复 Pylint 全部警告，清零代码规范问题；统一算术运算符空格规范（Flake8 E226）；统一代码格式规范。')

TAG('环境配置')
P('引入 SQLAlchemy 连接池管理数据库连接，配置 QueuePool（池大小 5，溢出 10）。数据库配置迁移至 .env 文件管理，支持环境变量覆盖。')

# ════════════════════════════════════════
# 7月14日
# ════════════════════════════════════════
doc.add_page_break()
H('第五天：2026年7月14日 — 功能增强 + 图表 + 校验', level=1)

TAG('输入校验系统')
P('引入 Pydantic 模型进行全表单输入校验。创建 StudentCreate、TeacherCreate、ClassCreate、CourseCreate、GradeRecord 等校验模型，统一通过 validate_or_error 函数调用。校验失败时通过 st.toast 显示中文错误提示。所有 text_input 配置 max_chars 参数与后端模型长度限制对齐。')

TAG('Plotly 交互式图表')
P('新增 Plotly Express 图表系统。数据概览页面包含 3 个饼图（学生/班级/教师状态分布）和 2 个柱状图（各专业班级数、各年级学生数）。班级学生名单页面增加学籍分和绩点分分布饼图。班级成绩统计增加成绩分布柱状图。教师查看课程学生页面增加成绩分布饼图与排行柱状图（左右并排）。学生我的成绩页面增加各科成绩柱状图（含 60 分及格线）。教师列表增加排课数和选课学生数排行柱状图。')

TAG('专业管理')
P('新增专业管理模块（CSV 存储方式）。data/majors.csv 包含 160+ 个专业，涵盖工、理、医、文、法、经、管、艺、体等学科门类。提供前端图形化增删操作，删除时自动检查是否有班级引用。班级管理联动专业下拉选择。')

TAG('班级管理优化')
P('年级和专业改为 selectbox 下拉选择，系统自动查询该年级+专业已有班级序号，自动生成班级名（格式：{年级}{专业}{序号}班）。修改模式中班级名/年级/专业只读显示，仅允许修改状态。')

TAG('登录体验优化')
P('登录页面新增快速选择下拉框，账号从数据库实时读取。输入校验区分三种格式：教务（admin）、教师工号（T 开头+数字）、学生学号（纯数字 8-12 位）。')

TAG('排课管理优化')
P('排课上限从 number_input 改为 text_input + max_chars=5，彻底解决 JavaScript 精度问题。排课时间选择从 date_input + 两个 selectbox 升级为 st.datetime_input 原生组件，砍掉 108 行冗余代码。')

TAG('连接池加固')
P('_make_page 包装器在异常时执行 conn.rollback() 后归还连接，防止未决事务污染连接池。CSV 文件读取统一使用 utf-8-sig 编码兼容 Windows Excel BOM 头。')

# ════════════════════════════════════════
# 7月15日
# ════════════════════════════════════════
doc.add_page_break()
H('第六天：2026年7月15日 — 大数据量支持', level=1)

TAG('大数据生成器开发')
P('创建 test/generate_large_data.py 一站式生成器，从基础数据（1000 学生/25 教师/54 排课）扩增至 20,000 学生、300 教师、329 班级、1,004 排课和 2,002,688 条选课记录。采用先删除触发器、再 LOAD DATA INFILE 极速导入、最后手动批量重算 GPA 的策略，整体用时约 30 秒。')

TAG('性能调优')
P('选课管理改为按学期+学号精准查询，避免 200 万条全量加载导致的页面卡顿。退选改为输入选课 ID 操作，不依赖下拉框。增加学期筛选过滤，默认只显示当前学期数据。')

TAG('班级/教师/专业扩展')
P('专业列表扩展至 160+；生成器从前 50 个专业创建 329 个班级；教师扩展至 300 人（含教授/副教授/讲师/助教等职称）；学期跨度扩展至 20 年（40 个学期）。')

TAG('数据校验与修复')
P('修复多处数据相关 Bug：生成器姓名库扩充防止死循环、唯一键冲突去重、空成绩用 \\N 而非空字符串、高斯分布截断防越界、CROSS JOIN 从 3 次减为 2 次避免 540 亿中间结果、排课容量调大防触发器拦截。')

# ════════════════════════════════════════
# 7月16日
# ════════════════════════════════════════
doc.add_page_break()
H('第七天：2026年7月16日 — 文档收尾 + 答辩准备', level=1)

TAG('答辩 PPT 重构')
P('使用 python-pptx 重构答辩 PPT 为 11 页精美排版，包含项目概述、系统架构、数据库设计、功能模块、特色功能、数据规模与性能优化、AI 辅助开发记录、总结等完整章节。采用深蓝主题配色、卡片式布局和图标装饰。')

TAG('课程设计说明书重构')
P('使用 python-docx 重构说明书为 9 章完整内容，更新全部数据规模（20000 学生/300 教师/200 万选课），补充新增功能描述（Plotly 图表、Pydantic 校验、专业管理、自动班级名）和性能优化方案。')

TAG('代码附录生成')
P('创建代码附录文档，整理全部 23 个 Python 和 SQL 源文件，采用等宽字体 + 灰色底纹排版。')

TAG('修复审查问题')
P('修复班级删除时学生引用问题（检查 COUNT 后在删除）、学生列表过滤已删除班级的学生（student_full_list 增加 AND c.is_deleted=0）、离职教师不能被排课（teacher_course_teachers 增加 AND t.status=1）。')

TAG('版本管理与打包')
P('全部 80+ 次提交通过 GitHub 进行版本管理。最终打包源码压缩包（57 个文件，4.2 MB），确保项目结构清晰可交付。')

# ── 保存 ──
doc.save(OUT)
sz = os.path.getsize(OUT) // 1024
print(f'OK: {sz} KB -> {OUT}')
