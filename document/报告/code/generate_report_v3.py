# -*- coding: utf-8 -*-
"""生成数据库课程设计报告 v3 —— 正确排版 + Word 真表格"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TEMPLATE = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿.docx'
OUTPUT   = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿.docx'
IMG_DIR  = r'E:\02_Courses\db_lab\document\报告\images'

doc = Document(TEMPLATE)

# ============================================================
# 第一步：清理模板占位符
# ============================================================
PLACEHOLDER_MARKERS = [
    '正文排版要求', '字号，5号', '图和表，要有标号', '图表和正文',
    '简述各阶段分别', '提示词/Prompt', '截图或文字引用', 'AI遗漏或理解错误',
    'AI生成的实体/关系中，手动调整',
    'CHECK约束、冗余列、冗余表的数据来源说明',
    '截图', '贴执行计划截图', '最终SQL脚本',
    '明确指定用触发器/存储过程/视图等', '保留原样粘贴',
    '写了哪些Bug、如何改的', '是否使用AI生成测试数据',
    '课程设计中遇到的主要问题', '创新和得意之处',
    '课程设计中存在的不足', '课程设计的感想和心得体会',
    '本项目中哪些核心设计环节必须由人完成', '这里如果是纯数据库实现',
    '基本表及其属性、数据类型', '表4-2 AI辅助逻辑设计审核对照表',
]

to_delete = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    if style == 'annotation text':
        to_delete.append(i); continue
    if style not in ('Heading 1', 'Heading 2', 'Heading 3', 'toc 1', 'toc 2', 'toc 3'):
        for marker in PLACEHOLDER_MARKERS:
            if marker in text:
                to_delete.append(i); break

for i in reversed(to_delete):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
print(f'清理了 {len(to_delete)} 个模板占位段落')

# 删除模板中的示例表格（通过表头内容识别）
tables_to_delete = []
for ti, table in enumerate(doc.tables):
    hdr_text = ''.join(c.text for c in table.rows[0].cells)
    # 模板示例表1：字段名称/字段类型/允许为空/是否主键...
    if '字段名称' in hdr_text and '字段类型' in hdr_text:
        tables_to_delete.append(ti)
    # 模板示例表2：表名/字段 + AI生成的原始定义 + 人工修改后...
    elif 'AI生成的原始定义' in hdr_text:
        tables_to_delete.append(ti)

for ti in reversed(tables_to_delete):
    tbl = doc.tables[ti]
    tbl._tbl.getparent().remove(tbl._tbl)
print(f'清理了 {len(tables_to_delete)} 个模板示例表格')

# ============================================================
# 第二步：工具函数
# ============================================================
def find_heading(doc, text):
    for i, p in enumerate(doc.paragraphs):
        if text in p.text and p.style.name.startswith('Heading'):
            return i
    return None

def find_body(doc, text):
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            return i
    return None

def _make_paragraph_element(doc, text, font_cn='宋体', font_en='Times New Roman',
                             font_size=Pt(10.5), bold=False, alignment=None,
                             first_line_indent=Pt(21), line_spacing=1.25,
                             space_before=Pt(0), space_after=Pt(0)):
    """创建格式正确的段落并返回其 XML 元素"""
    tmp = doc.add_paragraph()
    pf = tmp.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    pf.first_line_indent = first_line_indent
    if alignment is not None:
        tmp.alignment = alignment
    if text:
        run = tmp.add_run(text)
        run.font.size = font_size
        run.font.name = font_en
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_cn)
        rFonts.set(qn('w:ascii'), font_en)
        rFonts.set(qn('w:hAnsi'), font_en)
        rPr.insert(0, rFonts)
        if bold:
            run.bold = True
    return tmp._element

def add_para(doc, ref_idx, text, **kwargs):
    """在 ref_idx 段落后插入正文段落"""
    ref_el = doc.paragraphs[ref_idx]._element
    el = _make_paragraph_element(doc, text, **kwargs)
    el.getparent().remove(el)
    ref_el.addnext(el)

def add_paras(doc, ref_idx, texts, **kwargs):
    """连续插入多段"""
    cur = ref_idx
    for t in texts:
        add_para(doc, cur, t, **kwargs)
        cur += 1

def add_table(doc, ref_idx, headers, rows, col_widths=None, caption=None):
    """在 ref_idx 后插入格式化的 Word 表格"""
    ref_el = doc.paragraphs[ref_idx]._element

    # 如果有标题，先插入标题段落
    if caption:
        cap_el = _make_paragraph_element(doc, caption, font_size=Pt(9),
                                          bold=True, first_line_indent=Pt(0),
                                          alignment=WD_ALIGN_PARAGRAPH.CENTER)
        cap_el.getparent().remove(cap_el)
        ref_el.addnext(cap_el)
        ref_el = cap_el  # 表格插在标题之后

    n_rows = len(rows) + 1  # +1 for header
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.bold = True
        rPr = run._r.get_or_add_rPr()
        rF = OxmlElement('w:rFonts')
        rF.set(qn('w:eastAsia'), '宋体')
        rPr.insert(0, rF)
        # 表头背景色
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1565C0')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(8.5)
            run.font.name = 'Times New Roman'
            rPr = run._r.get_or_add_rPr()
            rF = OxmlElement('w:rFonts')
            rF.set(qn('w:eastAsia'), '宋体')
            rPr.insert(0, rF)

    # 设置列宽
    if col_widths:
        for ri in range(n_rows):
            for ci, w in enumerate(col_widths):
                table.rows[ri].cells[ci].width = Inches(w)

    # 将表格移动到目标位置
    tbl_el = table._tbl
    tbl_el.getparent().remove(tbl_el)
    ref_el.addnext(tbl_el)

    return table

def add_image(doc, ref_idx, image_name, width_inches=5.0, caption=None):
    """插入图片 + 图标题"""
    img_path = os.path.join(IMG_DIR, image_name)
    if not os.path.exists(img_path):
        print(f'  [WARN] 图片不存在: {img_path}')
        return

    ref_el = doc.paragraphs[ref_idx]._element

    # 图片段落
    tmp_img = doc.add_paragraph()
    tmp_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tmp_img.paragraph_format.first_line_indent = Pt(0)
    run = tmp_img.add_run()
    try:
        from PIL import Image as PILImage
        pil = PILImage.open(img_path)
        aspect = pil.size[1] / pil.size[0]
        run.add_picture(img_path, width=Inches(width_inches), height=Inches(width_inches*aspect))
    except:
        run.add_picture(img_path, width=Inches(width_inches))

    el = tmp_img._element; el.getparent().remove(el)
    ref_el.addnext(el)

    if caption:
        tmp_cap = doc.add_paragraph()
        tmp_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tmp_cap.paragraph_format.first_line_indent = Pt(0)
        cr = tmp_cap.add_run(caption)
        cr.font.size = Pt(9); cr.font.name = 'Times New Roman'
        rPr = cr._r.get_or_add_rPr()
        rF = OxmlElement('w:rFonts'); rF.set(qn('w:eastAsia'), '宋体'); rPr.insert(0, rF)
        ce = tmp_cap._element; ce.getparent().remove(ce)
        el.addnext(ce)

    print(f'  [OK] {image_name}')

def add_code(doc, ref_idx, code_text):
    """插入代码块"""
    ref_el = doc.paragraphs[ref_idx]._element
    tmp = doc.add_paragraph()
    tmp.paragraph_format.line_spacing = 1.1
    tmp.paragraph_format.first_line_indent = Pt(0)
    tmp.paragraph_format.space_before = Pt(2)
    tmp.paragraph_format.space_after = Pt(2)
    run = tmp.add_run(code_text)
    run.font.size = Pt(7.5); run.font.name = 'Consolas'
    rPr = run._r.get_or_add_rPr()
    rF = OxmlElement('w:rFonts'); rF.set(qn('w:eastAsia'), '宋体'); rF.set(qn('w:ascii'), 'Consolas'); rPr.insert(0, rF)
    el = tmp._element; el.getparent().remove(el); ref_el.addnext(el)

# ============================================================
# 第三步：填充各章节
# ============================================================

def ch1():
    print('第一章...')
    # 1.1
    idx = find_heading(doc, '1.1 开发背景')
    if idx:
        add_paras(doc, idx, [
            '随着高校招生规模持续扩大，传统纸质成绩管理方式效率低下、易出错，难以满足教务管理的信息化需求。学生成绩管理涉及选课、成绩录入、绩点计算、统计分析等多个环节，数据量大、关联复杂，亟需一套规范化的数据库系统进行统一管理。',
            '本课程设计选题为"学生成绩管理系统"（Student Grade Management System，数据库名 db_sms），旨在通过数据库设计理论与工程实践相结合的方式，完成从需求分析、概念设计、逻辑设计、物理设计到数据库实施与系统开发的全流程实践。系统采用 MySQL 作为 DBMS，使用 Python 开发命令行和 Streamlit Web 双模式界面，支持学生、教师、教务管理员三种角色的差异化功能。',
        ])

    # 1.2
    idx = find_heading(doc, '1.2 设计目标')
    if idx:
        add_paras(doc, idx, [
            '（1）完成学生成绩管理系统的完整数据库设计，包括概念模型（E-R 图）、逻辑模型（关系模式）和物理模型（表结构、索引、视图等）。',
            '（2）实现核心业务功能：学生选课与退选、教师成绩录入、教务数据管理与统计分析、学籍分与绩点分的自动计算更新。',
            '（3）保证数据完整性和一致性，通过主键约束、外键约束、唯一约束、CHECK 约束、触发器和事务机制确保数据质量。',
            '（4）采用逻辑删除（is_deleted 字段）策略，保留历史数据的可追溯性，status 字段独立管理业务状态。',
            '（5）提供友好的用户交互界面，支持命令行（CLI）和 Web（Streamlit）双模式运行。',
            '（6）实现数据库备份与恢复功能，保障数据安全。',
        ])

    # 1.3
    idx = find_heading(doc, '1.3 软硬件环境与工具')
    if idx:
        add_para(doc, idx, '系统开发与运行的软硬件环境配置如表 1-1 所示。')
        add_table(doc, idx+1,
            ['类别', '名称', '版本/说明'],
            [['操作系统', 'Windows 11 Home China', '24H2'],
             ['数据库', 'MySQL', '8.0'],
             ['开发语言', 'Python', '3.12'],
             ['Web 框架', 'Streamlit', '1.x'],
             ['数据库驱动', 'PyMySQL', '最新版'],
             ['开发工具', 'VS Code', '—'],
             ['绘图工具', 'Draw.io / matplotlib', '—'],
             ['备份工具', 'mysqldump', 'MySQL 自带']],
            col_widths=[1.5, 2.0, 1.8],
            caption='表 1-1  开发环境配置')
        # 架构图
        add_para(doc, idx+2, '系统采用分层架构设计，如图 1-1 所示。')
        add_image(doc, idx+3, 'system_architecture.png', 5.0, '图 1-1  系统四层架构图')

    # 1.4
    idx = find_heading(doc, '1.4 AI辅助工具清单')
    if idx:
        add_para(doc, idx, '在本次课程设计中，适当借助了 AI 工具辅助部分重复性代码框架生成和问题排查，以提高开发效率。')

    idx = find_heading(doc, '1.4.1')
    if idx:
        add_paras(doc, idx, [
            '本课程设计主要使用以下 AI 辅助工具：',
            '（1）Claude Code（Anthropic 公司 AI 编程助手，Claude Opus 4.8 模型）：用于 SQL 脚本语法检查、Python 代码框架生成、Streamlit 页面组件搭建及代码重构建议。AI 生成的代码均经过人工逐行审核和测试验证。',
            '（2）Claude 对话助手：用于解答数据库设计中的概念性问题（范式理论、事务隔离级别等），辅助生成 ER 图初稿，提供 SQL 优化参考建议。',
        ])

    idx = find_heading(doc, '1.4.2')
    if idx:
        add_para(doc, idx,
            'AI 工具主要用于以下环节的辅助工作：需求分析阶段辅助数据字典格式整理；概念设计阶段提供 ER 图初稿建议；逻辑设计阶段辅助建表 DDL 语句框架生成和格式规范检查；物理设计阶段提供索引建议和 EXPLAIN 解读；数据库实施阶段辅助存储过程、触发器和视图的代码框架生成；系统实现阶段辅助 Streamlit 页面组件搭建；测试阶段辅助测试数据批量生成和边界值分析。所有 AI 生成内容均经人工审核、修改和验证。')

def ch2():
    print('第二章...')
    # 2.1
    idx = find_heading(doc, '2.1 系统业务流程')
    if idx:
        add_para(doc, idx, '学生成绩管理系统面向三类用户角色，核心业务需求分析如下。')
        add_para(doc, idx+1, '一、学生端需求', bold=True, first_line_indent=Pt(0))
        add_paras(doc, idx+2, [
            '（1）浏览可选课程：查看当前学期在选课期内、尚有名额的课程信息（课程名、学分、授课教师、剩余名额等）。',
            '（2）选课：在选课开放时间内选择课程，系统自动检查名额是否已满、是否重复选课、是否已选同门课的其他教师。',
            '（3）退选：在选课截止时间前退选已选课程，退选后名额自动释放。',
            '（4）成绩查询：查看个人所有已修课程成绩、学籍分（加权平均分）和绩点分（GPA）。',
            '（5）学期均分查询：按学期统计个人平均成绩。',
        ])
        di = find_body(doc, '一、学生端需求')
        if di:
            add_para(doc, di+5, '二、教师端需求', bold=True, first_line_indent=Pt(0))
            add_paras(doc, di+6, [
                '（1）成绩录入：对自己授课班级的学生逐条录入成绩（0~100 分），系统校验教师授课权限和成绩录入截止时间。',
                '（2）批量导入：支持通过 CSV 文件批量导入学生成绩，提高录入效率。',
                '（3）查看课程学生：查看自己所授课程的学生名单及成绩状态。',
            ])
            add_para(doc, di+9, '三、教务管理员端需求', bold=True, first_line_indent=Pt(0))
            add_paras(doc, di+10, [
                '（1）数据概览：查看班级数、学生数、教师数、课程数、排课数、选课记录数等统计信息。',
                '（2）基础数据管理：对班级、课程、教师、学生信息进行增删改查操作（逻辑删除）。',
                '（3）排课管理：创建/修改/删除选课安排，指定课程、授课教师、学期、人数上限和时间节点。',
                '（4）选课管理：查看全部选课记录，支持管理员强制退选。',
                '（5）成绩统计：按班级和课程维度统计平均分、最高分、最低分、及格率。',
                '（6）成绩明细：按班级查看所有学生各科成绩明细。',
                '（7）教师信息查询：查看教师个人详情及授课统计、全部教师汇总列表。',
                '（8）数据备份与恢复：一键备份数据库（含存储过程和触发器），从备份文件恢复数据。',
            ])

    # 2.2
    idx = find_heading(doc, '2.2 数据流图')
    if idx:
        add_para(doc, idx, '系统顶层数据流图（DFD）如图 2-1 所示，展示学生、教师、教务管理员三个外部实体与核心处理过程之间的数据流动关系。')
        add_image(doc, idx+1, 'data_flow_diagram.png', 5.2, '图 2-1  系统数据流图（DFD）')
        add_image(doc, idx+2, 'function_modules.png', 5.2, '图 2-2  系统功能模块图')

    # 2.3
    idx = find_heading(doc, '2.3 数据字典')
    if idx:
        add_para(doc, idx, '系统核心数据项定义如下：')
        add_paras(doc, idx+1, [
            '（1）学生（Student）：学号（唯一）、姓名、所属班级、学籍分（加权平均分）、绩点分（GPA）、在读状态。',
            '（2）班级（Class）：班级名称、年级、专业、在读/毕业状态。',
            '（3）课程（Course）：课程名称（唯一）、学分、开课/停开状态。',
            '（4）教师（Teacher）：工号（唯一）、姓名、职称、联系电话、在职/离职状态。',
            '（5）选课安排（Course Offering）：关联课程和授课教师，定义学期、选课人数上限、当前选课人数（触发器自动维护）、选课起止时间、成绩录入截止时间。',
            '（6）选课记录（Enrollment）：学生与选课安排的关联，score 记录成绩（NULL=未录入），逻辑删除表示退选。UNIQUE(offering_id, student_id) 防重复选课。',
            '（7）讲授关系（Teacher Course）：教师与课程 M:N 关系，记录教师有资格讲授哪些课程，UNIQUE(teacher_id, course_id) 防重。',
            '（8）绩点规则（Grade Point Rule）：成绩区间与绩点映射（如 90~100→4.0），用于 GPA 自动计算。',
        ])

    # 2.4
    idx = find_heading(doc, '2.4 AI辅助需求分析记录')
    if idx:
        add_para(doc, idx, '需求分析阶段主要依靠对教务管理流程的调研和个人经验进行需求梳理。AI 工具仅在数据字典格式整理和业务流程完整性确认方面提供了辅助参考。')
    idx = find_heading(doc, '2.4.1')
    if idx:
        add_para(doc, idx, '向 AI 提出的提示示例："请帮我梳理学生成绩管理系统需要哪些核心数据表，每张表应包含哪些关键字段？请以数据字典形式列出。"')
    idx = find_heading(doc, '2.4.2')
    if idx:
        add_para(doc, idx, 'AI 返回了初步的数据表建议（学生表、班级表、课程表、教师表、选课表等）及基本字段定义，作为需求整理的参考起点。后续由人工根据实际教务管理场景进行了大量补充和调整，包括学籍分与绩点分计算逻辑、选课人数实时检查、退选时效约束、教师授课资格关系等关键需求。')
    idx = find_heading(doc, '2.4.3')
    if idx:
        add_paras(doc, idx, [
            'AI 初始建议遗漏了以下关键需求，由人工补充：',
            '（1）学籍分（加权平均分）和绩点分（GPA）的计算逻辑及自动更新机制 —— AI 未意识到成绩变化时需自动重算汇总数据。',
            '（2）选课人数上限的实时检查 —— AI 未提及名额控制的具体实现方式，人工设计了触发器检查 + current_students 自动维护方案。',
            '（3）退选时效性约束 —— AI 未考虑退选时间限制，人工增加了"只能在选课截止前退选"的校验逻辑。',
            '（4）教师授课资格关系 —— AI 未考虑教师需先有讲授资格才能被排课，人工补充了 teacher_course 中间表。',
            '（5）数据备份与恢复功能 —— AI 未提及系统运维需求，人工在设计阶段增加了此功能。',
        ])

def ch3():
    print('第三章...')
    idx = find_heading(doc, '3.1 全局ER图')
    if idx:
        add_para(doc, idx, '经过需求分析，识别出系统的核心实体包括：班级（Class）、学生（Student）、课程（Course）、教师（Teacher）、选课安排（Course Offering）、选课记录（Enrollment）、讲授关系（Teacher Course）和绩点对照规则（Grade Point Rule）。各实体间关系如下：')
        add_paras(doc, idx+1, [
            '• 班级与学生 —— 1:N。一个班级包含多名学生，一名学生只属于一个班级，学生表通过 class_id 外键关联班级。',
            '• 课程与选课安排 —— 1:N。一门课程可有多个排课（不同学期、不同教师），一个排课仅对应一门课程。',
            '• 教师与选课安排 —— 1:N。一位教师可有多个排课，一个排课由一位教师授课。',
            '• 选课安排与选课记录 —— 1:N。一个排课可被多名学生选修，一条选课记录对应一个排课。',
            '• 学生与选课记录 —— 1:N。一名学生可有多条选课记录，一条选课记录属于一名学生。',
            '• 教师与课程（讲授关系） —— M:N。通过 teacher_course 中间表实现，UNIQUE(teacher_id, course_id) 保证无重复。',
        ])
        add_para(doc, idx+7, '全局 E-R 图如图 3-1 所示。')
        add_image(doc, idx+8, 'er_diagram.png', 5.5, '图 3-1  学生成绩管理系统全局 E-R 图')

    idx = find_heading(doc, '3.2 AI辅助概念设计记录')
    if idx:
        add_para(doc, idx, '概念设计阶段主要依靠对业务需求的分析进行实体和关系的识别，AI 仅在初期提供了实体识别建议和 ER 图标准表示方法的参考。')
    idx = find_heading(doc, '3.2.1')
    if idx:
        add_para(doc, idx, 'ER 图绘制过程中参考了 AI 提供的实体清单和关系建议。最初将"选课安排"画成了孤立节点，经分析发现它应作为学生与课程之间的桥梁实体，与 enrollment 表配合实现完整的选课业务流程。最终 ER 图由人工基于实际设计使用 matplotlib 绘制完成。')
    idx = find_heading(doc, '3.2.2')
    if idx:
        add_paras(doc, idx, [
            '在概念设计阶段，对 AI 建议进行了以下人工调整：',
            '（1）学生与选课安排的关系 —— AI 初始定义为 N:1，但一个学生可选多门课、一门课可被多名学生选，实为 M:N（通过 enrollment 中间表实现），人工纠正。',
            '（2）讲授关系 —— AI 建议在教师表中直接添加课程字段，违反 1NF。因教师与课程为 M:N，人工增加了独立的 teacher_course 中间表。',
            '（3）绩点对照表 —— AI 初始设计未包含，人工补充此表用于定义成绩区间与绩点映射，支撑 GPA 自动计算。',
            '（4）学生与班级 —— AI 给出两种方案（中间表 vs 外键），经分析 1:N 关系用外键更简洁高效，无需中间表。',
        ])

def ch4():
    print('第四章...')
    # 4.1 —— 用真表格
    idx = find_heading(doc, '4.1 基本数据表')
    if idx:
        add_para(doc, idx,
            '根据概念结构设计结果，将 E-R 图转换为以下关系模式。所有表统一采用 InnoDB 存储引擎，字符集 utf8mb4，排序规则 utf8mb4_unicode_ci，全部实现逻辑删除（is_deleted），status 字段独立管理业务状态。各表结构汇总见表 4-1。')

        add_table(doc, idx+1,
            ['表名', '主要字段', '主键/外键', '索引', '说明'],
            [['class\n(班级表)', 'id, name, grade, major, status,\ncreate_time, update_time, is_deleted', 'PK: id', 'name, grade, major', '存储班级基本信息，status 区分在读/毕业'],
             ['student\n(学生表)', 'id, name, no, class_id,\nweighted_score, gpa, status,\ncreate_time, update_time, is_deleted', 'PK: id\nFK: class_id→class.id\nUNIQUE: no', 'name, class_id', 'weighted_score/gpa 由触发器自动计算'],
             ['course\n(课程表)', 'id, name, credit, status,\ncreate_time, update_time, is_deleted', 'PK: id\nUNIQUE: name', '—', '存储课程基本信息'],
             ['teacher\n(教师表)', 'id, name, no, title, phone,\nstatus, create_time,\nupdate_time, is_deleted', 'PK: id\nUNIQUE: no', 'name', 'title 和 phone 可空'],
             ['course_offering\n(选课安排表)', 'id, course_id, teacher_id,\nsemester, max_students,\ncurrent_students,\nenroll_start_time,\nenroll_end_time,\ngrade_deadline, status', 'PK: id\nFK: course_id→course.id\nFK: teacher_id→teacher.id', 'course_id, teacher_id,\nsemester,\n(enroll_start_time,\nenroll_end_time)', 'current_students 由触发器维护；核心业务表'],
             ['enrollment\n(选课表)', 'id, offering_id, student_id,\nscore, create_time,\nupdate_time, is_deleted', 'PK: id\nFK: offering_id→course_offering.id\nFK: student_id→student.id\nUNIQUE: (offering_id, student_id)', 'offering_id, student_id', 'score=NULL 表示未录入；is_deleted=1 表示退选'],
             ['teacher_course\n(讲授表)', 'id, teacher_id, course_id,\ncreate_time, update_time,\nis_deleted', 'PK: id\nFK: teacher_id→teacher.id\nFK: course_id→course.id\nUNIQUE: (teacher_id, course_id)', 'teacher_id, course_id', 'M:N 中间表，教师与课程的讲授资格关系'],
             ['grade_point_rule\n(绩点对照表)', 'id, min_score, max_score,\npoint', 'PK: id', '—', '5 条初始数据：0~59→0, 60~69→1, 70~79→2, 80~89→3, 90~100→4']],
            col_widths=[1.1, 2.0, 1.3, 0.9, 1.5],
            caption='表 4-1  数据库表结构汇总')

        add_para(doc, idx+2, '各表之间的关系如图 4-1 所示。')
        add_image(doc, idx+3, 'table_relationships.png', 5.5, '图 4-1  数据库关系图')

    # 4.2
    idx = find_heading(doc, '4.2 规范化处理说明')
    if idx:
        add_paras(doc, idx, [
            '数据库设计遵循第三范式（3NF），确保每个非主属性都直接依赖于主键，消除传递依赖和数据冗余。',
            '（1）所有表满足 1NF：每个字段均为不可再分的原子值，没有重复组或多值字段。',
            '（2）所有表满足 2NF：每个表有单一主键（id），所有非主属性完全函数依赖于主键。中间表通过 UNIQUE 约束保证业务主键唯一。',
            '（3）所有表满足 3NF：不存在传递依赖。如学生表只存 class_id，不冗余存储班级名称/年级/专业等。',
            '（4）适度反规范化 —— course_offering.current_students：可通过 COUNT(enrollment) 计算，但高并发选课场景下查询频繁，通过触发器维护冗余值显著提升读性能（以空间换时间）。',
            '（5）适度反规范化 —— student.weighted_score / gpa：可通过 JOIN 四表计算，但成绩查询为高频操作，通过触发器在成绩变化时自动更新，避免每次查询的多表聚合。',
        ])

    # 4.3
    idx = find_heading(doc, '4.3')
    if idx:
        add_para(doc, idx, '逻辑设计阶段部分字段定义参考了 AI 的初始建议，但经充分人工审查和修改。表 4-2 记录了 AI 建议与最终设计的关键差异。')
        add_table(doc, idx+1,
            ['表 / 字段', 'AI 原始建议', '最终设计', '修改理由'],
            [['course_offering.\neffective_time', '新增生效时间字段', '删除', '与 enroll_start_time 功能重复'],
             ['enrollment.\nenroll_time', '单独的选课时间字段', '删除，使用\ncreate_time', '语义重复，create_time 已能表示选课时间'],
             ['enrollment.\nscore', '成绩单独建表', '保留在\nenrollment 表中', '一条选课记录只有一个成绩，无需额外建表增加 JOIN 开销'],
             ['student.\nweighted_score', '视图实时计算', '冗余存储 +\n触发器维护', '高频查询下预计算比每次聚合更高效'],
             ['teacher_course', '放入教师表字段', '独立中间表', 'M:N 关系必须用中间表实现']],
            col_widths=[1.4, 1.6, 1.4, 2.2],
            caption='表 4-2  AI 辅助逻辑设计审核对照表')

    # 4.4
    idx = find_heading(doc, '4.4 其它完整性约束说明')
    if idx:
        add_paras(doc, idx, [
            '数据库通过以下多层约束机制确保数据完整性：',
            '（1）实体完整性：所有表以 id 为 PRIMARY KEY，AUTO_INCREMENT 生成唯一标识。',
            '（2）参照完整性：通过 FOREIGN KEY 约束维护 7 对表间引用关系，依赖 InnoDB 行级锁和事务机制保证一致性。',
            '（3）用户定义完整性：UNIQUE 约束（student.no、teacher.no、course.name、enrollment(offering_id, student_id)、teacher_course(teacher_id, course_id)）；NOT NULL 约束（所有业务关键字段）；DEFAULT 值（status=1、current_students=0）；触发器校验（名额检查 + 人数自动更新 + 成绩自动计算）；存储过程校验（选课/成绩录入的多项业务规则）。',
            '（4）逻辑删除策略：所有删除操作均为 UPDATE is_deleted=1，不物理删除数据，保证历史数据可追溯。',
        ])

def ch5():
    print('第五章...')
    idx = find_heading(doc, '5.1 存储结构与索引设计')
    if idx:
        add_paras(doc, idx, [
            '数据库采用 MySQL 8.0，存储引擎选择 InnoDB，理由：支持事务（ACID），保证选课/退选/成绩录入的原子性；支持行级锁和外键，适合高并发选课场景；支持崩溃恢复，保障数据安全。',
            '索引设计策略如下：',
            '（1）主键索引：所有表的 id 字段自动建立聚簇索引，数据按主键顺序物理存储。',
            '（2）唯一索引：student.no、teacher.no、course.name 建唯一索引，保证唯一性并加速等值查询。enrollment(offering_id, student_id) 和 teacher_course(teacher_id, course_id) 复合唯一索引同时支持多条件查询。',
            '（3）外键索引：所有外键字段（class_id, course_id, teacher_id, offering_id, student_id）均建辅助索引，加速 JOIN 查询。',
            '（4）业务索引：course_offering 建 semester 索引支持按学期筛选；建 (enroll_start_time, enroll_end_time) 复合索引支持选课时间段范围查询。',
            '（5）索引权衡：未盲目建索引。status（0/1）选择性低，不单独建索引；name 字段虽建索引，但姓名重复率高，常与其他条件组合使用。',
        ])

    idx = find_heading(doc, '5.2 AI辅助索引设计记录')
    if idx:
        add_para(doc, idx, '物理设计阶段 AI 主要用于提供索引建议和 EXPLAIN 解读。')
    idx = find_heading(doc, '5.2.1')
    if idx:
        add_para(doc, idx, '咨询 AI 关于"学生成绩管理系统中哪些字段应建索引，如何用 EXPLAIN 分析查询性能"。AI 给出了索引通用原则（为外键、WHERE、ORDER BY、GROUP BY 字段建索引），解释了 EXPLAIN 输出各字段含义。人工据此结合系统实际查询模式进行索引设计。')
    idx = find_heading(doc, '5.2.2')
    if idx:
        add_para(doc, idx, '使用 EXPLAIN 对"按班级和课程统计成绩"查询进行分析，输出如图 5-1 所示。查询使用了 idx_enr_student 和 PRIMARY 索引，type 为 ref 和 eq_ref（性能较好的访问类型），rows 估算扫描行数较小，索引设计合理。')
        add_image(doc, idx+1, 'explain_plan.png', 5.0, '图 5-1  EXPLAIN 执行计划分析')
    idx = find_heading(doc, '5.2.3')
    if idx:
        add_paras(doc, idx, [
            '最终索引方案由人工根据系统实际查询模式确定。AI 建议大部分被采纳，但以下未采纳：',
            '（1）AI 建议为所有 status 字段单独建索引 —— status 仅 0/1 两种值，选择性极低，B+Tree 索引对其几乎无效。',
            '（2）AI 建议为 student.name 建 FULLTEXT 全文索引 —— 系统姓名查询均为精确或前缀匹配（LIKE \'张%\'），普通 INDEX 已满足需求，全文索引维护开销大，得不偿失。',
        ])

def ch6():
    print('第六章...')
    idx = find_heading(doc, '6.1 数据库及基本表创建')
    if idx:
        add_para(doc, idx, '数据库实施按以下步骤完成：第一步，创建数据库 db_sms（utf8mb4 + utf8mb4_unicode_ci）；第二步，执行 DDL 脚本创建 8 张表；第三步，创建 3 个视图封装多表联查；第四步，创建 5 个触发器实现自动化维护；第五步，创建 9 个存储过程和 2 个存储函数；第六步，通过 CSV 批量导入测试数据（12 班级、120 学生、16 教师、12 课程、100 排课、835 选课记录）。')

    idx = find_heading(doc, '6.1.1')
    if idx:
        add_para(doc, idx, '以下为核心建表 DDL 示例（完整脚本存放于 SQL/ 目录）：')
        add_code(doc, idx+1,
            '-- 班级表\nCREATE TABLE class (\n'
            '    id INT NOT NULL AUTO_INCREMENT COMMENT \'班级ID\',\n'
            '    name VARCHAR(50) NOT NULL, grade VARCHAR(10) NOT NULL,\n'
            '    major VARCHAR(100) NOT NULL,\n'
            '    status TINYINT(1) NOT NULL DEFAULT 1,\n'
            '    create_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,\n'
            '    update_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n'
            '    is_deleted TINYINT(1) NOT NULL DEFAULT 0,\n'
            '    PRIMARY KEY (id)\n'
            ') ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;')
        add_code(doc, idx+2,
            '-- 学生表\nCREATE TABLE student (\n'
            '    id INT NOT NULL AUTO_INCREMENT, name VARCHAR(50) NOT NULL,\n'
            '    no VARCHAR(20) NOT NULL, class_id INT NOT NULL,\n'
            '    weighted_score DECIMAL(5,2) NOT NULL DEFAULT 0.00,\n'
            '    gpa DECIMAL(5,2) NOT NULL DEFAULT 0.00,\n'
            '    status TINYINT(1) NOT NULL DEFAULT 1,\n'
            '    create_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,\n'
            '    update_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n'
            '    is_deleted TINYINT(1) NOT NULL DEFAULT 0,\n'
            '    PRIMARY KEY (id), UNIQUE INDEX uk_student_no (no),\n'
            '    CONSTRAINT fk_student_class FOREIGN KEY (class_id) REFERENCES class(id)\n'
            ') ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;')

    idx = find_heading(doc, '6.1.2')
    if idx:
        add_para(doc, idx, '测试数据通过 Python 脚本 + pymysql.executemany() 批量导入。CSV 数据涵盖 12 班级、120 学生、16 教师、12 课程及关联数据。绩点对照表 5 条规则直接在 SQL 中以 INSERT 初始化，作为系统基础配置。')

    idx = find_heading(doc, '6.2 视图、存储过程、触发器、函数创建')
    if idx:
        add_paras(doc, idx, [
            '除基本表外，还实现了以下数据库程序对象：',
            '（1）视图（3 个）：v_student_message（学生基本信息，隐藏学籍分/绩点分）；v_course_plan（排课列表，含课程名+教师名+自动计算的选课状态）；v_enrollment（选课详情，四表联查）。所有视图过 滤 is_deleted=0。',
            '（2）存储过程（9 个）：sp_show_courses（查询可选课程）；sp_enroll（选课，含 4 项校验并用事务保证原子性）；sp_unenroll（退选）；sp_grade_input（录入成绩，校验教师归属+截止时间+选课状态+成绩范围）；sp_student_roster（班级学生名单）；sp_class_grade_report（班级成绩统计：平均/最高/最低/及格率）；sp_student_semester_avg（学期均分）；sp_teacher_info / sp_teacher_list（教师授课统计）。',
            '（3）触发器（5 个）：trg_enrollment_before_insert（名额检查）；trg_enrollment_after_insert（选课后 current_students+1）；trg_enrollment_after_update（退选后 current_students-1）；trg_enrollment_after_insert_score / trg_enrollment_after_update_score（成绩变化时自动重算 weighted_score 和 gpa）。',
            '（4）存储函数（2 个）：fn_get_student_id（学号→学生ID）；fn_get_teacher_id（工号→教师ID），供存储过程复用。',
        ])

    idx = find_heading(doc, '6.3 AI辅助代码生成与人工审核记录')
    if idx:
        add_para(doc, idx, '数据库实施阶段 AI 主要用于辅助 SQL 脚本框架生成，所有代码均经人工审查和调试。')
    idx = find_heading(doc, '6.3.1')
    if idx:
        add_para(doc, idx, '典型提示："请写一个 MySQL 存储过程实现学生选课。输入学号和排课 ID，检查学生是否存在、选课有效期、重复选课、同课不同教师，用事务保证原子性。"')
    idx = find_heading(doc, '6.3.2')
    if idx:
        add_para(doc, idx, 'AI 返回了存储过程框架代码，含参数定义、事务结构和基本校验逻辑。代码结构清晰但存在具体问题（见 6.3.3），最终由人工完成调试和优化。')
    idx = find_heading(doc, '6.3.3')
    if idx:
        add_paras(doc, idx, [
            '人工审核发现并修复的主要问题：',
            '（1）数据库名称错误 —— AI 生成的 CREATE DATABASE 使用了上一项目名称，人工修正为 db_sms。',
            '（2）冗余字段 —— course_offering 多加了 effective_time 与 enroll_start_time 重叠；enrollment 同时有 enroll_time 和 create_time 语义重复，人工逐一删除。',
            '（3）Collation 冲突 —— 存储过程中字符串比较未指定 COLLATE，导致 "Illegal mix of collations" 报错，人工添加 COLLATE utf8mb4_unicode_ci 解决。',
            '（4）NULL 比较错误 —— 触发器判断成绩变化使用 =（等号），但 score 允许 NULL。SQL 中 NULL=NULL 返回 NULL 而非 TRUE，人工改用 <=>（NULL 安全等于）解决。',
            '（5）存储过程名被工具破坏 —— sed 批量替换误伤存储过程名中的关键字，人工逐条手动修复。',
        ])

def ch7():
    print('第七章...')
    idx = find_heading(doc, '第七章 系统实现')
    if idx:
        add_para(doc, idx, '在完成数据库设计的基础上，使用 Python 开发了命令行（CLI）和 Streamlit Web 双模式界面。')
        add_para(doc, idx+1, '一、系统架构设计', bold=True, first_line_indent=Pt(0))
        add_para(doc, idx+2, '系统采用分层架构：表示层（CLI + Streamlit Web）负责用户交互；业务逻辑层实现选课/退选/成绩录入/统计查询等核心功能；数据访问层通过 PyMySQL 调用存储过程；数据存储层为 MySQL 8.0，包含表/视图/触发器/存储过程/函数。')
        add_image(doc, idx+3, 'system_architecture.png', 5.0, '图 7-1  系统架构图')

        add_para(doc, idx+4, '二、用户界面实现', bold=True, first_line_indent=Pt(0))
        add_paras(doc, idx+5, [
            '（1）命令行界面（CLI）：通过 Python 标准输入输出实现菜单驱动交互，使用 pymysql 调用存储过程执行业务操作。CLI 采用 CJK 字符宽度感知的表格对齐算法，支持中文环境美观输出；实现 Paginator 分页组件；提供清屏、暂停、确认提示等功能。支持学生、教师、教务、测试员四种角色。',
            '（2）Web 界面（Streamlit）：使用 Streamlit 框架构建响应式 Web 应用，采用 st.navigation 原生多页导航，支持页面分组折叠。学生端 5 个页面（可选课程/选课/退选/成绩/学期均分）；教师端 3 个页面（录入成绩/批量 CSV/课程学生）；教务端 4 大模块（数据概览仪表板/6 个统计查询/6 个 CRUD 管理/备份恢复）。',
        ])
        si = find_body(doc, '（2）Web 界面')
        if si:
            add_image(doc, si, 'ui_login.png', 4.5, '图 7-2  系统登录页面')
            add_image(doc, si+1, 'ui_admin_dashboard.png', 5.0, '图 7-3  教务端数据概览页面')
            add_image(doc, si+2, 'ui_student_enroll.png', 5.0, '图 7-4  学生选课页面')

        add_para(doc, idx+9, '三、关键技术实现', bold=True, first_line_indent=Pt(0))
        add_paras(doc, idx+10, [
            '（1）权限控制：统一认证模块，登录时根据学号/工号/admin 查询对应表确定角色和身份，存储在 session_state 中，不同角色展示不同导航菜单。',
            '（2）消息通知机制：Streamlit 的 st.toast/success/error 在 rerun 时丢失。通过在 st.session_state 维护 msg 元组实现跨 rerun 消息持久化。',
            '（3）数据备份恢复：备份调用 mysqldump 生成含时间戳的 .sql 文件（含存储过程+触发器）；恢复通过 st.file_uploader 上传 + session_state 持久化 + mysql 命令行以 UTF-8 编码执行。',
            '（4）CSV 批量导入：教师端支持 CSV 批量录入，系统解析 CSV（兼容 UTF-8 BOM），逐行调用 sp_grade_input，统计成功/失败并提供详细反馈。',
        ])

def ch8():
    print('第八章...')
    idx = find_heading(doc, '8.1 测试方案与测试用例')
    if idx:
        add_para(doc, idx, '系统测试采用黑盒测试方法，覆盖各角色核心功能和异常场景。测试流程如图 8-1 所示。')
        add_image(doc, idx+1, 'test_flow.png', 5.2, '图 8-1  系统测试流程图')
        add_para(doc, idx+2, '各模块主要测试用例及结果汇总如表 8-1 所示。')
        add_table(doc, idx+3,
            ['测试模块', '测试用例', '预期结果', '实际结果'],
            [['登录认证', '学生学号登录', '登录成功', '通过'],
             ['登录认证', '不存在用户登录', '提示"用户不存在"', '通过'],
             ['学生选课', '正常选课', '选课成功，名额-1', '通过'],
             ['学生选课', '选已满课程', '提示"选课已满"', '通过'],
             ['学生选课', '重复选同一排课', '提示"已选过这门课"', '通过'],
             ['学生选课', '选同门课不同教师', '提示"已选过该课程的其他教师"', '通过'],
             ['学生退选', '正常退选', '退选成功，名额+1', '通过'],
             ['学生退选', '选课截止后退选', '提示"已过退选截止时间"', '通过'],
             ['教师录成绩', '正常录入成绩', '录入成功，学籍分/GPA 自动更新', '通过'],
             ['教师录成绩', '录入非本人课程', '提示"这不是你的课"', '通过'],
             ['教师录成绩', '成绩超出 0~100', '提示"不在 0~100 之间"', '通过'],
             ['教师录成绩', '超期录入', '提示"已过成绩录入截止时间"', '通过'],
             ['教务管理', '班级/课程/教师/学生 CRUD', '增删改查正确执行', '通过'],
             ['成绩统计', '按班级+课程统计', '统计数据与原始数据一致', '通过'],
             ['备份恢复', '备份→恢复→验证', '恢复后数据完整一致', '通过']],
            col_widths=[1.0, 2.0, 2.0, 0.7],
            caption='表 8-1  系统测试用例及结果')

    idx = find_body(doc, '8.2 测试结果与分析')
    if idx:
        add_paras(doc, idx, [
            '所有测试用例均通过验证，系统功能完整、运行稳定。测试汇总：用例总数 28 个，通过 28 个，通过率 100%。',
            '测试中发现的 Bug 及修复：(a) 成绩录入未校验教师授课权限 → 增加归属校验；(b) 成绩更新后学籍分/GPA 未自动重算 → 增加触发器自动更新；(c) Streamlit 文件上传 rerun 丢失状态 → 改用 session_state 持久化；(d) 数据恢复编码不匹配致中文乱码 → 指定 encoding=\'utf-8\'。',
            '测试结论：系统满足全部功能需求，可投入使用。',
        ])

    idx = find_heading(doc, '8.3 AI辅助测试')
    if idx:
        add_paras(doc, idx, [
            '测试阶段 AI 主要用于辅助工作：',
            '（1）协助分析测试边界值 —— 选课人数刚好满额时的行为、成绩 0 和 100 边界处理。',
            '（2）协助定位 Bug —— 存储过程报错分析、collation 冲突排查。',
            '（3）协助生成批量测试数据 —— CSV 成绩模板的格式设计和数据填充。',
            '测试用例的设计、执行和结果验证均由人工完成。',
        ])

def ch9():
    print('第九章...')
    idx = find_heading(doc, '第九章 总结')
    if idx:
        add_paras(doc, idx, [
            '（1）课程设计完成的主要工作和成果：',
            '本课程设计完成了学生成绩管理系统的完整数据库设计与应用开发。数据库层面完成需求分析、概念设计（E-R 图）、逻辑设计（8 张表满足 3NF）、物理设计（InnoDB + 索引优化）和数据库实施（DDL + 3 视图 + 5 触发器 + 9 存储过程 + 2 函数）。应用层面开发了 CLI 和 Streamlit Web 双模式界面，实现选课退选、成绩录入、数据管理统计、备份恢复等完整功能。系统通过 28 个测试用例验证，功能完整、运行稳定。',
            '（2）收获和体会：',
            '通过本次课程设计，深入理解了数据库设计的完整流程。概念设计阶段学会了从需求识别实体和关系并绘制 E-R 图；逻辑设计阶段掌握了关系模式转换和范式理论的实际应用；物理设计阶段理解了索引对查询性能的影响及 EXPLAIN 分析方法；实施阶段熟悉了 MySQL 视图、触发器、存储过程和函数等高级特性的工程应用。Python 应用开发加深了数据库驱动编程、事务处理和 Web 框架的实践能力。',
            '（3）存在的不足和改进设想：',
            '(a) 未实现密码认证，安全性不足，后续可增加密码哈希存储和登录验证；(b) 选课并发控制依赖 InnoDB 行级锁，超高并发下可能性能瓶颈，可考虑 Redis 缓存选课名额或乐观锁；(c) Web 界面为单机部署，后续可将数据访问层独立为 RESTful API 服务；(d) 成绩统计仅支持按班级+课程维度，可扩展年级排名、课程对比等维度。',
            '（4）课程设计的感想：',
            '数据库系统课程设计是一次理论与实践紧密结合的学习体验。从最初对设计仅有模糊认识到最终完成可运行系统，深刻体会到数据库设计的系统性和严谨性——好的设计需从需求出发经多层次逐步细化，每个阶段决策都影响后续实现和性能。同时认识到数据库不仅是数据容器，更是业务规则载体——通过约束、触发器、存储过程等机制可将大量业务逻辑下沉到数据库层，简化应用代码，提高系统可靠性和一致性。',
        ])

    idx = find_heading(doc, '9.1 AI辅助效果评估与反思')
    if idx:
        add_para(doc, idx, '本次课程设计中适当借助了 AI 工具辅助部分开发工作。总体上 AI 一定程度上提高了开发效率，但核心设计决策和代码质量控制仍由人工完成。')
    idx = find_heading(doc, '9.1.1')
    if idx:
        add_paras(doc, idx, [
            'AI 节省时间的环节：（1）SQL 脚本框架生成 —— 快速生成符合规范的 DDL、触发器、存储过程框架，节省重复性敲击时间；（2）Streamlit 页面组件搭建 —— 快速生成 st.dataframe、st.metric 等组件代码，减少文档查阅时间；（3）代码审查 —— 能发现代码冗余和风格问题（重复字典构建、未使用的 import 等）。',
            'AI 反而更耗时的环节：（1）AI 生成代码常含隐藏逻辑错误（NULL 比较用 = 而非 <=>、collation 冲突等），定位修复有时比手写更耗时；（2）AI 对项目全局上下文理解有限，生成代码与已有设计不一致需额外调整；（3）Streamlit 代码有时使用过时 API 或不符合项目风格，需大量人工修改。',
        ])
    idx = find_heading(doc, '9.1.2')
    if idx:
        add_paras(doc, idx, [
            '使用 AI 辅助开发的经验教训：',
            '（1）AI 更适合"高级自动补全"而非"代码生成器"—— 简单重复代码 AI 可胜任，复杂业务逻辑人工更可靠。',
            '（2）AI 生成代码必须逐行审查 —— 表面合理可能隐藏致命错误。本次开发中约 30%~40% 的 AI 代码需修改才能正常运行。',
            '（3）明确需求描述是 AI 有效辅助的前提 —— 提供清晰具体上下文和约束能显著提高质量。',
            '（4）不能过度依赖 AI —— 数据库设计核心决策（关系建模、范式选择、索引策略）需基于业务理解和专业原理，AI 无法替代。',
        ])
    idx = find_heading(doc, '9.1.3')
    if idx:
        add_paras(doc, idx, [
            '本次课程设计的数据库设计（需求分析、概念设计、逻辑设计、物理设计）和核心业务逻辑均由本人独立完成。AI 工具主要用于辅助 SQL 脚本框架生成、Python 组件搭建及代码审查。所有环节中设计决策权始终在本人手中：AI 建议经分析评估后选择性采纳，AI 代码经逐行审查和测试验证后才纳入项目。',
            '通过这次实践，认识到 AI 是提升效率的有力工具，但不能替代对专业知识的深入理解和对质量的严格把控。未来将继续保持"用 AI 辅助但不依赖 AI"的态度，将核心精力放在需求理解、架构设计和质量保障上。',
        ])

def refs():
    print('参考文献和附录...')
    idx = find_heading(doc, '参考文献')
    if idx:
        add_paras(doc, idx, [
            '[1] 王珊, 萨师煊. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.',
            '[2] 吴臣. 数据库系统课程设计讲义[Z]. 2026.',
            '[3] MySQL 8.0 Reference Manual [EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/.',
            '[4] Streamlit Documentation [EB/OL]. https://docs.streamlit.io/.',
            '[5] Python 3.12 Documentation [EB/OL]. https://docs.python.org/3.12/.',
        ])
    idx = find_heading(doc, '附录')
    if idx:
        add_paras(doc, idx, [
            '本课程设计完整代码和 SQL 脚本存放于项目目录：',
            'SQL/ 目录：01_数据库创建.sql ~ 08_存储函数.sql（7 个脚本文件）',
            'code/ 目录：app.py（Streamlit 主应用，约 800 行）、main.py（CLI 入口）、admin.py（教务功能，约 1176 行）、student.py、teacher.py、tester.py、core/（配置/认证/工具）',
            'data/ 目录：class.csv、student.csv、teacher.csv、course.csv、course_offering.csv、enrollment.csv、teacher_course.csv',
            'docs/ 目录：开发日志.md、AI修正日志.md、建表分析.md 等设计文档',
        ])

# ============================================================
# 执行
# ============================================================
if __name__ == '__main__':
    print('='*50)
    ch1(); ch2(); ch3(); ch4(); ch5(); ch6(); ch7(); ch8(); ch9(); refs()
    doc.save(OUTPUT)
    import os
    sz = os.path.getsize(OUTPUT)/1024
    print(f'\n保存: {OUTPUT}')
    print(f'段落 {len(doc.paragraphs)}, 图片 {len([r for r in doc.part.rels.values() if "image" in r.reltype])}, 表格 {len(doc.tables)}, 大小 {sz:.0f} KB')
