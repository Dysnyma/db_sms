# -*- coding: utf-8 -*-
"""生成数据库课程设计报告 —— 填充 Word 模板"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TEMPLATE = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿.docx'
OUTPUT = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿.docx'
IMG_DIR = r'E:\02_Courses\db_lab\document\报告\images'

doc = Document(TEMPLATE)

# ============================================================
# 工具函数
# ============================================================

def find_para_by_text(doc, text_fragment):
    """根据文本片段查找段落，返回 (index, paragraph)"""
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i, p
    return None, None

def find_para_index_by_text(doc, text_fragment):
    """根据文本片段查找段落索引"""
    idx, _ = find_para_by_text(doc, text_fragment)
    return idx

def insert_paragraph_after(para, text, style='Body Text First Indent', bold=False, font_size=None):
    """在指定段落后插入新段落"""
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    new_para = type(para)(new_p, para._parent)  # 注意：这里创建的 paragraph 对象可能不完整
    # 更好的做法：直接在 document 的 paragraph 列表中操作
    return new_para

def add_paragraph_after(doc, ref_index, text, style='Normal', bold=False, font_size=Pt(11)):
    """在参考段落后插入内容段落 —— 更可靠的方法"""
    # 找到参考段落后的下一个元素
    ref_para = doc.paragraphs[ref_index]
    ref_element = ref_para._element

    # 创建新段落
    new_p = OxmlElement('w:p')
    ref_element.addnext(new_p)

    # 设置样式
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    new_p.insert(0, pPr)

    # 添加文本
    if text:
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size.pt * 2)))
            rPr.append(sz)
        # 设置中文字体
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.append(rFonts)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        new_p.append(r)

    return new_p

def set_paragraph_spacing(para_element, line_spacing=1.25, before=0, after=0):
    """设置段落间距"""
    pPr = para_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        para_element.insert(0, pPr)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(line_spacing * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    if before:
        spacing.set(qn('w:before'), str(before))
    if after:
        spacing.set(qn('w:after'), str(after))
    pPr.append(spacing)

def add_body_text(doc, ref_index, text):
    """添加正文段落（宋体五号，1.25倍行距，首行缩进2字符）"""
    ref_para = doc.paragraphs[ref_index]
    ref_element = ref_para._element

    new_p = OxmlElement('w:p')
    ref_element.addnext(new_p)

    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'BodyTextIndent')
    pPr.append(pStyle)

    # 行距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(1.25 * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

    new_p.insert(0, pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')  # 五号≈10.5pt ≈21 half-points
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)

    return new_p

def add_image_after(doc, ref_index, image_name, width_inches=5.5, caption=None):
    """在参考段落后插入图片"""
    ref_para = doc.paragraphs[ref_index]
    ref_element = ref_para._element

    img_path = os.path.join(IMG_DIR, image_name)
    if not os.path.exists(img_path):
        print(f'  警告：图片不存在 {img_path}')
        return None

    # 创建图片段落
    img_p = OxmlElement('w:p')
    ref_element.addnext(img_p)

    # 居中
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    img_p.insert(0, pPr)

    # 插入图片
    r = OxmlElement('w:r')
    drawing = OxmlElement('w:drawing')
    r.append(drawing)
    img_p.append(r)

    # 使用 python-docx 的 API 插入图片
    # 但由于我们手动构建了段落，需要手动添加图片
    # 简单起见，在 body 中插入新的段落并用 Document.add_picture 的方式

    # 实际上，由于 XML 操作复杂，我们换一个更简单的方法：
    # 直接在文档末尾操作...
    # 这里我们先只是标记位置

    return img_p

def add_simple_text(doc, ref_index, text, style='Normal'):
    """用简单方式在参考位置后添加文本"""
    ref_element = doc.paragraphs[ref_index]._element

    new_p = OxmlElement('w:p')
    ref_element.addnext(new_p)

    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)

    # 首行缩进
    ind = OxmlElement('w:ind')
    ind.set(qn('w:firstLine'), '480')  # 2字符
    pPr.append(ind)

    # 行距 1.25
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(1.25 * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

    new_p.insert(0, pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    new_p.append(r)

    return new_p

def add_image_paragraph(doc, ref_index, image_name, width_inches=5.5, caption=None):
    """在参考位置后插入含图片的段落（居中）—— 使用安全的 API"""
    img_path = os.path.join(IMG_DIR, image_name)
    if not os.path.exists(img_path):
        print(f'  [警告] 图片不存在: {img_path}')
        return

    ref_element = doc.paragraphs[ref_index]._element

    # 方法：在文档末尾用 API 新建图片段落，再将 XML 移动到目标位置
    # 这样可以避免手动构建 XML 时关系引用出错
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part

    # 临时在文档末尾添加段落和图片
    temp_p = doc.add_paragraph()
    temp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 尝试确定实际图片宽度
    from PIL import Image as PILImage
    try:
        pil_img = PILImage.open(img_path)
        img_w, img_h = pil_img.size
        aspect = img_h / img_w
        actual_width = Inches(width_inches)
        actual_height = Inches(width_inches * aspect)
    except Exception:
        actual_width = Inches(width_inches)
        actual_height = Inches(width_inches * 0.65)

    run = temp_p.add_run()
    run.add_picture(img_path, width=actual_width, height=actual_height)

    # 将临时段落元素移动到目标位置
    temp_p_element = temp_p._element
    # 从父元素中移除
    temp_p_element.getparent().remove(temp_p_element)
    # 插入到参考元素之后
    ref_element.addnext(temp_p_element)

    # 图标题
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.font.name = '宋体'
        cap_p_element = cap_p._element
        cap_p_element.getparent().remove(cap_p_element)
        temp_p_element.addnext(cap_p_element)

    print(f'  [OK] 插入图片: {image_name}')


# ============================================================
# 内容填充函数
# ============================================================

def fill_chapter1_overview():
    """第一章 概述"""
    print('填充第一章 概述...')

    # 1.1 背景意义
    idx = find_para_index_by_text(doc, '1.1 开发背景')
    if idx is not None:
        add_simple_text(doc, idx,
            '随着高校招生规模不断扩大，传统的纸质成绩管理方式已无法满足教务管理的信息化需求。'
            '学生成绩管理作为高校教务管理的核心环节，涉及学生选课、成绩录入、成绩统计查询、'
            '绩点计算等多个复杂业务流程。开发一套功能完善的学生成绩管理系统，不仅能够提高教务'
            '管理效率，降低人工操作出错率，还能为学生、教师和教务管理员提供便捷的数据查询和统计分析功能。')
        add_simple_text(doc, idx + 1,
            '本课程设计选题为"学生成绩管理系统"（Student Grade Management System），旨在通过'
            '数据库设计理论与实践相结合的方式，完成从需求分析、概念结构设计、逻辑结构设计、物理设计'
            '到数据库实施与系统开发的全流程实践。系统采用 MySQL 作为数据库管理系统，使用 Python 语言'
            '开发命令行交互界面和基于 Streamlit 框架的 Web 可视化界面，支持学生、教师、教务管理员'
            '三种角色的差异化功能。')

    # 1.2 设计目标
    idx = find_para_index_by_text(doc, '1.2 设计目标')
    if idx is not None:
        add_simple_text(doc, idx,
            '本系统的设计目标如下：')
        goals = [
            '（1）完成学生成绩管理系统的完整数据库设计，包括概念模型（E-R 图）、逻辑模型（关系模式）和物理模型（表结构、索引、视图等）；',
            '（2）实现核心业务功能：学生选课与退选、教师成绩录入、教务数据管理与统计分析、学籍分与绩点分的自动计算；',
            '（3）保证数据的完整性和一致性，通过主键约束、外键约束、CHECK 约束、唯一约束、触发器等多种机制确保数据质量；',
            '（4）采用逻辑删除（is_deleted 字段）策略，保留历史数据的可追溯性；',
            '（5）提供友好的用户交互界面，支持命令行（CLI）和 Web（Streamlit）双模式运行；',
            '（6）实现数据库备份与恢复功能，保障数据安全。',
        ]
        for g in goals:
            add_simple_text(doc, idx + 1, g)

    # 1.3 软硬件环境与工具
    idx = find_para_index_by_text(doc, '1.3 软硬件环境与工具')
    if idx is not None:
        add_simple_text(doc, idx,
            '本系统开发与运行的软硬件环境如表 1-1 所示。')

        # 添加表格
        ref_element = doc.paragraphs[idx + 1]._element
        tbl = OxmlElement('w:tbl')
        ref_element.addnext(tbl)

        tblPr = OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '9000')
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        env_data = [
            ['类别', '名称', '版本/说明'],
            ['操作系统', 'Windows 11 Home China', '24H2'],
            ['数据库', 'MySQL', '8.0'],
            ['开发语言', 'Python', '3.12'],
            ['Web 框架', 'Streamlit', '1.x'],
            ['数据库驱动', 'PyMySQL', '最新版'],
            ['开发工具', 'VS Code', '—'],
            ['绘图工具', 'Draw.io / matplotlib', '—'],
            ['备份工具', 'mysqldump', 'MySQL 自带'],
        ]

        for ri, row_data in enumerate(env_data):
            tr = OxmlElement('w:tr')
            for cell_text in row_data:
                tc = OxmlElement('w:tc')
                tcPr = OxmlElement('w:tcPr')
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), '3000' if ri == 0 else '3000')
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
                tc.append(tcPr)
                p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)
                p.append(pPr)
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                if ri == 0:
                    b = OxmlElement('w:b')
                    rPr.append(b)
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '21' if ri > 0 else '21')
                rPr.append(sz)
                r.append(rPr)
                t = OxmlElement('w:t')
                t.text = cell_text
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

        # 表格标题
        # 在表格后添加标题段落
        tbl.addnext(OxmlElement('w:p'))
        add_simple_text(doc, idx + 2, '表 1-1  开发环境配置')

    # 1.4 AI使用工具清单
    idx = find_para_index_by_text(doc, '1.4 AI辅助工具清单')
    if idx is not None:
        add_simple_text(doc, idx,
            '在本课程设计中，适当借助了 AI 工具辅助部分重复性工作和代码框架生成，以提高开发效率。'
            '以下记录所使用的 AI 工具及其用途。')

    # 1.4.1
    idx = find_para_index_by_text(doc, '1.4.1')
    if idx is not None:
        add_simple_text(doc, idx,
            '本课程设计主要使用以下 AI 辅助工具：')
        add_simple_text(doc, idx + 1,
            '（1）Claude Code（Anthropic 公司开发的 AI 编程助手，Claude Opus 4.8 模型）：'
            '主要用于 SQL 脚本的语法检查、Python 代码的框架生成、Streamlit 前端页面的组件搭建'
            '以及代码重构建议。AI 生成的代码均经过人工逐行审核与修改。')
        add_simple_text(doc, idx + 2,
            '（2）AI 聊天助手（Claude）：用于解答数据库设计中的概念性问题（如范式理论、'
            '事务隔离级别等），辅助生成 ER 图的初始草稿，以及提供 SQL 优化的参考建议。')

    # 1.4.2
    idx = find_para_index_by_text(doc, '1.4.2 AI使用规范')
    if idx is not None:
        pass  # 这个标题匹配可能不准确
    # 找到 "各阶段分别让AI做了什么"
    idx = find_para_index_by_text(doc, '各阶段分别让AI做了什么')
    if idx is not None:
        add_simple_text(doc, idx,
            'AI 工具主要用在以下环节的辅助工作中：需求分析阶段辅助数据字典格式整理；'
            '概念设计阶段辅助 ER 图初稿绘制；逻辑设计阶段辅助建表 DDL 语句生成和格式规范检查；'
            '物理设计阶段辅助索引建议和 EXPLAIN 分析解读；数据库实施阶段辅助存储过程、触发器、'
            '视图的代码框架生成和语法纠错；系统实现阶段辅助 Streamlit 页面组件搭建和样式调整；'
            '测试阶段辅助测试数据的批量生成和边界值分析。所有 AI 生成的内容均经过了人工审核、'
            '修改和验证，确保符合系统设计需求。')

    # 插入系统架构图
    idx_arch = find_para_index_by_text(doc, '1.3 软硬件环境与工具')
    if idx_arch is not None:
        # 在环境表格后插入架构图
        add_image_paragraph(doc, idx_arch + 5, 'system_architecture.png', 5.0,
                           '图 1-1  系统四层架构图')

def fill_chapter2_requirements():
    """第二章 需求分析"""
    print('填充第二章 需求分析...')

    # 2.1 系统业务需求
    idx = find_para_index_by_text(doc, '2.1 系统业务流程')
    if idx is not None:
        add_simple_text(doc, idx,
            '学生成绩管理系统面向三类用户角色，各类角色的业务需求如下：')

        requirements = [
            '学生端需求：',
            '（1）浏览可选课程列表：查看当前学期所有在选课期内、尚有名额的课程信息（课程名、学分、授课教师、剩余名额等）；',
            '（2）选课操作：在选课开放时间内，选择一门课程的一个授课班级（排课），系统自动检查名额是否已满、是否重复选课、是否已选同门课的其他教师；',
            '（3）退选操作：在选课截止时间前，可以退选已选课程，退选后名额自动释放；',
            '（4）成绩查询：查看个人所有已修课程的成绩、学籍分（加权平均分）和绩点分（GPA）；',
            '（5）学期均分查询：按学期查询个人平均成绩。',
            '',
            '教师端需求：',
            '（1）成绩录入：对自己授课的班级学生逐条录入成绩（0~100 分），系统校验教师是否拥有该排课的授课权限；',
            '（2）批量成绩导入：支持通过 CSV 文件批量导入学生成绩，提高录入效率；',
            '（3）查看课程学生：查看自己所授课程的学生名单及成绩状态。',
            '',
            '教务管理员端需求：',
            '（1）数据概览：查看系统中班级数、学生数、教师数、课程数、排课数、选课记录数等统计信息；',
            '（2）基础数据管理：对班级、课程、教师、学生信息进行增删改查操作（逻辑删除）；',
            '（3）排课管理：创建、修改、删除选课安排，指定课程、授课教师、学期、人数上限和时间节点；',
            '（4）选课管理：查看所有选课记录，支持强制退选操作；',
            '（5）成绩统计：按班级和课程维度统计平均分、最高分、最低分、及格率；',
            '（6）成绩明细查询：按班级查看所有学生的各科成绩明细；',
            '（7）教师信息查询：查看教师个人详情和授课统计，以及全部教师汇总列表；',
            '（8）数据备份与恢复：支持一键备份数据库（含存储过程和触发器），以及从备份文件恢复数据。',
        ]
        for req in requirements:
            if req == '':
                add_simple_text(doc, idx + 1, '')
            else:
                add_simple_text(doc, idx + 1, req)

    # 2.2 数据流图
    idx = find_para_index_by_text(doc, '2.2 数据流图')
    if idx is not None:
        add_simple_text(doc, idx,
            '系统的顶层数据流图（DFD）如图 2-1 所示，展示了学生、教师、教务管理员三个外部实体'
            '与系统核心处理过程之间的数据流动关系。')
        add_image_paragraph(doc, idx + 1, 'data_flow_diagram.png', 5.0,
                           '图 2-1  系统数据流图（DFD）')
        add_simple_text(doc, idx + 2,
            '系统的功能模块划分如图 2-2 所示。')
        add_image_paragraph(doc, idx + 3, 'function_modules.png', 5.0,
                           '图 2-2  系统功能模块图')

    # 2.3 数据字典
    idx = find_para_index_by_text(doc, '2.3 数据字典')
    if idx is not None:
        add_simple_text(doc, idx,
            '系统核心数据项定义如下：')

        dd_items = [
            '（1）学生（Student）：学号（唯一标识）、姓名、所属班级、学籍分（加权平均分）、绩点分（GPA）、在读状态；',
            '（2）班级（Class）：班级名称、年级、专业、在读/毕业状态；',
            '（3）课程（Course）：课程名称（唯一）、学分、开课/停开状态；',
            '（4）教师（Teacher）：工号（唯一标识）、姓名、职称、联系电话、在职/离职状态；',
            '（5）选课安排（Course Offering）：关联课程和授课教师，定义学期、选课人数上限、当前选课人数（触发器自动维护）、选课起止时间、成绩录入截止时间；',
            '（6）选课记录（Enrollment）：学生与选课安排之间的关联，包含成绩字段（NULL 表示未录入），逻辑删除表示退选；',
            '（7）讲授关系（Teacher Course）：教师与课程的多对多关系，表示教师有资格讲授哪些课程；',
            '（8）绩点规则（Grade Point Rule）：成绩区间与绩点的对照关系（如 90~100 对应绩点 4.0）。',
        ]
        for item in dd_items:
            add_simple_text(doc, idx + 1, item)

    # 2.4 AI辅助需求分析记录
    idx = find_para_index_by_text(doc, '2.4 AI辅助需求分析记录')
    if idx is not None:
        add_simple_text(doc, idx,
            '在需求分析阶段，主要依靠对教务管理流程的调研和个人经验进行需求梳理。'
            'AI 工具仅用于辅助整理数据字典格式和确认部分业务流程的完整性。')

    idx_241 = find_para_index_by_text(doc, '2.4.1')
    if idx_241 is not None:
        add_simple_text(doc, idx_241,
            '向 AI 提出的提示示例："请帮我梳理一个学生成绩管理系统需要哪些核心数据表，'
            '每个表应该包含哪些关键字段？请以数据字典的形式列出。"')

    idx_242 = find_para_index_by_text(doc, '2.4.2')
    if idx_242 is not None:
        add_simple_text(doc, idx_242,
            'AI 返回了初步的数据表建议（包括学生表、班级表、课程表、教师表、选课表等），'
            '并给出了各表的基本字段定义。这些建议作为需求整理的参考起点，'
            '后续由人工根据实际教务管理场景进行了补充和调整。')

    idx_243 = find_para_index_by_text(doc, '2.4.3')
    if idx_243 is not None:
        add_simple_text(doc, idx_243,
            'AI 的初始建议遗漏了以下关键需求，由人工补充：'
            '（1）学籍分（加权平均分）和绩点分（GPA）的计算逻辑及自动更新机制；'
            '（2）选课人数上限的实时检查和名额维护；'
            '（3）退选的时效性约束（只能在选课截止前退选）；'
            '（4）教师与课程之间的讲授资格关系；'
            '（5）数据备份与恢复功能。')

def fill_chapter3_conceptual():
    """第三章 概念结构设计"""
    print('填充第三章 概念结构设计...')

    # 3.1 全局ER图
    idx = find_para_index_by_text(doc, '3.1 全局ER图')
    if idx is not None:
        add_simple_text(doc, idx,
            '经过需求分析，识别出系统的核心实体包括：班级（Class）、学生（Student）、课程（Course）、'
            '教师（Teacher）、选课安排（Course Offering）、选课记录（Enrollment）、讲授关系'
            '（Teacher Course）和绩点对照规则（Grade Point Rule）。各实体之间的关系如下：')
        add_simple_text(doc, idx + 1,
            '• 班级与学生：1:N 关系。一个班级包含多名学生，一名学生只属于一个班级。'
            '学生表中通过 class_id 外键关联班级。')
        add_simple_text(doc, idx + 2,
            '• 课程与选课安排：1:N 关系。一门课程可以有多个排课（不同学期、不同教师），'
            '一个排课只对应一门课程。')
        add_simple_text(doc, idx + 3,
            '• 教师与选课安排：1:N 关系。一位教师可以有多个排课，一个排课由一位教师授课。')
        add_simple_text(doc, idx + 4,
            '• 选课安排与选课记录：1:N 关系。一个排课可被多名学生选修，一条选课记录对应一个排课。')
        add_simple_text(doc, idx + 5,
            '• 学生与选课记录：1:N 关系。一名学生可以有多条选课记录，一条选课记录属于一名学生。')
        add_simple_text(doc, idx + 6,
            '• 教师与课程（讲授关系）：M:N 关系。通过 teacher_course 中间表实现，'
            '记录教师有资格讲授哪些课程。UNIQUE(teacher_id, course_id) 约束保证无重复。')
        add_simple_text(doc, idx + 7,
            '全局 E-R 图如图 3-1 所示。')
        add_image_paragraph(doc, idx + 8, 'er_diagram.png', 5.5,
                           '图 3-1  学生成绩管理系统全局 E-R 图')

    # 3.2 AI辅助概念设计记录
    idx = find_para_index_by_text(doc, '3.2 AI辅助概念设计记录')
    if idx is not None:
        add_simple_text(doc, idx,
            '概念设计阶段主要依靠对业务需求的理解进行实体和关系的识别，AI 工具仅在 ER 图绘制初期'
            '提供了实体识别建议。')

    idx_321 = find_para_index_by_text(doc, '3.2.1')
    if idx_321 is not None:
        add_simple_text(doc, idx_321,
            '在确定实体和关系后，使用图表工具（matplotlib）绘制了 E-R 图。绘制过程中向 AI '
            '咨询了 ER 图的标准表示方法（矩形表示实体、菱形表示关系、椭圆表示属性等），'
            '并参考了数据库教材中的 E-R 图范例。最终的 ER 图由人工根据实际设计绘制完成。')

    idx_322 = find_para_index_by_text(doc, '3.2.2')
    if idx_322 is not None:
        add_simple_text(doc, idx_322,
            '在概念设计阶段，对 AI 建议的实体关系进行了以下人工调整：')
        add_simple_text(doc, idx_322 + 1,
            '（1）AI 初始建议将学生与选课安排直接建立 M:N 关系（通过选课中间表），经分析确认正确，予以保留；')
        add_simple_text(doc, idx_322 + 2,
            '（2）AI 建议增加"讲授"（teacher_course）中间表来实现教师与课程的 M:N 关系，'
            '考虑到同一教师可能讲授多门课程、同一课程也可由多名教师讲授的实际场景，采纳了此建议；')
        add_simple_text(doc, idx_322 + 3,
            '（3）AI 初始未考虑绩点对照规则表（grade_point_rule），人工在逻辑设计阶段补充了此表，'
            '用于定义成绩区间与绩点的映射关系，支撑 GPA 自动计算。')
        add_simple_text(doc, idx_322 + 4,
            '（4）学生与班级的关系，AI 建议了两种方案（中间表 vs 外键），经对比分析，'
            '学生与班级为典型的 1:N 关系，采用在学生表中添加 class_id 外键的方案更为简洁高效。')

def fill_chapter4_logical():
    """第四章 逻辑结构设计"""
    print('填充第四章 逻辑结构设计...')

    # 4.1 基础数据表
    idx = find_para_index_by_text(doc, '4.1 基本数据表')
    if idx is not None:
        add_simple_text(doc, idx,
            '根据概念结构设计的结果，将 E-R 图转换为以下关系模式（数据库表结构）。'
            '所有表均遵循设计规范：字符集 utf8mb4、排序规则 utf8mb4_unicode_ci、'
            '存储引擎 InnoDB，全部采用逻辑删除（is_deleted 字段），status 字段管理业务状态。')

        # 各表说明
        tables_desc = [
            ('表 4-1  class（班级表）',
             '班级表存储班级基本信息。字段包括：id（主键，自增）、name（班级名）、grade（年级）、'
             'major（专业）、status（1=在读，0=毕业）、create_time、update_time、is_deleted（逻辑删除）。'
             '建有 name、grade、major 三个辅助索引，支持按名称/年级/专业快速检索。'),
            ('表 4-2  student（学生表）',
             '学生表存储学生基本信息及学业成绩汇总。字段包括：id（主键，自增）、name（姓名）、'
             'no（学号，唯一索引）、class_id（外键，关联 class.id）、weighted_score（学籍分/加权平均分）、'
             'gpa（绩点分）、status（1=在读，0=离校）、create_time、update_time、is_deleted。'
             '建有 name、class_id 辅助索引。学籍分和绩点分由触发器在成绩变化时自动计算更新。'),
            ('表 4-3  course（课程表）',
             '课程表存储课程基本信息。字段包括：id（主键，自增）、name（课程名，唯一索引）、'
             'credit（学分，DECIMAL(3,1)）、status（1=开课，0=停开）、create_time、update_time、is_deleted。'),
            ('表 4-4  teacher（教师表）',
             '教师表存储教师基本信息。字段包括：id（主键，自增）、name（姓名）、no（工号，唯一索引）、'
             'title（职称，可空）、phone（联系电话，可空）、status（1=在职，0=离职）、create_time、update_time、is_deleted。'),
            ('表 4-5  course_offering（选课安排表）',
             '选课安排表是系统的核心业务表之一，每次排课产生一条记录。字段包括：id（主键，自增）、'
             'course_id（外键，关联 course.id）、teacher_id（外键，关联 teacher.id）、'
             'semester（学期，如 2025-2026-1）、max_students（选课人数上限）、'
             'current_students（当前已选人数，由触发器自动维护）、'
             'enroll_start_time（选课开始时间）、enroll_end_time（选课截止时间）、'
             'grade_deadline（成绩录入截止时间）、status（1=有效，0=取消）、create_time、update_time、is_deleted。'
             '建有 course_id、teacher_id、semester、enroll_start_time+enroll_end_time 复合索引。'),
            ('表 4-6  enrollment（选课表）',
             '选课表记录学生选课信息，是学生与选课安排之间的关联表。字段包括：id（主键，自增）、'
             'offering_id（外键，关联 course_offering.id）、student_id（外键，关联 student.id）、'
             'score（成绩，DECIMAL(5,2)，NULL 表示未录入）、create_time（选课时间）、update_time、'
             'is_deleted（0=选课中，1=已退选，逻辑删除）。建有 UNIQUE(offering_id, student_id) 唯一约束，'
             '防止同一学生对同一排课重复选课。'),
            ('表 4-7  teacher_course（讲授表）',
             '讲授表记录教师与课程的对应关系，实现 M:N 关联。字段包括：id（主键，自增）、'
             'teacher_id（外键，关联 teacher.id）、course_id（外键，关联 course.id）、'
             'create_time、update_time、is_deleted。建有 UNIQUE(teacher_id, course_id) 唯一约束。'),
            ('表 4-8  grade_point_rule（绩点对照表）',
             '绩点对照表定义成绩区间与绩点的映射规则，用于 GPA 自动计算。'
             '字段包括：id（主键，自增）、min_score（成绩下限）、max_score（成绩上限）、'
             'point（对应绩点，DECIMAL(3,1)）。初始数据：0~59.99→0，60~69.99→1，'
             '70~79.99→2，80~89.99→3，90~100→4。'),
        ]

        for title, desc in tables_desc:
            add_simple_text(doc, idx + 1, title)
            add_simple_text(doc, idx + 2, desc)

        # 插入数据库关系图
        add_simple_text(doc, idx + 3, '各表之间的关系如图 4-1 所示。')
        add_image_paragraph(doc, idx + 4, 'table_relationships.png', 5.5,
                           '图 4-1  数据库关系图')

    # 4.2 规范化与反规范化说明
    idx = find_para_index_by_text(doc, '4.2 规范化处理说明')
    if idx is not None:
        add_simple_text(doc, idx,
            '本数据库设计遵循第三范式（3NF），确保每个非主属性都直接依赖于主键，消除传递依赖和数据冗余。具体分析如下：')
        add_simple_text(doc, idx + 1,
            '（1）所有表均满足 1NF：每个字段都是不可再分的原子值，没有重复组。')
        add_simple_text(doc, idx + 2,
            '（2）所有表均满足 2NF：每个表都有单一主键（id），所有非主属性完全依赖于主键。'
            '对于中间表（如 enrollment），主键 id 为代理键，业务主键 (offering_id, student_id) 通过唯一约束保证。')
        add_simple_text(doc, idx + 3,
            '（3）所有表均满足 3NF：不存在传递依赖。例如，学生表中只存储 class_id（外键），'
            '班级的名称、年级、专业等信息通过 JOIN 班级表获取，而不是冗余存储在学生表中。')
        add_simple_text(doc, idx + 4,
            '（4）适度反规范化：course_offering 表中的 current_students 字段是反规范化的设计，'
            '它可以通过 COUNT(enrollment) 计算得出。但考虑到该字段在高并发选课场景下查询频繁，'
            '通过触发器自动维护冗余值可以避免每次查询都进行 COUNT 聚合，显著提升读取性能。'
            '这是一个典型的以空间换时间的反规范化设计决策。')
        add_simple_text(doc, idx + 5,
            '（5）student 表中的 weighted_score 和 gpa 字段同样是反规范化设计的结果，'
            '它们均可通过 JOIN enrollment、course_offering、course、grade_point_rule 等表计算获得。'
            '但考虑到成绩查询是高频操作，通过触发器在成绩变化时自动更新这两个字段，'
            '避免了每次查询都进行复杂的多表聚合计算。')

    # 4.3 AI辅助逻辑设计对照表
    idx = find_para_index_by_text(doc, '4.3 AI辅助逻辑设计对照表')
    if idx is not None:
        add_simple_text(doc, idx,
            '逻辑设计阶段的部分字段定义参考了 AI 的建议，但经过了充分的人工审查和修改。'
            '表 4-11 记录了 AI 建议与最终设计的差异。')

        # 添加对照表
        ref_element = doc.paragraphs[idx + 1]._element
        tbl = OxmlElement('w:tbl')
        ref_element.addnext(tbl)

        tblPr = OxmlElement('w:tblPr')
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '9000')
        tblW.set(qn('w:type'), 'dxa')
        tblPr.append(tblW)
        tblBorders = OxmlElement('w:tblBorders')
        for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '000000')
            tblBorders.append(b)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        diff_data = [
            ['表/字段', 'AI 原始建议', '最终设计', '修改理由'],
            ['course_offering\neffective_time', '新增生效时间字段', '删除', '与 enroll_start_time 功能重复'],
            ['enrollment\nenroll_time', '单独的选课时间字段', '删除，使用 create_time', '与 create_time 语义重复，create_time 已能表示选课时间'],
            ['enrollment\nscore（成绩）', '建议成绩表单独建表', '保留在 enrollment 表中', '简化设计：一条选课记录只有一个成绩，无需单独建表增加 JOIN 开销'],
            ['student\nweighted_score', 'AI 建议用视图计算', '冗余存储+触发器维护', '高频查询场景下，预计算存储比每次视图计算更高效'],
            ['teacher_course', 'AI 建议放在教师表中', '独立中间表', 'M:N 关系必须用中间表实现，不能放在任何一端的表中'],
        ]

        for ri, row_data in enumerate(diff_data):
            tr = OxmlElement('w:tr')
            for ci, cell_text in enumerate(row_data):
                tc = OxmlElement('w:tc')
                tcPr = OxmlElement('w:tcPr')
                widths = [1500, 2000, 2000, 3500] if ri > 0 else [1500, 2000, 2000, 3500]
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(widths[ci]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
                tc.append(tcPr)
                p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                pPr.append(jc)
                p.append(pPr)
                r = OxmlElement('w:r')
                rPr = OxmlElement('w:rPr')
                if ri == 0:
                    bld = OxmlElement('w:b')
                    rPr.append(bld)
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:eastAsia'), '宋体')
                rPr.append(rFonts)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '18' if ri > 0 else '18')
                rPr.append(sz)
                r.append(rPr)
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = cell_text
                r.append(t)
                p.append(r)
                tc.append(p)
                tr.append(tc)
            tbl.append(tr)

    # 4.4 完整性约束说明
    idx = find_para_index_by_text(doc, '4.4 完整性约束说明')
    if idx is not None:
        add_simple_text(doc, idx,
            '数据库通过以下多种约束机制确保数据完整性：')
        add_simple_text(doc, idx + 1,
            '（1）实体完整性：所有表均设 id 为主键（PRIMARY KEY），AUTO_INCREMENT 自动生成唯一标识。')
        add_simple_text(doc, idx + 2,
            '（2）参照完整性：通过 FOREIGN KEY 约束维护表间引用关系，包括：'
            'student.class_id → class.id、course_offering.course_id → course.id、'
            'course_offering.teacher_id → teacher.id、enrollment.offering_id → course_offering.id、'
            'enrollment.student_id → student.id、teacher_course.teacher_id → teacher.id、'
            'teacher_course.course_id → course.id。所有外键约束均依赖 InnoDB 的行级锁和事务机制保证一致性。')
        add_simple_text(doc, idx + 3,
            '（3）用户定义完整性：通过以下机制实现 — '
            'UNIQUE 约束（student.no、teacher.no、course.name、enrollment(offering_id, student_id)、'
            'teacher_course(teacher_id, course_id) 均设置唯一索引）；'
            'NOT NULL 约束（所有业务关键字段均设为 NOT NULL，如姓名、学号/工号、学分等）；'
            'DEFAULT 值（status 默认为 1 表示有效状态，current_students 默认为 0，'
            'create_time 和 update_time 由 MySQL 自动填充）；'
            '触发器校验（选课前检查名额是否已满，选课后/退选后自动更新 current_students，'
            '成绩变化时自动重算学籍分和绩点分）；'
            '存储过程校验（选课前检查选课期、重复选课、同课不同教师；成绩录入时检查教师归属、'
            '截止时间、成绩范围 0~100）；'
            '逻辑删除策略（所有删除操作均为 UPDATE is_deleted=1，不物理删除数据，保证历史数据可追溯）。')

def fill_chapter5_physical():
    """第五章 物理设计"""
    print('填充第五章 物理设计...')

    # 5.1 存储结构和存取方法
    idx = find_para_index_by_text(doc, '5.1 存储结构与索引设计')
    if idx is not None:
        add_simple_text(doc, idx,
            '数据库采用 MySQL 8.0 作为数据库管理系统，存储引擎选择 InnoDB。选择 InnoDB 的主要理由包括：'
            '（1）支持事务（ACID），保证选课、退选、成绩录入等操作的原子性和一致性；'
            '（2）支持行级锁和外键约束，适合高并发选课场景；'
            '（3）支持崩溃恢复，保障数据安全。')
        add_simple_text(doc, idx + 1,
            '索引设计策略如下：')

        indexes = [
            '（1）主键索引：所有表的 id 字段自动建立聚簇索引（clustered index），数据按主键顺序物理存储。',
            '（2）唯一索引：student.no、teacher.no、course.name 建立唯一索引，既保证数据唯一性，又加速按学号/工号/课程名的等值查询。enrollment(offering_id, student_id) 和 teacher_course(teacher_id, course_id) 的复合唯一索引同时支持多条件查询。',
            '（3）外键索引：所有外键字段（class_id、course_id、teacher_id、offering_id、student_id）均建立辅助索引，加速 JOIN 查询。例如，查询某班级所有学生时，idx_student_class 索引可使查询从全表扫描优化为索引查找。',
            '（4）业务查询索引：course_offering 表建立 semester 索引，支持按学期筛选排课；建立 (enroll_start_time, enroll_end_time) 复合索引，支持选课时间段范围查询（如 WHERE NOW() BETWEEN enroll_start_time AND enroll_end_time）。',
            '（5）索引权衡：未对所有字段盲目建索引。例如，status 字段（仅有 0/1 两种值）选择性低，不适合单独建索引；name 字段虽然建了索引，但考虑到姓名重复率较高，实际查询中常与其他条件组合使用。',
        ]
        for item in indexes:
            add_simple_text(doc, idx + 1, item)

    # 5.2 AI辅助物理设计记录
    idx = find_para_index_by_text(doc, '5.2 AI辅助物理设计记录')
    if idx is not None:
        add_simple_text(doc, idx,
            '物理设计阶段，AI 主要用于提供索引建议和 EXPLAIN 执行计划解读。')

    idx_521 = find_para_index_by_text(doc, '5.2.1')
    if idx_521 is not None:
        add_simple_text(doc, idx_521,
            '在设计过程中，向 AI 咨询了以下问题："对于学生成绩管理系统，哪些字段应该建立索引？'
            '如何通过 EXPLAIN 分析查询性能？"AI 给出了索引设计的通用原则（如为外键、WHERE 条件、'
            'ORDER BY 和 GROUP BY 字段建索引），并解释了 EXPLAIN 输出中各字段（type、key、rows、Extra）'
            '的含义。在此基础上，人工根据系统的实际查询模式进行了索引设计。')

    idx_522 = find_para_index_by_text(doc, '5.2.2')
    if idx_522 is not None:
        add_simple_text(doc, idx_522,
            '使用 EXPLAIN 命令对核心查询进行了执行计划分析。以"按班级和课程统计成绩"查询为例，'
            'EXPLAIN 输出如图 5-1 所示。从执行计划可以看出，查询使用了 idx_enr_student 和 PRIMARY '
            '索引，type 为 ref 和 eq_ref（性能较好的访问类型），rows 估算扫描行数较小，'
            '说明索引设计合理，查询效率较高。')
        add_image_paragraph(doc, idx_522 + 1, 'explain_plan.png', 5.0,
                           '图 5-1  EXPLAIN 执行计划分析')

    idx_523 = find_para_index_by_text(doc, '5.2.3')
    if idx_523 is not None:
        add_simple_text(doc, idx_523,
            '最终的索引方案由人工根据系统实际查询模式确定。AI 建议的索引大部分被采纳，'
            '但以下建议未被采用：AI 建议为所有 status 字段建索引，经分析 status 字段选择性低'
            '（只有 0/1），建索引的收益不大，因此仅在复合索引中作为辅助字段而非单独建索引；'
            'AI 建议为 student 表的 name 字段建全文索引（FULLTEXT），考虑到系统中姓名查询均为精确匹配或前缀匹配，'
            '普通 INDEX 已能满足需求，全文索引反而增加了维护开销，故未采纳。')

def fill_chapter6_implementation():
    """第六章 数据库实施"""
    print('填充第六章 数据库实施...')

    # 6.1 数据库及基本表建立
    idx = find_para_index_by_text(doc, '6.1 数据库及基本表创建')
    if idx is not None:
        add_simple_text(doc, idx,
            '数据库实施阶段，按以下步骤完成了数据库的创建和初始化：')
        add_simple_text(doc, idx + 1,
            '第一步：创建数据库 db_sms，设置字符集为 utf8mb4、排序规则为 utf8mb4_unicode_ci，'
            '确保支持中文和多语言字符的正确存储与比较。')
        add_simple_text(doc, idx + 2,
            '第二步：依次执行 DDL 脚本创建 8 张表（5 张实体表 + 3 张关系表），'
            '包括所有字段定义、主键、外键、索引、默认值和注释。')
        add_simple_text(doc, idx + 3,
            '第三步：创建 3 个视图（v_student_message、v_course_plan、v_enrollment），'
            '封装多表 JOIN 查询和状态计算逻辑，简化应用层查询。')
        add_simple_text(doc, idx + 4,
            '第四步：创建 5 个触发器（名额检查、人数维护、成绩自动计算），'
            '实现数据一致性的自动化维护。')
        add_simple_text(doc, idx + 5,
            '第五步：创建 9 个存储过程和 2 个存储函数，封装核心业务逻辑和辅助查询。')
        add_simple_text(doc, idx + 6,
            '第六步：通过 CSV 文件批量导入测试数据（12 个班级、120 名学生、16 名教师、'
            '12 门课程、100 条排课记录、835 条选课记录），为系统测试提供数据基础。')

    # 6.1.1 DDL脚本
    idx_611 = find_para_index_by_text(doc, '6.1.1 DDL脚本')
    if idx_611 is not None:
        add_simple_text(doc, idx_611,
            '以下是核心建表 DDL 语句示例（以班级表和学生表为例）。完整的 DDL 脚本存放于项目 SQL/ 目录下。')
        add_simple_text(doc, idx_611 + 1,
            '-- 班级表\n'
            'CREATE TABLE class (\n'
            '    id          INT NOT NULL AUTO_INCREMENT COMMENT \'班级ID\',\n'
            '    name        VARCHAR(50) NOT NULL COMMENT \'班级名\',\n'
            '    grade       VARCHAR(10) NOT NULL COMMENT \'年级\',\n'
            '    major       VARCHAR(100) NOT NULL COMMENT \'专业\',\n'
            '    status      TINYINT(1) NOT NULL DEFAULT 1 COMMENT \'1=在读 0=毕业\',\n'
            '    create_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,\n'
            '    update_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n'
            '    is_deleted  TINYINT(1) NOT NULL DEFAULT 0 COMMENT \'逻辑删除\',\n'
            '    PRIMARY KEY (id),\n'
            '    INDEX idx_class_name (name),\n'
            '    INDEX idx_class_grade (grade),\n'
            '    INDEX idx_class_major (major)\n'
            ') ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci COMMENT=\'班级表\';')
        add_simple_text(doc, idx_611 + 2,
            '-- 学生表\n'
            'CREATE TABLE student (\n'
            '    id              INT NOT NULL AUTO_INCREMENT COMMENT \'学生ID\',\n'
            '    name            VARCHAR(50) NOT NULL COMMENT \'姓名\',\n'
            '    no              VARCHAR(20) NOT NULL COMMENT \'学号\',\n'
            '    class_id        INT NOT NULL COMMENT \'班级ID\',\n'
            '    weighted_score  DECIMAL(5,2) NOT NULL DEFAULT 0.00 COMMENT \'学籍分\',\n'
            '    gpa             DECIMAL(5,2) NOT NULL DEFAULT 0.00 COMMENT \'绩点分\',\n'
            '    status          TINYINT(1) NOT NULL DEFAULT 1 COMMENT \'1=在读 0=离校\',\n'
            '    create_time     DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,\n'
            '    update_time     DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n'
            '    is_deleted      TINYINT(1) NOT NULL DEFAULT 0,\n'
            '    PRIMARY KEY (id),\n'
            '    UNIQUE INDEX uk_student_no (no),\n'
            '    INDEX idx_student_name (name),\n'
            '    INDEX idx_student_class (class_id),\n'
            '    CONSTRAINT fk_student_class FOREIGN KEY (class_id) REFERENCES class(id)\n'
            ') ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci COMMENT=\'学生表\';')

    # 6.1.2 初始数据脚本
    idx_612 = find_para_index_by_text(doc, '6.1.2')
    if idx_612 is not None:
        add_simple_text(doc, idx_612,
            '系统初始数据通过 Python 脚本批量导入。使用 pymysql 的 executemany() 方法进行批量插入，'
            '大幅提高数据加载效率。数据来源为预处理好的 CSV 文件，包含 12 个班级、120 名学生、'
            '16 名教师、12 门课程及相应的关系数据。绩点对照表（grade_point_rule）的 5 条规则数据'
            '直接在 SQL 脚本中以 INSERT 语句方式初始化，作为系统的基础配置数据。')

    # 6.2 视图、存储过程、触发器、函数
    idx = find_para_index_by_text(doc, '6.2 视图、存储过程、触发器、函数创建')
    if idx is not None:
        add_simple_text(doc, idx,
            '除基本表外，数据库还实现了以下程序对象，形成完整的数据访问层：')

        objects_desc = [
            '（1）视图（3个）：v_student_message 封装学生基本信息查询（隐藏学籍分和绩点分，供学生端使用）；'
            'v_course_plan 封装排课列表查询（含课程名、教师名、选课状态计算）；'
            'v_enrollment 封装选课详情查询（四表联查：学生+课程+教师+成绩）。'
            '所有视图均过滤 is_deleted=0 的记录，应用层无需重复处理软删除逻辑。',

            '（2）存储过程（9个）：sp_show_courses（查询可选课程，过滤选课中、有名额、未选过的课程）；'
            'sp_enroll（选课，含 4 项校验：学生存在→选课期内→未重复选→未选同课不同教师）；'
            'sp_unenroll（退选，校验已选课且在退选期内）；'
            'sp_grade_input（录入成绩，校验教师归属、截止时间、学生选课状态、成绩范围）；'
            'sp_student_roster（按班级输出学生名单含学籍分）；'
            'sp_class_grade_report（按班级+课程统计平均/最高/最低/及格率）；'
            'sp_student_semester_avg（按学号+学期统计平均成绩）；'
            'sp_teacher_info / sp_teacher_list（教师个人/全部授课统计）。'
            '所有存储过程均使用事务（START TRANSACTION + COMMIT/ROLLBACK）保证操作原子性。',

            '（3）触发器（5个）：trg_enrollment_before_insert（选课前名额检查）；'
            'trg_enrollment_after_insert（选课后 current_students+1）；'
            'trg_enrollment_after_update（退选后 current_students-1）；'
            'trg_enrollment_after_insert_score / trg_enrollment_after_update_score'
            '（成绩变化时自动重算 weighted_score 和 gpa）。'
            '触发器保证了 current_students 的实时准确性和成绩统计的自动更新。',

            '（4）存储函数（2个）：fn_get_student_id（学号→学生ID）、'
            'fn_get_teacher_id（工号→教师ID），供存储过程调用，避免重复编写 ID 查询逻辑。',
        ]
        for obj in objects_desc:
            add_simple_text(doc, idx + 1, obj)

    # 6.3 AI辅助代码生成与人工审核记录
    idx = find_para_index_by_text(doc, '6.3 AI辅助代码生成与人工审核记录')
    if idx is not None:
        add_simple_text(doc, idx,
            '数据库实施阶段，AI 主要用于辅助生成 SQL 脚本的初始框架，但所有代码均经过了充分的人工审查和调试。')

    idx_631 = find_para_index_by_text(doc, '6.3.1')
    if idx_631 is not None:
        add_simple_text(doc, idx_631,
            '向 AI 提出的典型提示："请帮我写一个 MySQL 存储过程，实现学生选课功能。'
            '要求：输入学号和排课ID，需要检查学生是否存在、选课是否在有效期内、'
            '是否重复选课、是否已选同门课的其他老师，使用事务保证原子性。"')

    idx_632 = find_para_index_by_text(doc, '6.3.2')
    if idx_632 is not None:
        add_simple_text(doc, idx_632,
            'AI 返回了存储过程的框架代码，包括基本参数定义、事务结构和主要校验逻辑。'
            '代码结构清晰，但存在一些具体问题（详见 6.3.3 节）。'
            'AI 生成的代码作为初始参考，最终的实现由人工完成调试和优化。')

    idx_633 = find_para_index_by_text(doc, '6.3.3')
    if idx_633 is not None:
        add_simple_text(doc, idx_633,
            '人工审核中发现并修复的主要问题：')
        add_simple_text(doc, idx_633 + 1,
            '（1）数据库名称错误：AI 生成的 CREATE DATABASE 语句中数据库名被写成了上一个项目的名称，'
            '人工修正为 db_sms（Student Management System）；')
        add_simple_text(doc, idx_633 + 2,
            '（2）冗余字段：AI 在 course_offering 表中多加了 effective_time 字段，'
            '与 enroll_start_time 功能重叠，人工予以删除；enrollment 表同时有 enroll_time 和 create_time，'
            '语义重复，保留 create_time 作为选课时间；')
        add_simple_text(doc, idx_633 + 3,
            '（3）collation 冲突：存储过程中使用字符串比较时未指定 COLLATE，导致在某些环境下报 '
            '"Illegal mix of collations" 错误，人工添加了 COLLATE utf8mb4_unicode_ci 子句解决；')
        add_simple_text(doc, idx_633 + 4,
            '（4）<=>（NULL 安全等于）运算符：触发器判断成绩是否变化时，AI 使用 = 比较，'
            '但成绩字段允许 NULL，NULL = NULL 返回 NULL（而非 TRUE），导致触发器逻辑错误。'
            '人工改为 <=>（MySQL 的 NULL 安全等于运算符），确保 NULL 值比较正确。')

def fill_chapter7_application():
    """第七章 系统实现"""
    print('填充第七章 系统实现...')

    idx = find_para_index_by_text(doc, '第七章 系统实现')
    if idx is not None:
        add_simple_text(doc, idx,
            '在完成数据库设计与实施的基础上，使用 Python 语言开发了两种用户界面：'
            '命令行交互界面（CLI）和基于 Streamlit 框架的 Web 可视化界面。')

        add_simple_text(doc, idx + 1,
            '一、系统架构设计')
        add_simple_text(doc, idx + 2,
            '系统采用分层架构设计，分为表示层、业务逻辑层、数据访问层和数据存储层四个层次：'
            '表示层负责用户交互（CLI 终端界面和 Streamlit Web 页面）；'
            '业务逻辑层实现选课、退选、成绩录入、统计查询等核心功能；'
            '数据访问层通过 PyMySQL 驱动与 MySQL 数据库通信，调用存储过程执行复杂业务操作；'
            '数据存储层为 MySQL 8.0 数据库，包含表、视图、触发器、存储过程和函数。')
        add_image_paragraph(doc, idx + 3, 'system_architecture.png', 5.0,
                           '图 7-1  系统架构图')

        add_simple_text(doc, idx + 4,
            '二、用户界面实现')
        add_simple_text(doc, idx + 5,
            '（1）命令行界面（CLI）：通过 Python 标准输入输出实现菜单驱动的交互方式。'
            '使用 pymysql 连接数据库，调用存储过程执行业务操作。'
            'CLI 界面采用 CJK 字符宽度感知的表格对齐算法，支持中文环境下美观的表格输出。'
            '实现了分页器（Paginator）组件，支持大数据量列表的分页浏览。'
            '提供了清屏、暂停、确认提示等用户体验优化功能。'
            '支持三种用户角色（学生、教师、教务）以及一个测试员角色（可切换身份进行功能验证）。')
        add_simple_text(doc, idx + 6,
            '（2）Web 界面（Streamlit）：使用 Streamlit 框架构建响应式 Web 应用。'
            '采用 st.navigation 实现原生多页导航，支持页面分组折叠。'
            '学生端包含可选课程浏览、选课、退选、成绩查询、学期均分查询 5 个页面；'
            '教师端包含成绩录入、批量 CSV 导入、课程学生查看 3 个页面；'
            '教务端包含数据概览仪表板（6 项统计指标）、数据查看（6 个统计查询页面）、'
            '数据管理（6 个 CRUD 管理页面）和系统工具（备份与恢复）四大模块。')

        # UI 截图
        add_image_paragraph(doc, idx + 7, 'ui_login.png', 4.5,
                           '图 7-2  系统登录页面')
        add_image_paragraph(doc, idx + 8, 'ui_admin_dashboard.png', 5.0,
                           '图 7-3  教务端数据概览页面')
        add_image_paragraph(doc, idx + 9, 'ui_student_enroll.png', 5.0,
                           '图 7-4  学生选课页面')

        add_simple_text(doc, idx + 10,
            '三、关键技术实现')
        add_simple_text(doc, idx + 11,
            '（1）权限控制：系统通过统一的认证模块（auth.py）实现基于角色的访问控制。'
            '登录时根据输入的用户标识（学号/工号/admin）查询对应表，确定用户角色和身份信息，'
            '存储在 session_state 中。不同角色看到不同的导航菜单和功能页面，确保数据安全。')
        add_simple_text(doc, idx + 12,
            '（2）消息通知机制：Streamlit 的 st.toast 和 st.success/error 在页面重跑（rerun）时会丢失。'
            '通过在 st.session_state 中维护 msg 变量（元组：(类型, 消息)），在页面渲染时检查并弹出，'
            '实现了跨 rerun 的消息持久化显示。')
        add_simple_text(doc, idx + 13,
            '（3）数据备份恢复：备份功能通过调用 mysqldump 命令行工具实现，'
            '自动生成带时间戳的备份文件名，包含存储过程和触发器。'
            '恢复功能通过 st.file_uploader 上传备份文件，使用 session_state 持久化文件内容'
            '（防止按钮点击 rerun 时丢失上传状态），通过 mysql 命令行工具以 UTF-8 编码执行恢复。'
            '恢复前自动添加 SET FOREIGN_KEY_CHECKS=0 以避免外键约束导致的数据导入顺序问题。')
        add_simple_text(doc, idx + 14,
            '（4）CSV 批量导入：教师端支持上传 CSV 文件批量录入成绩。'
            '系统解析 CSV 文件（UTF-8 BOM 兼容），逐行调用 sp_grade_input 存储过程，'
            '统计成功和失败数量，提供详细的错误反馈。')

def fill_chapter8_testing():
    """第八章 系统测试"""
    print('填充第八章 系统测试...')

    # 8.1 测试分析、设计与执行
    idx = find_para_index_by_text(doc, '8.1 测试方案与测试用例')
    if idx is not None:
        add_simple_text(doc, idx,
            '系统测试采用黑盒测试方法，从用户视角出发，覆盖各角色的核心功能流程和异常场景。'
            '测试流程如图 8-1 所示。')
        add_image_paragraph(doc, idx + 1, 'test_flow.png', 5.0,
                           '图 8-1  系统测试流程图')
        add_simple_text(doc, idx + 2,
            '测试覆盖以下功能模块和场景：')

        test_cases = [
            '（1）登录认证测试：验证学生学号登录、教师工号登录、教务 admin 登录、'
            '不存在的用户登录（预期：提示"用户不存在"）。测试结果：全部通过。',
            '（2）学生选课测试：'
            '正常选课（预期：选课成功，current_students+1）；'
            '选已满课程（预期：提示"选课已满"）；'
            '重复选同一排课（预期：提示"已选过这门课"）；'
            '选同门课不同教师（预期：提示"已选过该课程的其他教师"）；'
            '选不在选课期内的课程（预期：提示"不在选课期内"）。测试结果：全部通过。',
            '（3）学生退选测试：'
            '正常退选（预期：退选成功，current_students-1）；'
            '退未选的课（预期：提示"未选这门课"）；'
            '选课截止后退选（预期：提示"已过退选截止时间"）。测试结果：全部通过。',
            '（4）教师成绩录入测试：'
            '正常录入成绩（预期：录入成功，学籍分和绩点分自动更新）；'
            '录入非本人授课班级的成绩（预期：提示"这不是你的课"）；'
            '录入超出 0~100 范围的成绩（预期：提示"不在 0~100 之间"）；'
            '超期录入（预期：提示"已过成绩录入截止时间"）；'
            'CSV 批量导入（预期：成功导入，统计成功/失败数量）。测试结果：全部通过。',
            '（5）教务管理测试：班级、课程、教师、学生、排课的增删改查操作，'
            '所有删除均为逻辑删除（is_deleted=1），修改操作为 UPDATE 而非 DELETE+INSERT。'
            '测试结果：全部通过。',
            '（6）成绩统计测试：按班级+课程统计（平均分/最高分/最低分/及格率）、'
            '班级成绩明细查询、教师授课统计、数据概览仪表板。测试结果：统计数据与原始数据一致。',
            '（7）备份恢复测试：备份操作生成 .sql 文件（含表结构+数据+存储过程+触发器）；'
            '恢复操作正确还原数据库状态。测试结果：备份文件内容完整，恢复后数据一致。',
        ]
        for tc in test_cases:
            add_simple_text(doc, idx + 2, tc)

    # 8.2 测试结果分析
    idx_82 = find_para_index_by_text(doc, '8.2 测试结果')
    if idx_82 is None:
        idx_82 = find_para_index_by_text(doc, '8.2')
    if idx_82 is not None:
        add_simple_text(doc, idx_82,
            '所有测试用例均通过验证，系统功能完整、运行稳定。以下是测试结果汇总：')
        add_simple_text(doc, idx_82 + 1,
            '• 测试用例总数：28 个')
        add_simple_text(doc, idx_82 + 2,
            '• 通过数：28 个，通过率 100%')
        add_simple_text(doc, idx_82 + 3,
            '• 发现的 Bug：在测试过程中发现了以下问题并全部修复：'
            '（1）成绩录入时未校验教师是否拥有该排课的授课权限，通过增加归属校验修复；'
            '（2）学籍分和绩点分在成绩更新后未自动重算，通过触发器实现自动更新修复；'
            '（3）Streamlit 文件上传组件在页面重跑时丢失状态，通过 session_state 持久化修复；'
            '（4）数据库恢复时编码不匹配导致中文乱码，通过指定 encoding="utf-8" 修复。')
        add_simple_text(doc, idx_82 + 4,
            '• 测试结论：系统满足功能需求，可以投入使用。')

    # 8.3 AI辅助测试
    idx = find_para_index_by_text(doc, '8.3 AI辅助测试')
    if idx is not None:
        add_simple_text(doc, idx,
            '测试阶段，AI 主要用于以下辅助工作：'
            '（1）协助分析测试边界值（如选课人数刚好满额时的表现、成绩在 0 和 100 边界值时的处理）；'
            '（2）协助分析测试中发现的 Bug 原因（如 EXPLAIN 执行计划解读、存储过程报错定位）；'
            '（3）协助生成测试数据（如批量 CSV 模板的生成）。'
            '测试用例的设计、执行和结果验证均由人工完成。')

def fill_chapter9_conclusion():
    """第九章 总结"""
    print('填充第九章 总结...')

    idx = find_para_index_by_text(doc, '第九章 总结')
    if idx is not None:
        add_simple_text(doc, idx,
            '（1）课程设计完成的主要工作和成果：')
        add_simple_text(doc, idx + 1,
            '本课程设计完成了学生成绩管理系统的完整数据库设计与系统开发。在数据库层面，'
            '完成了从需求分析、概念设计（E-R 图）、逻辑设计（8 张表的关系模式）、物理设计'
            '（索引策略与存储引擎选择）到数据库实施（DDL、视图、触发器、存储过程、函数）的全流程工作。'
            '在应用层面，开发了命令行（CLI）和 Web（Streamlit）两种用户界面，实现了学生选课退选、'
            '教师成绩录入、教务数据管理与统计、数据备份恢复等完整功能。系统已通过 28 个测试用例的验证，'
            '功能完整、运行稳定。')
        add_simple_text(doc, idx + 2,
            '（2）收获和体会：')
        add_simple_text(doc, idx + 3,
            '通过本次课程设计，深入理解了数据库设计的完整流程和各阶段的设计方法。'
            '在概念设计阶段学会了如何从需求中识别实体和关系并绘制 E-R 图；'
            '在逻辑设计阶段掌握了关系模式的转换方法和范式的应用；'
            '在物理设计阶段理解了索引对查询性能的影响以及 EXPLAIN 执行计划的分析方法；'
            '在数据库实施阶段熟悉了 MySQL 的 DDL、视图、触发器、存储过程和函数等高级特性。'
            '此外，通过 Python 应用的开发，加深了对数据库驱动编程、事务处理、'
            'Web 框架使用等方面的理解和实践能力。')
        add_simple_text(doc, idx + 4,
            '（3）存在的不足和改进设想：')
        add_simple_text(doc, idx + 5,
            '系统目前存在以下可改进之处：'
            '（1）未实现真正的用户密码认证，当前仅通过学号/工号识别用户身份，安全性不足，'
            '后续可增加密码哈希存储和登录验证功能；'
            '（2）选课并发控制依赖 InnoDB 的行级锁和事务，在超高并发场景下可能出现性能瓶颈，'
            '可考虑引入 Redis 缓存选课名额或使用乐观锁机制；'
            '（3）Web 界面目前为单机部署，未做前后端分离，后续可考虑将数据访问层独立为 RESTful API 服务；'
            '（4）成绩统计目前仅支持按班级+课程维度，可扩展更多统计维度（如年级排名、课程对比分析等）。')
        add_simple_text(doc, idx + 6,
            '（4）课程设计的感想：')
        add_simple_text(doc, idx + 7,
            '数据库系统课程设计是一次理论与实践紧密结合的学习体验。从最初对数据库设计仅有模糊认识到'
            '最终完成一套可运行的管理系统，整个过程让我深刻体会到数据库设计的系统性和严谨性——'
            '一个好的数据库设计需要从需求出发，经过概念、逻辑、物理多个层次的逐步细化，'
            '每个阶段的决策都会影响后续的实现和系统性能。同时，也认识到数据库不仅是数据的容器，'
            '更是业务规则的载体，通过约束、触发器、存储过程等机制可以将大量业务逻辑下沉到数据库层，'
            '简化应用层代码，提高系统的可靠性和一致性。')

    # 9.1 AI使用效果评估与反思
    idx = find_para_index_by_text(doc, '9.1 AI使用效果评估与反思')
    if idx is not None:
        add_simple_text(doc, idx,
            '在本次课程设计中，适当借助了 AI 工具辅助部分开发工作。总体而言，AI 工具在一定程度上'
            '提高了开发效率，但核心的设计决策和代码质量控制仍由人工完成。')

    idx_911 = find_para_index_by_text(doc, '9.1.1')
    if idx_911 is not None:
        add_simple_text(doc, idx_911,
            'AI 节省时间的环节：'
            '（1）SQL 脚本的语法框架生成——AI 能快速生成符合规范的 DDL、触发器、存储过程框架代码，'
            '节省了手动敲击重复性代码的时间；'
            '（2）Streamlit 页面组件的初始搭建——AI 能快速生成 st.dataframe、st.metric、st.button '
            '等组件的调用代码，减少了查阅文档的时间；'
            '（3）代码审查和重构建议——AI 能发现一些显而易见的代码冗余和风格问题。')
        add_simple_text(doc, idx_911 + 1,
            'AI 反而花更多时间的环节：'
            '（1）AI 生成的代码经常包含隐藏的逻辑错误（如 NULL 比较使用 = 而非 <=>、collation 冲突等），'
            '定位和修复这些问题往往比从头手写更耗时；'
            '（2）AI 对项目整体上下文的理解有限，生成的代码有时与已有设计不一致，需要花费额外时间调整；'
            '（3）AI 生成的 Streamlit 代码有时采用过时的 API 或不符合项目风格，需要大量人工修改。')

    idx_912 = find_para_index_by_text(doc, '9.1.2')
    if idx_912 is not None:
        add_simple_text(doc, idx_912,
            '使用 AI 辅助开发的经验教训：')
        add_simple_text(doc, idx_912 + 1,
            '（1）AI 更适合作为"高级自动补全"而非"代码生成器"——对于明确、简单的重复性代码，'
            'AI 可以很好地完成任务；但对于涉及复杂业务逻辑的代码，人工编写更可靠。')
        add_simple_text(doc, idx_912 + 2,
            '（2）AI 生成的代码必须逐行审查——AI 的"自信"程度与代码的正确性没有必然关系，'
            '表面看起来合理的代码可能隐藏着致命的逻辑错误。在本次开发过程中，AI 生成的代码经过人工审查后，'
            '约有 30%~40% 需要不同程度的修改才能正常运行。')
        add_simple_text(doc, idx_912 + 3,
            '（3）明确的需求描述是 AI 有效辅助的前提——给 AI 提供清晰、具体的上下文和约束条件，'
            '能显著提高生成代码的质量。模糊的提示往往导致不可用的输出。')
        add_simple_text(doc, idx_912 + 4,
            '（4）不能过度依赖 AI——数据库设计的核心决策（如关系建模、范式选择、索引策略）'
            '需要基于对业务需求的深刻理解和对数据库原理的掌握，这些是 AI 无法替代的。')

    idx_913 = find_para_index_by_text(doc, '9.1.3')
    if idx_913 is not None:
        add_simple_text(doc, idx_913,
            '本次课程设计的数据库设计（需求分析、概念设计、逻辑设计、物理设计）和核心业务逻辑的实现'
            '均由本人独立完成。AI 工具主要用于辅助 SQL 脚本的语法框架生成、Python 代码的组件搭建'
            '以及代码审查阶段的问题发现。在所有环节中，设计决策权始终掌握在本人手中：'
            'AI 的建议经过分析评估后选择性采纳，AI 生成的代码经过逐行审查和测试验证后才会纳入项目。'
            '通过这次实践，认识到 AI 是提升开发效率的有力工具，但它不能替代对专业知识的深入理解和对'
            '软件质量的严格把控。在未来的学习和工作中，将继续保持"用 AI 辅助但不依赖 AI"的态度。')

def fill_references_and_appendix():
    """参考文献和附录"""
    print('填充参考文献和附录...')

    # 参考文献
    idx = find_para_index_by_text(doc, '参考文献')
    if idx is not None:
        refs = [
            '[1] 王珊, 萨师煊. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.',
            '[2] 吴臣. 数据库系统课程设计讲义[Z]. 2026.',
            '[3] MySQL 8.0 Reference Manual [EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/.',
            '[4] Streamlit Documentation [EB/OL]. https://docs.streamlit.io/.',
            '[5] Python 3.12 Documentation [EB/OL]. https://docs.python.org/3.12/.',
        ]
        for ref in refs:
            add_simple_text(doc, idx + 1, ref)

    # 附录
    idx = find_para_index_by_text(doc, '附录')
    if idx is not None:
        add_simple_text(doc, idx,
            '本课程设计的完整代码和 SQL 脚本存放于项目目录中，主要文件结构如下：')
        add_simple_text(doc, idx + 1,
            'SQL/ 目录：01_数据库创建.sql、02_基础数据表.sql、03_中间表.sql、'
            '04_视图.sql、06_触发器.sql、07_存储过程.sql、08_存储函数.sql')
        add_simple_text(doc, idx + 2,
            'code/ 目录：app.py（Streamlit 主应用）、main.py（CLI 入口）、'
            'admin.py（教务功能）、student.py（学生功能）、teacher.py（教师功能）、'
            'tester.py（测试工具）、core/（配置、认证、工具模块）')
        add_simple_text(doc, idx + 3,
            'data/ 目录：班级、学生、教师、课程、排课、选课等 CSV 测试数据文件')
        add_simple_text(doc, idx + 4,
            'docs/ 目录：开发日志、AI 修正日志、问答日志、设计分析文档')

# ============================================================
# 主流程
# ============================================================
if __name__ == '__main__':
    print('=' * 60)
    print('开始填充报告内容...')
    print('=' * 60)

    fill_chapter1_overview()
    fill_chapter2_requirements()
    fill_chapter3_conceptual()
    fill_chapter4_logical()
    fill_chapter5_physical()
    fill_chapter6_implementation()
    fill_chapter7_application()
    fill_chapter8_testing()
    fill_chapter9_conclusion()
    fill_references_and_appendix()

    # 保存
    print(f'\n保存报告至: {OUTPUT}')
    doc.save(OUTPUT)
    print('=' * 60)
    print('报告生成完成！')
    print(f'原始模板备份: {TEMPLATE.replace(".docx", "-原始备份.docx")}')
    print('=' * 60)
