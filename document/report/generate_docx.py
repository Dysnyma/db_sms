"""
生成课程设计说明书 .docx
用法：python document/report/generate_docx.py
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(3)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def set_cn_font(run, font_name):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        set_cn_font(run, '黑体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
        elif level == 3:
            run.font.size = Pt(12)
    return h

def add_para(text, bold=False, indent=False, size=12, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()  # spacing
    return table


# ════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据库系统课程设计')
run.font.size = Pt(26)
run.bold = True
run.font.name = '黑体'
run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('学生成绩管理系统')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
set_cn_font(run, '黑体')

for _ in range(6):
    doc.add_paragraph()

info = [('专    业：', '计算机科学与技术'),
        ('学生姓名：', '蔡坤灿'),
        ('学    号：', ''),
        ('指导教师：', ''),
        ('日    期：', '2026年7月')]
for label, val in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + val)
    run.font.size = Pt(14)
    set_cn_font(run, '宋体')

doc.add_page_break()

# ════════════════════════════════════════════
# 目录页
# ════════════════════════════════════════════
add_heading_styled('目  录', level=1)
toc = ['第一章 概述', '第二章 需求分析', '第三章 概念结构设计',
       '第四章 逻辑结构设计', '第五章 物理设计', '第六章 数据库实施',
       '第七章 系统实现与测试', '第八章 系统运行与维护', '第九章 总结',
       '参考文献', '附录']
for item in toc:
    add_para(item)

doc.add_page_break()

# ════════════════════════════════════════════
# 第一章 概述
# ════════════════════════════════════════════
add_heading_styled('第一章 概述', level=1)

add_heading_styled('1.1 项目背景', level=2)
add_para('本课程设计选题为题目二「学生成绩管理系统」。系统通过分析学校日常教学管理中的课程、选课、学生、班级、教师和成绩等核心数据，设计并实现一个集教师授课管理、学生选课管理、成绩管理与统计等功能于一体的小型数据库应用系统。', indent=True)
add_para('选题要求使用传统的人工设计方法结合AI辅助工具完成数据库设计的六个阶段：需求分析、概念结构设计、逻辑结构设计、物理结构设计、数据库实施、数据库运行与维护。', indent=True)

add_heading_styled('1.2 项目目标', level=2)
for goal in ['深入理解数据库系统理论知识，掌握数据库设计的完整流程',
             '完成功能完整的学生成绩管理系统，含ER图、表结构、视图/触发器/存储过程',
             '使用 Python 分别实现 CLI 和 Streamlit 双界面版本',
             '合理使用 AI 辅助工具，记录 AI 使用过程与人工修正对照']:
    add_para(f'  {goal}', indent=True)

add_heading_styled('1.3 软硬件环境与工具', level=2)
add_table(['类别', '名称', '版本/说明'],
          [['数据库', 'MySQL', '8.0'],
           ['编程语言', 'Python', '3.12'],
           ['CLI框架', '标准库 + pymysql', '—'],
           ['Web框架', 'Streamlit', '1.x'],
           ['AI工具', 'Claude Code', '路由B：AI增强IDE'],
           ['图表库', 'Plotly', '交互式图表'],
           ['数据校验', 'Pydantic', '输入模型验证'],
           ['连接池', 'SQLAlchemy QueuePool', '5池+10溢出'],
           ['操作系统', 'Windows 11', '10.0.26100'],
           ['版本管理', 'Git', '—']])

add_heading_styled('1.4 AI工具使用清单', level=2)
add_para('本课程设计选择「路由B：AI增强IDE全栈辅助开发」方案，使用Claude Code作为AI辅助工具。')
add_table(['设计阶段', 'AI作用'],
          [['需求分析', '阅读任务书，提取实体与业务规则'],
           ['概念结构设计', '讨论实体关系建模决策'],
           ['逻辑结构设计', '辅助编写建表DDL，人工逐行审核修正'],
           ['物理结构设计', '分析查询模式，建议索引策略'],
           ['数据库实施', '生成SQL代码（视图/触发器/SP），人工修复'],
           ['应用层开发', '编写Python代码，辅助调试重构']])

doc.add_page_break()

# ════════════════════════════════════════════
# 第二章 需求分析
# ════════════════════════════════════════════
add_heading_styled('第二章 需求分析', level=1)

add_heading_styled('2.1 系统业务需求', level=2)

add_heading_styled('（1）实体与关系', level=3)
for item in ['一个学生属于一个班级，一个班级有多个学生（1:N）',
             '一个学生可选修多门课程，一门课程可由多个学生选修（M:N），通过选课中间表实现',
             '一门课程可由多个教师教授，一个教师可教授多门课程（M:N），通过讲授表和排课表实现']:
    add_para(f'  • {item}')

add_heading_styled('（2）选课机制', level=3)
for item in ['排课安排有有效选课时间、选课截止时间和成绩录入截止时间',
             '选课人数不能超过上限（max_students），满员时抛出异常',
             '同一名学生不能选择同一门课程中两个不同教师的班级',
             '学生在选课截止时间前可以退选',
             '选课/退选后触发器自动维护 current_students']:
    add_para(f'  • {item}')

add_heading_styled('（3）成绩管理', level=3)
for item in ['教师对其所授课程的选课学生进行成绩录入',
             '成绩录入后触发器自动实时计算学籍分和绩点分（GPA）',
             '成绩录入支持逐条录入和批量CSV录入',
             '成绩录入截止时间后不能再修改']:
    add_para(f'  • {item}')

add_heading_styled('（4）统计报表', level=3)
add_para('所有统计报表使用存储过程实现：')
for item in ['sp_student_roster — 班级学生学籍分名单',
             'sp_class_grade_report — 班级成绩统计（均分/最高/最低/及格率）',
             'sp_student_semester_avg — 学生学期均分',
             'sp_teacher_info / sp_teacher_list — 教师授课统计']:
    add_para(f'  • {item}')

add_heading_styled('（5）角色权限', level=3)
add_table(['角色', '功能'],
          [['学生', '查可选课程、选课、退选、查成绩和学籍分、学期均分'],
           ['教师', '逐条/批量录入成绩、查看所授课程学生名单（含图表）'],
           ['教务管理员', '数据概览、6大实体CRUD、专业管理、排课管理、选课管理、统计报表、备份恢复']])

add_heading_styled('（6）数据维护', level=3)
for item in ['全部采用逻辑删除（is_deleted字段）',
             '学籍分和选课人数使用触发器自动计算',
             '数据库备份和恢复功能（mysqldump --routines --triggers）',
             '班级删除时检查是否有在读学生，防止数据孤儿']:
    add_para(f'  • {item}')

add_heading_styled('2.2 数据字典', level=2)
add_para('数据字典包含全部8张表的字段定义、类型、约束和说明。详见第四章的逻辑结构设计。')

doc.add_page_break()

# ════════════════════════════════════════════
# 第三章 概念结构设计
# ════════════════════════════════════════════
add_heading_styled('第三章 概念结构设计', level=1)

add_heading_styled('3.1 全局ER图', level=2)

add_heading_styled('实体清单', level=3)
for item in ['班级（class）— 学生的归属单位',
             '学生（student）— 选课和成绩的主体',
             '课程（course）— 教学课程',
             '教师（teacher）— 授课教师',
             '排课（course_offering）— 课程与教师的排课安排']:
    add_para(f'  • {item}')

add_heading_styled('关系清单', level=3)
for item in ['选课（enrollment）— 学生与排课的M:N关系',
             '讲授（teacher_course）— 教师与课程的M:N关系',
             '绩点规则（grade_point_rule）— 成绩→绩点映射']:
    add_para(f'  • {item}')

add_para('class ──1:N── student ──M:N── course_offering ──N:1── course / ──N:1── teacher', bold=True)

add_heading_styled('关键设计决策', level=3)
for item in ['学生与班级：1:N — student表加class_id外键，不建中间表',
             '学生与课程：M:N — 通过enrollment中间表实现，同时记录成绩',
             '课程与教师：M:N — 通过teacher_course（讲授资格）和course_offering（实际排课）两层实现',
             '全部逻辑删除 — is_deleted管数据状态，status管业务状态',
             '学籍分和绩点分 — 由触发器自动维护，不通过应用程序计算']:
    add_para(f'  • {item}')

add_heading_styled('3.2 AI辅助概念设计记录', level=2)
add_para('AI帮助确认学生-班级是1:N关系不需中间表；建议使用teacher_course和course_offering双层设计。')
add_table(['AI生成', '人工修改', '理由'],
          [['course_offering含effective_time', '删除', '与enroll_start_time语义重复'],
           ['enrollment同时有enroll_time和create_time', '删除enroll_time', 'create_time就是选课时间'],
           ['遗漏绩点规则表', '新增grade_point_rule', 'GPA计算需要成绩-绩点映射']])

doc.add_page_break()

# ════════════════════════════════════════════
# 第四章 逻辑结构设计
# ════════════════════════════════════════════
add_heading_styled('第四章 逻辑结构设计', level=1)

add_heading_styled('4.1 基础数据表', level=2)
add_para('共设计8张表：class、student、course、teacher、course_offering、enrollment、teacher_course、grade_point_rule。')
add_para('设计原则：表名全小写+下划线、utf8mb4字符集、InnoDB引擎、全部逻辑删除、INT自增主键。')

# 班级表
add_para('表4-1 班级表（class）', bold=True)
add_table(['字段', '类型', '约束', '说明'],
          [['id', 'INT', 'PK, AUTO_INCREMENT', '班级ID'],
           ['name', 'VARCHAR(50)', 'NOT NULL', '班级名'],
           ['grade', 'VARCHAR(10)', 'NOT NULL', '年级'],
           ['major', 'VARCHAR(100)', 'NOT NULL', '专业'],
           ['status', 'TINYINT(1)', 'DEFAULT 1', '1=在读, 0=毕业'],
           ['is_deleted', 'TINYINT(1)', 'DEFAULT 0', '逻辑删除']])

add_para('表4-2 学生表（student）', bold=True)
add_table(['字段', '类型', '约束', '说明'],
          [['id', 'INT', 'PK, AUTO_INCREMENT', '学生ID'],
           ['name', 'VARCHAR(50)', 'NOT NULL', '姓名'],
           ['no', 'VARCHAR(20)', 'NOT NULL, UNIQUE', '学号'],
           ['class_id', 'INT', 'FK→class(id)', '所属班级'],
           ['weighted_score', 'DECIMAL(5,2)', 'DEFAULT 0.00', '学籍分'],
           ['gpa', 'DECIMAL(5,2)', 'DEFAULT 0.00', '绩点分'],
           ['status', 'TINYINT(1)', 'DEFAULT 1', '1=在读, 0=离校'],
           ['is_deleted', 'TINYINT(1)', 'DEFAULT 0', '逻辑删除']])

add_para('表4-3 课程表（course）', bold=True)
add_table(['字段', '类型', '约束', '说明'],
          [['id', 'INT', 'PK, AUTO_INCREMENT', '课程ID'],
           ['name', 'VARCHAR(100)', 'NOT NULL, UNIQUE', '课程名'],
           ['credit', 'DECIMAL(3,1)', 'NOT NULL', '学分'],
           ['status', 'TINYINT(1)', 'DEFAULT 1', '1=开课, 0=停开'],
           ['is_deleted', 'TINYINT(1)', 'DEFAULT 0', '逻辑删除']])

add_para('表4-4 教师表（teacher）', bold=True)
add_table(['字段', '类型', '约束', '说明'],
          [['id', 'INT', 'PK, AUTO_INCREMENT', '教师ID'],
           ['name', 'VARCHAR(50)', 'NOT NULL', '姓名'],
           ['no', 'VARCHAR(20)', 'NOT NULL, UNIQUE', '工号'],
           ['title', 'VARCHAR(50)', 'NULLABLE', '职称'],
           ['phone', 'VARCHAR(20)', 'NULLABLE', '电话'],
           ['status', 'TINYINT(1)', 'DEFAULT 1', '1=在职, 0=离职'],
           ['is_deleted', 'TINYINT(1)', 'DEFAULT 0', '逻辑删除']])

add_para('表4-5 排课表（course_offering）', bold=True)
add_table(['字段', '类型', '约束', '说明'],
          [['id', 'INT', 'PK, AUTO_INCREMENT', '排课ID'],
           ['course_id', 'INT', 'FK→course(id)', '课程'],
           ['teacher_id', 'INT', 'FK→teacher(id)', '教师'],
           ['semester', 'VARCHAR(20)', 'NOT NULL', '学期'],
           ['max_students', 'INT', 'NOT NULL', '选课上限'],
           ['current_students', 'INT', 'DEFAULT 0', '当前人数'],
           ['enroll_start_time', 'DATETIME', 'NOT NULL', '选课开始'],
           ['enroll_end_time', 'DATETIME', 'NOT NULL', '选课截止'],
           ['grade_deadline', 'DATETIME', 'NOT NULL', '成绩截止'],
           ['status', 'TINYINT(1)', 'DEFAULT 1', '1=有效, 0=取消'],
           ['is_deleted', 'TINYINT(1)', 'DEFAULT 0', '逻辑删除']])

add_para('表4-6 选课表（enrollment）', bold=True)
add_table(['字段', '类型', '约束', '说明'],
          [['id', 'INT', 'PK, AUTO_INCREMENT', '选课ID'],
           ['offering_id', 'INT', 'FK, UNIQUE组合', '排课'],
           ['student_id', 'INT', 'FK, UNIQUE组合', '学生'],
           ['score', 'DECIMAL(5,2)', 'NULLABLE', '成绩（NULL=未录入）'],
           ['is_deleted', 'TINYINT(1)', 'DEFAULT 0', '退选标记']])

add_para('表4-7 讲师关系表（teacher_course）', bold=True)
add_table(['字段', '类型', '约束', '说明'],
          [['id', 'INT', 'PK, AUTO_INCREMENT', '关系ID'],
           ['teacher_id', 'INT', 'FK→teacher(id), UNIQUE组合', '教师'],
           ['course_id', 'INT', 'FK→course(id), UNIQUE组合', '课程']])

add_para('表4-8 绩点规则表（grade_point_rule）', bold=True)
add_table(['字段', '类型', '约束', '说明'],
          [['id', 'INT', 'PK, AUTO_INCREMENT', '规则ID'],
           ['min_score', 'DECIMAL(5,2)', 'NOT NULL', '成绩下限'],
           ['max_score', 'DECIMAL(5,2)', 'NOT NULL', '成绩上限'],
           ['point', 'DECIMAL(3,1)', 'NOT NULL', '对应绩点']])

add_para('绩点规则数据：0-59→0, 60-69→1, 70-79→2, 80-89→3, 90-100→4。')

add_heading_styled('4.2 规范化设计说明', level=2)
for item in ['1NF：所有字段均为原子值，无重复组',
             '2NF：每个非主键字段完全依赖于主键，无部分函数依赖',
             '3NF：不存在传递依赖，通过class_id关联class表获取班级信息',
             '约束完整性：全部外键定义FOREIGN KEY约束，学号/工号定义UNIQUE索引']:
    add_para(f'  • {item}')

add_heading_styled('4.3 AI辅助逻辑设计修正对照表', level=2)
add_table(['字段/表名', 'AI原始方案', '人工修正后', '修正理由'],
          [['course_offering.effective_time', '多了一个时间字段', '删除', '与enroll_start_time重复'],
           ['enrollment.enroll_time', '和create_time同时存在', '保留create_time', '语义重复'],
           ['sp_show_courses', '仅过滤已选的plan_id', '新增course_id NOT IN', '同课程不同教师也显示'],
           ['sp_enroll', '未检查同一课程其他教师', '新增第4步校验', '防止选同一课程两个老师'],
           ['teacher_course表', 'AI未明确提及', '人工新增', '排课时需校验讲授资格']])

add_heading_styled('4.4 视图设计', level=2)
add_para('v_student_message — 学生基本信息视图；v_course_plan — 选课安排列表视图（含选课状态）；v_enrollment — 选课详情视图。', indent=True)

add_heading_styled('4.5 存储过程设计', level=2)
add_table(['编号', '存储过程', '参数', '功能'],
          [['SP1', 'sp_show_courses', 'p_student_no', '查询可选课程'],
           ['SP2', 'sp_enroll', 'p_student_no, p_plan_id', '选课（4步校验+事务）'],
           ['SP3', 'sp_unenroll', 'p_student_no, p_plan_id', '退选'],
           ['SP4', 'sp_grade_input', 'p_teacher_no, p_plan_id, p_student_no, p_score', '录入成绩（三重身份校验）'],
           ['SP5', 'sp_student_roster', 'p_class_id', '班级学生学籍分名单'],
           ['SP6', 'sp_class_grade_report', 'p_class_id, p_course_id', '班级成绩统计'],
           ['SP7', 'sp_student_semester_avg', 'p_student_no, p_semester', '学生学期均分'],
           ['SP8', 'sp_teacher_info', 'p_teacher_no', '单个教师授课统计'],
           ['SP9', 'sp_teacher_list', '无', '全部教师授课统计']])

add_heading_styled('4.6 触发器设计', level=2)
add_table(['编号', '触发器', '触发时机', '功能'],
          [['TRG1', 'trg_enrollment_before_insert', 'BEFORE INSERT', '选课前检查名额（FOR UPDATE行锁）'],
           ['TRG2', 'trg_enrollment_after_insert', 'AFTER INSERT', '选课后current_students+1'],
           ['TRG3', 'trg_enrollment_after_update', 'AFTER UPDATE', '退选时current_students-1'],
           ['TRG4', 'trg_enrollment_after_insert_score', 'AFTER INSERT', '带成绩时重算GPA'],
           ['TRG5', 'trg_enrollment_after_update_score', 'AFTER UPDATE', '成绩变化时重算GPA']])

add_para('学籍分公式：weighted_score = Σ(成绩×学分) / Σ(学分)', bold=True)
add_para('GPA公式：gpa = Σ(学分×绩点) / Σ(学分)', bold=True)

add_heading_styled('4.7 函数设计', level=2)
add_para('fn_get_student_id（学号→学生ID）、fn_get_teacher_id（工号→教师ID），供所有存储过程复用。')

doc.add_page_break()

# ════════════════════════════════════════════
# 第五章 物理设计
# ════════════════════════════════════════════
add_heading_styled('第五章 物理设计', level=1)

add_heading_styled('5.1 存储结构', level=2)
add_table(['配置项', '值', '理由'],
          [['存储引擎', 'InnoDB', '支持事务、行级锁、外键约束'],
           ['字符集', 'utf8mb4', '完全支持中文和4字节Unicode字符'],
           ['排序规则', 'utf8mb4_unicode_ci', 'Unicode标准排序'],
           ['主键策略', 'INT AUTO_INCREMENT', '4字节整型，查询效率高']])

add_heading_styled('5.2 索引设计', level=2)
add_table(['表名', '索引字段', '类型', '理由'],
          [['class', 'name, grade, major', '普通索引', '按班级/年级/专业查询'],
           ['student', 'no', '唯一索引', '学号唯一性约束，最频繁查询条件'],
           ['student', 'class_id', '普通索引', '按班级查询学生名单'],
           ['course', 'name', '唯一索引', '课程名唯一性约束'],
           ['teacher', 'no', '唯一索引', '工号唯一性约束'],
           ['course_offering', 'teacher_id', '普通索引', '按教师查排课（高频）'],
           ['course_offering', 'semester', '普通索引', '按学期筛选'],
           ['enrollment', '(offering_id, student_id)', '唯一索引', '防重复选课'],
           ['enrollment', 'student_id', '普通索引', '按学生查选课（高频）']])

add_heading_styled('5.3 AI辅助物理设计记录', level=2)
add_para('AI分析了高频查询模式后建议索引策略，经EXPLAIN验证全部采纳，type均为ref或const级别，无全表扫描。')

doc.add_page_break()

# ════════════════════════════════════════════
# 第六章 数据库实施
# ════════════════════════════════════════════
add_heading_styled('第六章 数据库实施', level=1)

add_heading_styled('6.1 数据库创建', level=2)
add_para('SQL脚本执行顺序：01_数据库创建.sql → 02_基础数据表.sql → 03_中间表.sql → 04_视图.sql → 06_触发器.sql → 07_存储过程.sql → 08_存储函数.sql')
add_para('执行方式：python src/reset_data.py')

add_heading_styled('6.2 初始化数据', level=2)
add_para('使用 data/generate_data.py 生成基础仿真数据，通过 src/import_data.py 导入。')
add_table(['CSV文件', '数据量'],
          [['class.csv', '24条'],
           ['student.csv', '1000条'],
           ['teacher.csv', '25条'],
           ['course.csv', '18条'],
           ['enrollment.csv', '~3000条']])

add_heading_styled('6.3 大规横数据生成', level=2)
add_para('对于性能测试场景，提供 test/generate_large_data.py 一站式生成器：')
add_table(['实体', '规模'],
          [['学生', '20,000人'],
           ['教师', '300人'],
           ['班级', '329个'],
           ['排课', '1,004条'],
           ['选课', '2,002,688条']])
add_para('使用 LOAD DATA LOCAL INFILE 极速导入，配合预删除触发器 + 批量GPA重算策略，200万条选课导入耗时约30秒。')

add_heading_styled('6.4 AI辅助编码与人工审核记录', level=2)
add_table(['修正类型', '问题', '修正'],
          [['Bug', 'sp_enroll未检查同一门课不同老师', '新增JOIN过滤'],
           ['Bug', '1267 collation冲突', '统一utf8mb4_unicode_ci'],
           ['优化', 'tester.py 175行与角色代码重复', '精简为83行跳板'],
           ['优化', '备份缺少routines/triggers', '新增--routines --triggers'],
           ['优化', '输入校验缺失', '全部表单增加Pydantic校验+max_chars'],
           ['优化', 'st.tabs+st.form组件干扰', '改为st.radio+st.button']])

doc.add_page_break()

# ════════════════════════════════════════════
# 第七章 系统实现与测试
# ════════════════════════════════════════════
add_heading_styled('第七章 系统实现与测试', level=1)

add_heading_styled('7.1 系统架构', level=2)
add_para('系统采用双界面架构——命令行界面（CLI）和 Web 界面（Streamlit），共享同一套数据库和业务函数。')
add_para('项目结构：')
for item in ['src/app.py — Streamlit Web入口（22个页面）',
             'src/pages/admin_page.py — 教务14页面+Plotly图表',
             'src/pages/teacher.py — 教师3页面+成绩饼图/排行',
             'src/pages/student.py — 学生5页面+各科成绩柱状图',
             'sql/ — 数据库DDL/DML（7个文件）',
             'data/ — CSV数据（含专业列表26个）',
             'test/ — 大规横数据生成器']:
    add_para(f'  • {item}')

add_heading_styled('7.2 功能模块', level=2)

add_heading_styled('学生端（5功能）', level=3)
add_table(['功能', '实现方式'],
          [['查询可选课程', 'sp_show_courses → st.dataframe表格'],
           ['选课', 'selectbox选择+按钮确认'],
           ['退选', 'selectbox选择+按钮确认'],
           ['查看我的成绩', 'st.dataframe + 各科成绩柱状图（Plotly，含及格线）'],
           ['学期均分查询', 'selectbox下拉学期选择 → 存储过程']])

add_heading_styled('教师端（3功能）', level=3)
add_table(['功能', '实现方式'],
          [['录入成绩', 'selectbox选排课→输入学号+成绩+按钮'],
           ['批量CSV录入', 'st.file_uploader上传→循环sp_grade_input'],
           ['查看课程学生', 'selectbox选排课→成绩分布饼图+排行柱状图+未录入折叠名单+统计指标']])

add_heading_styled('教务端（16功能）', level=3)
add_para('查询统计类（6个）：数据概览（含Plotly图表）、班级学生名单（含学籍分饼图）、班级成绩统计、班级成绩明细（含课程均分柱状图）、教师信息、教师列表（含排课/学生排行）。')
add_para('数据管理类（8个）：班级管理（年级/专业下拉+自动生成班级名）、课程管理、教师管理（职称下拉）、学生管理、排课管理（datetime_input）、选课管理（按学号+学期查询）、专业管理（CSV存储）、选课管理（输入退选ID）。')
add_para('系统工具类（2个）：备份数据（mysqldump）、恢复数据。')

add_heading_styled('7.3 输入校验', level=2)
add_para('全表单使用 Pydantic 模型进行统一校验，搭配 st.toast 显示中文错误提示。')
add_para('text_input 全部配置 max_chars 参数，与后端模型长度限制对齐。')
add_para('登录输入区分三种格式：教务(admin)、教师工号(T+数字)、学生学号(纯数字8-12位)。')
add_para('排课上限限制最多5位字符（text_input + max_chars=5），并配合后端 int() 转换校验。')

add_heading_styled('7.4 测试结果', level=2)
add_para('业务规则测试全部通过：选课名额满员异常、同课异师拦截、current_students实时更新、GPA自动重算。')
add_para('数据规模：班级24个→329个，学生1000人→20000人，选课2990条→2002688条。')
add_para('并发安全：FOR UPDATE行级锁保证选课名额不超卖，UNIQUE索引防重复选课，事务保证原子性。')

doc.add_page_break()

# ════════════════════════════════════════════
# 第八章 系统运行与维护
# ════════════════════════════════════════════
add_heading_styled('第八章 系统运行与维护', level=1)

add_heading_styled('8.1 备份与恢复', level=2)
add_para('备份命令：mysqldump -u root --databases db_sms --routines --triggers --add-drop-table')
add_para('恢复：使用 FOREIGN_KEY_CHECKS=0 包裹导入，避免外键冲突。')
add_para('CLI和Streamlit双界面均提供备份和恢复功能，备份文件自动以时间戳命名保存到 backup/ 目录。')

add_heading_styled('8.2 数据独立性演示', level=2)
add_para('通过修改触发器实现 GPA 计算逻辑变更，不修改任何应用层代码——体现了数据独立性。')

doc.add_page_break()

# ════════════════════════════════════════════
# 第九章 总结
# ════════════════════════════════════════════
add_heading_styled('第九章 总结', level=1)

add_heading_styled('9.1 课程设计成果', level=2)
add_table(['成果', '详情'],
          [['数据库设计', '8张表、3个视图、5个触发器、9个存储过程、2个函数'],
           ['CLI应用', '4角色命令行界面，约1000行Python代码'],
           ['Web应用', '22个Streamlit页面，含Plotly交互式图表'],
           ['测试数据', '20000学生/300教师/329班级/200万选课记录'],
           ['AI辅助', 'Claude Code全流程辅助，详细记录AI使用过程']])

add_heading_styled('9.2 收获与体会', level=2)
for item in ['深入理解数据库设计完整6阶段流程',
             '掌握ER图设计、规范化理论、索引优化、视图/触发器/存储过程的实际应用',
             '体验了数据独立性的实际意义',
             '学会版本管理（Git）和AI协作方法']:
    add_para(f'  • {item}')

add_heading_styled('9.3 AI工具使用效果与反思', level=2)
add_para('AI有效应用：DDL/SQL生成节省约70%时间，CRUD代码节省约80%，Bug定位快速。')
add_para('AI局限：对Streamlit组件行为理解不够深入，st.form方案经历3轮迭代。')
add_para('反思：AI是强大辅助工具，但核心设计决策必须由人工主导，AI生成代码必须逐行审核。')

# ════════════════════════════════════════════
# 参考文献
# ════════════════════════════════════════════
add_heading_styled('参考文献', level=1)
for ref in ['[1] 数据库系统概论[M].',
            '[2] 《数据库系统课程设计》任务书-2026-任.docx',
            '[3] MySQL 8.0 Reference Manual',
            '[4] Streamlit API Reference',
            '[5] pymysql Documentation',
            '[6] Draw.io (diagrams.net)',
            '[7] 数据库系统课程设计 说明书规范-任2026.docx']:
    add_para(ref)

# ════════════════════════════════════════════
# 保存
# ════════════════════════════════════════════
output = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      '《数据库系统课程设计》说明书.docx')
doc.save(output)
print(f'✅ 说明书已保存: {output}')
