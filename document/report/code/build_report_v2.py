# -*- coding: utf-8 -*-
"""完整报告 v2 —— 基于任务书要求全新撰写"""
import sys,io; sys.stdout=io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn; from docx.oxml import OxmlElement
import os

SRC = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿-原始备份.docx'
OUT = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书-蔡坤灿-v2.docx'
IMG = r'E:\02_Courses\db_lab\document\报告\images'

doc = Document(SRC)

# ====== 清理模板占位符 ======
P = [
    '正文排版要求','字号，5号','图和表，要有标号','图表和正文的引用',
    '简述各阶段分别用AI做了什么','提示词/Prompt','截图或文字引用',
    'AI遗漏或理解错误了哪些内容','AI生成的实体/关系中，手动调整',
    '基本表及其属性、数据类型及长度','表4-1 XXX表',
    'CHECK约束、冗余列、冗余表的数据来源说明',
    '截图','贴执行计划截图','最终SQL脚本',
    '明确指定用触发器/存储过程/视图等','保留原样粘贴',
    '写了哪些Bug、如何改的','是否使用AI生成测试数据',
    '课程设计中遇到的主要问题','创新和得意之处',
    '课程设计中存在的不足','课程设计的感想和心得体会',
    '本项目中哪些核心设计环节必须由人完成',
    '这里如果是纯数据库实现','表4-2 AI辅助逻辑设计审核对照表',
    '豆包最开始把学生','我最初画的ER图',
    '一个学生只有一张消费卡','有特殊处理比如增加冗余列降低了范式等需要该章节',
]
td = []
for i,p in enumerate(doc.paragraphs):
    s = p.style.name
    if s == 'annotation text': td.append(i); continue
    if s not in ('Heading 1','Heading 2','Heading 3','toc 1','toc 2','toc 3'):
        for m in P:
            if m in p.text: td.append(i); break
for i in reversed(td): doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
print(f'清理 {len(td)} 段')

td2 = []
for ti,t in enumerate(doc.tables):
    h = ''.join(c.text for c in t.rows[0].cells)
    if '字段名称' in h and '字段类型' in h: td2.append(ti)
    if 'AI生成的原始定义' in h: td2.append(ti)
for ti in reversed(td2): doc.tables[ti]._tbl.getparent().remove(doc.tables[ti]._tbl)
print(f'清理 {len(td2)} 表格')

# ====== 工具函数 ======
def fh(text):
    for i,p in enumerate(doc.paragraphs):
        if text in p.text and p.style.name.startswith('Heading'): return i
    return None
def ft(text):
    for i,p in enumerate(doc.paragraphs):
        if text in p.text: return i
    return None

def _el(text, cn='宋体', en='Times New Roman', sz=Pt(10.5), bold=False,
        align=None, indent=True, lsp=1.25, sb=0, sa=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.line_spacing = lsp
    pf.space_before = Pt(sb); pf.space_after = Pt(sa)
    pf.first_line_indent = Pt(21) if indent else Pt(0)
    if align is not None: p.alignment = align
    if text:
        r = p.add_run(text); r.font.size = sz; r.font.name = en; r.bold = bold
        rPr = r._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts')
        rf.set(qn('w:eastAsia'),cn); rf.set(qn('w:ascii'),en); rf.set(qn('w:hAnsi'),en)
        rPr.insert(0,rf)
    return p._element

def ap(ref, text, **kw):
    e = _el(text,**kw); e.getparent().remove(e)
    doc.paragraphs[ref]._element.addnext(e)

def an(ref, texts, **kw):
    c = ref
    for t in texts: ap(c,t,**kw); c+=1

def ah2(ref, text):
    e = _el(text,cn='黑体',en='Arial',sz=Pt(14),bold=True,indent=False,sb=12,sa=6)
    e.getparent().remove(e)
    pPr = e.find(qn('w:pPr'))
    if pPr is None: pPr=OxmlElement('w:pPr'); e.insert(0,pPr)
    ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'),'Heading2'); pPr.append(ps)
    doc.paragraphs[ref]._element.addnext(e)

def ah3(ref, text):
    e = _el(text,cn='黑体',en='Arial',sz=Pt(12),bold=True,indent=False,sb=6,sa=4)
    e.getparent().remove(e)
    pPr = e.find(qn('w:pPr'))
    if pPr is None: pPr=OxmlElement('w:pPr'); e.insert(0,pPr)
    ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'),'Heading3'); pPr.append(ps)
    doc.paragraphs[ref]._element.addnext(e)

def aimg(ref, name, w=5.0, cap=None):
    path = os.path.join(IMG,name)
    if not os.path.exists(path): print(f'  [!]{name}'); return
    re = doc.paragraphs[ref]._element
    ip = doc.add_paragraph(); ip.alignment=WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.first_line_indent=Pt(0)
    r=ip.add_run()
    try:
        from PIL import Image; im=Image.open(path); a=im.size[1]/im.size[0]
        r.add_picture(path,width=Inches(w),height=Inches(w*a))
    except: r.add_picture(path,width=Inches(w))
    ie=ip._element; ie.getparent().remove(ie); re.addnext(ie)
    if cap:
        cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent=Pt(0)
        cr=cp.add_run(cap); cr.font.size=Pt(9); cr.font.name='Times New Roman'
        rPr=cr._r.get_or_add_rPr(); rf=OxmlElement('w:rFonts')
        rf.set(qn('w:eastAsia'),'宋体'); rPr.insert(0,rf)
        ce=cp._element; ce.getparent().remove(ce); ie.addnext(ce)
    print(f'  [img] {name}')

def atbl(ref, hdrs, rows, cw, caption=None):
    re = doc.paragraphs[ref]._element
    if caption:
        ce=_el(caption,sz=Pt(9),bold=True,indent=False,align=WD_ALIGN_PARAGRAPH.CENTER)
        ce.getparent().remove(ce); re.addnext(ce); re=ce
    nc=len(hdrs); nr=len(rows)+1
    t=doc.add_table(rows=nr,cols=nc,style='Table Grid')
    t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for ci,h in enumerate(hdrs):
        c=t.rows[0].cells[ci]; c.text=''
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(h); r.font.size=Pt(9); r.font.name='Times New Roman'; r.bold=True
        rPr=r._r.get_or_add_rPr(); rf=OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'),'宋体'); rPr.insert(0,rf)
        sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),'1565C0'); sh.set(qn('w:val'),'clear')
        c._tc.get_or_add_tcPr().append(sh); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=''
            p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(str(val)); r.font.size=Pt(8.5); r.font.name='Times New Roman'
            rPr=r._r.get_or_add_rPr(); rf=OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'),'宋体'); rPr.insert(0,rf)
    if cw:
        for ri in range(nr):
            for ci,w in enumerate(cw): t.rows[ri].cells[ci].width=Inches(w)
    te=t._tbl; te.getparent().remove(te); re.addnext(te)
    return t

def acode(ref, code):
    e=_el(code,en='Consolas',sz=Pt(7.5),indent=False,lsp=1.1)
    e.getparent().remove(e); doc.paragraphs[ref]._element.addnext(e)

def aquote(ref, text):
    """引用块——用于 AI 提示词示例"""
    e=_el(text,sz=Pt(10),indent=False,lsp=1.15)
    e.getparent().remove(e)
    # 加左边距和灰色
    pPr=e.find(qn('w:pPr'))
    if pPr is None: pPr=OxmlElement('w:pPr'); e.insert(0,pPr)
    ind=OxmlElement('w:ind'); ind.set(qn('w:left'),'420'); pPr.append(ind)
    doc.paragraphs[ref]._element.addnext(e)

# ====== 正文 ======

def ch1():
    print('Ch1...')
    i=fh('1.1 开发背景')
    if i:
        an(i,[
            '随着高校招生规模持续扩大，传统纸质成绩管理方式效率低下、易出错，难以满足教务管理的信息化需求。学生成绩管理涉及选课排课、成绩录入、绩点计算、统计分析等多个环节，数据量大、关联复杂，亟需一套规范化的数据库系统进行统一管理。',
            '本课程设计选题为"学生成绩管理系统"（Student Grade Management System，数据库名 db_sms），对应任务书题目二。旨在通过完整的数据库设计流程——需求分析、概念设计、逻辑设计、物理设计、数据库实施与系统开发——将数据库理论知识与工程实践相结合。系统采用 MySQL 8.0 作为 DBMS，使用 Python 语言开发命令行和基于 Streamlit 框架的 Web 双模式界面，支持学生、教师、教务管理员三种角色。',
        ])
    i=fh('1.2 设计目标')
    if i:
        an(i,[
            '基于任务书 7.2 节的基本需求描述，本系统的设计目标如下：',
            '（1）完成学生成绩管理系统的完整数据库设计，覆盖从 E-R 概念模型到物理表结构的全部设计阶段。',
            '（2）实现核心业务功能：学生选课与退选（含名额实时检查和时效控制）、教师成绩录入与批量导入、教务数据管理与统计分析、学籍分和绩点分的自动计算更新。',
            '（3）通过主键约束、外键约束、唯一约束、触发器和事务机制等多层机制，保证数据的完整性和一致性。',
            '（4）采用逻辑删除（is_deleted 字段）策略，保留数据痕迹；status 字段独立管理业务状态。',
            '（5）设计必要的视图实现外模式，设计触发器维护冗余数据和自定义完整性，设计存储过程封装复杂业务逻辑和统计报表。',
            '（6）支持命令行（CLI）和 Web（Streamlit）双模式运行，覆盖三种用户角色。',
            '（7）实现数据库备份与恢复功能。',
        ])
    i=fh('1.3 软硬件环境与工具')
    if i:
        ap(i,'系统开发与运行的软硬件环境配置如表 1-1 所示。')
        atbl(i+1,
            ['类别','名称','版本/说明'],
            [['操作系统','Windows 11 Home China','24H2'],
             ['数据库管理系统','MySQL','8.0'],
             ['开发语言','Python','3.12'],
             ['Web 框架','Streamlit','1.x'],
             ['数据库驱动','PyMySQL','最新版'],
             ['IDE','VS Code','—'],
             ['ER 图绘制','matplotlib','3.x'],
             ['备份工具','mysqldump','MySQL 自带']],
            [1.6,2.0,1.5], caption='表 1-1  开发环境配置')
        ap(i+2,'系统采用分层架构设计，如图 1-1 所示。')
        aimg(i+3,'system_architecture.png',5.0,'图 1-1  系统四层架构图')

    # -------- 1.4 AI --------
    i=fh('1.4 AI辅助工具清单')
    if i:
        ap(i,'根据任务书附录一"AI辅助工具链"的要求，本课程设计选择了规定的 AI 使用路径，并在各阶段遵循"先判断、再告知、后记录"的操作流程。以下说明路径选择、所用工具和各阶段使用策略。')
    i=fh('1.4.1')
    if i:
        an(i,[
            '任务书附录一给出了两种 AI 使用路径。路径 A 为"专业 AI 数据库工具链"——使用 itBuilder、ERD Online、Chat2DB 等面向数据库设计场景的专用 AI 工具，侧重于数据库设计本身的质量和规范性。路径 B 为"AI 编程 IDE 全栈开发工具链"——使用 Cursor、GitHub Copilot、通义灵码等通用型 AI 编程助手，侧重于数据库和前后端应用的全栈开发。',
            '我选择的是路径 B。选择理由：本次课程设计不仅涉及数据库设计，还需要使用 Python + Streamlit 完成 Web 应用开发，路径 B 的通用型 AI 编程助手更适合这种全栈场景。同时，路径 B 的核心理念——"AI 生成代码后人工逐段审查、测试和修改"——与本课程"主体采用传统工具链人工主导完成"的要求一致。',
            '使用的 AI 工具为 Claude Code（Anthropic 公司开发，基于 Claude Opus 4.8 模型，通过 VS Code 集成终端使用）。该工具定位与路径 B 中列举的 Cursor、GitHub Copilot 类似，在代码语法检查、框架生成、代码审查和文档查阅方面提供智能辅助。此外，在部分概念性问题上使用了 Claude 网页对话版进行资料查阅。',
        ])
    i=fh('1.4.2')
    if i:
        ap(i,'任务书要求使用 AI 辅助设计时必须遵循"先判断、再告知、后记录"的操作流程，即：在向 AI 提问之前，先自行确定该功能应使用哪种数据库对象（触发器、存储过程、视图还是函数），然后在提示词中明确指定，禁止使用模糊指令让 AI 替代自己做技术决策。整个课程设计过程中，我严格遵循了这一原则。各阶段的具体使用策略如下：')
        an(i+1,[
            '需求分析阶段：业务需求主要来源于对学校教务流程的实际了解以及任务书 7.2 节给出的详细描述，AI 仅辅助整理了数据字典的格式。',
            '概念设计阶段：实体识别和关系分析由自己根据需求文档完成。AI 在 ER 图标准符号规范方面提供了参考，并提醒我注意到了"成绩和绩点的转换关系"这一隐含需求（由此增加了 grade_point_rule 实体）。最终 ER 图由自己使用 matplotlib 绘制。',
            '逻辑设计阶段：先按照教材转换规则将 ER 图转换为 8 张表的关系模式并手写 DDL，然后让 AI 辅助检查语法规范性（字段类型是否合理、约束定义是否完整、字符集设置是否正确）。关键的设计决策（如关系模式的确定、外键的设置、索引的初步规划）由自己做出。',
            '物理设计阶段：索引方案由自己根据系统查询模式逐项分析确定。对于关键查询，先用 EXPLAIN 分析执行计划，再将 EXPLAIN 输出和原始 SQL 一起提供给 AI 帮助解读。AI 给出的两个索引建议（为 status 建索引、为 name 建全文索引）经分析后未采纳，理由在第五章详述。',
            '数据库实施阶段：严格遵循前置决策原则。对于选课、退选、成绩录入等需要事务控制的操作，先判断应使用存储过程，再在提示词中明确指定"创建存储过程 sp_xxx"。对于名额自动检查、成绩自动计算等由数据变更触发的工作，先判断应使用触发器，再明确指定触发类型和时机。AI 生成的代码约有 30%~40% 在后续调试中被修改或重写。',
            '系统实现阶段：Python 应用代码整体由自己编写。AI 帮助快速上手 Streamlit 框架（了解 st.navigation、session_state 等组件的用法），但页面布局、交互逻辑和权限控制由自己设计和实现。',
            '测试阶段：测试用例全部由自己根据业务规则设计。AI 仅辅助生成了部分 CSV 批量测试数据文件。',
        ])

def ch2():
    print('Ch2...')
    i=fh('2.1 系统业务流程')
    if i:
        ap(i,'根据任务书 7.2.1 节的基本需求描述，学生成绩管理系统面向三类用户角色，核心业务需求分析如下。')
        ap(i+1,'一、学生端需求',bold=True,indent=False)
        an(i+2,[
            '（1）浏览可选课程：查看当前学期在选课期内、尚有名额的排课信息，包括课程名、学分、授课教师、剩余名额等。',
            '（2）选课：在选课开放时间内选择某门课程的一个排课（同一门课程只能选一个教师），系统自动检查名额是否已满、是否重复选课、是否已选同门课的其他教师。',
            '（3）退选：在选课到期时间之前，可以退选已选课程并重新选课，退选后名额自动释放。',
            '（4）成绩查询：查看个人所有已修课程的成绩、学籍分（加权平均分）和绩点分（GPA）。',
            '（5）学期均分查询：按学期统计个人各课程的平均成绩。',
        ])
        j=ft('一、学生端需求')
        if j:
            ap(j+5,'二、教师端需求',bold=True,indent=False)
            an(j+6,[
                '（1）成绩录入：对自己所授课程班级的学生录入成绩（0~100 分）。录入时需校验教师身份、排课归属和成绩录入截止时间。支持逐条录入和 CSV 文件批量导入两种方式，批量导入时一次保存。',
                '（2）查看课程学生：查看自己所授排课的学生名单及成绩录入状态（已录入/未录入）。',
            ])
            ap(j+9,'三、教务管理员端需求',bold=True,indent=False)
            an(j+10,[
                '（1）数据概览：查看系统中班级数、学生数、教师数、课程数、排课数和选课记录数的统计概览。',
                '（2）基础数据管理：对班级、课程、教师、学生信息进行增删改查操作。所有删除采用逻辑删除，不物理删除数据。',
                '（3）排课管理：创建、修改、删除选课安排。为每次排课指定课程、授课教师、学期、选课人数上限，以及选课开始时间、选课截止时间和成绩录入截止时间三个关键时间节点。',
                '（4）选课管理：查看全部选课记录，支持管理员对异常选课进行强制退选。',
                '（5）成绩统计与报表：按班级和课程统计平均分、最高分、最低分、及格率（通过存储过程实现）；按班级输出学生成绩明细；按班级输出含学籍分的学生名单；查询教师个人及全部教师的授课统计信息。',
                '（6）数据备份与恢复：一键将数据库（含表结构、数据、存储过程和触发器）备份到指定文件，以及从备份文件恢复数据。',
            ])
    i=fh('2.2 数据流图')
    if i:
        ap(i,'系统顶层数据流图（DFD）如图 2-1 所示，展示了学生、教师、教务管理员三个外部实体与系统核心处理过程之间的数据流动关系。')
        aimg(i+1,'data_flow_diagram.png',5.0,'图 2-1  系统数据流图（DFD）')
        aimg(i+2,'function_modules.png',5.0,'图 2-2  系统功能模块图')
    i=fh('2.3 数据字典')
    if i:
        ap(i,'系统核心数据项定义如下：')
        an(i+1,[
            '（1）学生（Student）：学号（唯一标识）、姓名、所属班级、学籍分（加权平均分）、绩点分（GPA）、在读状态。',
            '（2）班级（Class）：班级名称、年级、专业、在读/毕业状态。',
            '（3）课程（Course）：课程名称（唯一）、学分、开课/停开状态。',
            '（4）教师（Teacher）：工号（唯一标识）、姓名、职称、联系电话、在职/离职状态。',
            '（5）选课安排（Course Offering）：关联某门课程和某位授课教师，指定学期、选课人数上限、当前已选人数（由触发器自动维护），以及选课开始时间、选课截止时间和成绩录入截止时间。',
            '（6）选课记录（Enrollment）：记录学生与选课安排之间的关联。score 字段存储成绩（NULL 表示尚未录入），通过逻辑删除实现退选功能。设有 UNIQUE(offering_id, student_id) 约束防止同一学生对同一排课重复选课。',
            '（7）讲授关系（Teacher Course）：教师与课程的 M:N 关系表，记录每位教师有资格讲授哪些课程。UNIQUE(teacher_id, course_id) 保证同一教师对同一课程只有一条资格记录。',
            '（8）绩点规则（Grade Point Rule）：定义成绩区间与绩点的映射关系（0~59.99→0、60~69.99→1、70~79.99→2、80~89.99→3、90~100→4），来源于任务书表 2-1。用于 GPA 的自动计算。',
        ])

    # 2.4 AI需求分析
    i=fh('2.4 AI辅助需求分析记录')
    if i:
        ap(i,'需求分析阶段，业务需求主要来自两个方面：一是任务书 7.2.1 节给出的详细基本需求描述——包括学生选课规则、成绩计算方式、角色划分等，这些需求已经相当具体和明确；二是自己对学校教务管理流程的实际了解。AI 在此阶段的参与程度很低，仅用于将整理好的数据项按统一格式排列成数据字典，属于纯格式整理工作。核心需求的分析、理解和确认均由人工完成。')
    i=fh('2.4.1')
    if i:
        ap(i,'在整理数据字典阶段，向 AI 提出的提示词如下：')
        aquote(i,'"以下是我整理的学生成绩管理系统的数据项清单，请帮我校对数据字典的格式规范性，确保每个数据项的描述包含：中文名称、英文名称、唯一标识说明、关键属性和状态说明。"')
    i=fh('2.4.2')
    if i:
        ap(i,'AI 返回了格式统一的数据字典条目，每条包含中英文名称和关键属性描述，格式规范。但 AI 的输出中存在以下需要人工调整的地方：AI 对"选课安排"实体的理解不够准确——它将其简单地描述为"课程的上课安排"，遗漏了名额控制和三个时间节点（选课开始、选课截止、成绩录入截止）这些关键的业务属性。这些内容由人工根据任务书第 120~125 条需求进行了补充。')
    i=fh('2.4.3')
    if i:
        an(i,[
            'AI 在此阶段最明显的不足是对业务细节的遗漏，以下由人工补充和修正：',
            '（1）选课安排实体的时间维度——AI 未意识到一次排课需要同时设置选课开始时间、选课截止时间和成绩录入截止时间三个独立的时间节点，人工根据任务书补充了这三个时间属性。',
            '（2）成绩与绩点的转换——AI 未提及需要绩点对照表来支持 GPA 计算。人工根据任务书 7.2.2 节"运维需求变更"中关于绩点分的要求，在数据字典中增加了 grade_point_rule 实体。',
            '（3）退选的逻辑删除实现——AI 将退选描述为"删除选课记录"，而任务书明确要求使用逻辑删除来保留数据痕迹。人工将退选的设计改为通过 is_deleted 字段实现。',
        ])

def ch3():
    print('Ch3...')
    i=fh('3.1 全局ER图')
    if i:
        ap(i,'经过需求分析，识别出系统包含以下核心实体：班级（Class）、学生（Student）、课程（Course）、教师（Teacher）、选课安排（Course Offering）、选课记录（Enrollment）、讲授关系（Teacher Course）和绩点对照规则（Grade Point Rule）。各实体间的关系分析如下：')
        an(i+1,[
            '• 班级与学生：1:N。一个班级包含多名学生，一名学生只属于一个班级。学生表通过 class_id 外键关联班级。',
            '• 课程与选课安排：1:N。一门课程可以有多次排课（不同学期、由不同教师讲授），但一次排课只对应一门课程。',
            '• 教师与选课安排：1:N。一位教师可以有多次排课，但一次排课只由一位教师授课。',
            '• 选课安排与选课记录：1:N。一次排课可被多名学生选修（在人数上限内），一条选课记录对应一次排课。',
            '• 学生与选课记录：1:N。一名学生可以选修多门课程（多条选课记录），一条选课记录属于一名学生。',
            '• 教师与课程（讲授关系）：M:N。一位教师可以讲授多门课程，一门课程也可以由多位教师讲授。通过 teacher_course 中间表实现，UNIQUE(teacher_id, course_id) 约束保证无重复。',
        ])
        ap(i+7,'全局 E-R 图如图 3-1 所示。')
        aimg(i+8,'er_diagram.png',5.5,'图 3-1  学生成绩管理系统全局 E-R 图')

    i=fh('3.2 AI辅助概念设计记录')
    if i:
        ap(i,'概念设计阶段，主要工作——实体识别、属性分析和关系确定——由自己基于需求分析的结果独立完成。AI 在此阶段提供了两方面的辅助：ER 图符号规范的确认，以及实体完整性方面的补充建议。')
    i=fh('3.2.1')
    if i:
        ap(i,'在绘制 ER 图初稿时，我先根据需求文档自己列出了实体清单并初步确定了它们之间的关系，然后用 matplotlib 开始绘图。在绘图过程中，向 AI 咨询了 ER 图的标准符号规范（如"矩形表示实体、菱形表示关系、椭圆表示属性"这一基本规则是否正确），确认自己的画法符合规范。AI 在浏览我的实体清单后，提出了一个我没有考虑到的点：成绩和绩点的对应关系是否需要单独的实体来表示？这促使我重新审视任务书表 2-1 的内容，最终增加了 grade_point_rule 实体。最终 ER 图由自己使用 matplotlib 逐一绘制完成。')
    i=fh('3.2.2')
    if i:
        an(i,[
            '在概念设计阶段，我对 AI 的建议进行了以下人工审核和调整：',
            '（1）实体关系基数——AI 初始给出的某些关系基数不够准确。例如选课安排与选课记录的关系，AI 最初描述为 N:1，但实际场景是一次排课可容纳多名学生选课（在名额限制内），应为 1:N。人工根据任务书第 119~126 条的业务描述进行了纠正。',
            '（2）教师与课程的讲授关系——AI 建议直接在教师表中增加"可授课程"字段来实现。但教师与课程是 M:N 关系，在教师表中嵌入多值字段违反第一范式。人工按教材规范将讲授关系抽取为独立的 teacher_course 中间表。',
            '（3）绩点对照表的增加——如 3.2.1 所述，这一实体是在 AI 的提醒下补充的，但表的结构设计和数据填充（5 个成绩区间）由人工根据任务书表 2-1 完成。',
        ])

def ch4():
    print('Ch4...')
    i=fh('4.1 基本数据表')
    if i:
        ap(i,'根据概念结构设计的结果，将 E-R 图转换为以下关系模式（数据库表）。所有表统一采用 InnoDB 存储引擎，字符集 utf8mb4，排序规则 utf8mb4_unicode_ci。全部实现逻辑删除（is_deleted 字段），status 字段独立管理业务状态。对于多属性组合主键的情况，统一增加自增 ID 列作为主键，将实际业务主键通过 UNIQUE 索引设为候选码。各表结构汇总如表 4-1 所示。')
        atbl(i+1,
            ['表名','主要字段','主键 / 外键','索引','说明'],
            [['class\n(班级表)','id, name, grade, major,\nstatus, create_time,\nupdate_time, is_deleted','PK: id','name, grade, major','存储班级基本信息'],
             ['student\n(学生表)','id, name, no, class_id,\nweighted_score, gpa,\nstatus, create_time,\nupdate_time, is_deleted','PK: id\nFK: class_id→class.id\nUNIQUE: no','name, class_id','weighted_score 和 gpa\n由触发器自动维护'],
             ['course\n(课程表)','id, name, credit,\nstatus, create_time,\nupdate_time, is_deleted','PK: id\nUNIQUE: name','—','课程基本信息'],
             ['teacher\n(教师表)','id, name, no, title,\nphone, status,\ncreate_time, update_time,\nis_deleted','PK: id\nUNIQUE: no','name','title 和 phone 可为空'],
             ['course_offering\n(选课安排表)','id, course_id, teacher_id,\nsemester, max_students,\ncurrent_students,\nenroll_start_time,\nenroll_end_time,\ngrade_deadline, status','PK: id\nFK: course_id→course.id\nFK: teacher_id→teacher.id','course_id, teacher_id,\nsemester,\n(enroll_start_time,\nenroll_end_time)','核心业务表。current_\nstudents 由触发器维护'],
             ['enrollment\n(选课表)','id, offering_id,\nstudent_id, score,\ncreate_time, update_time,\nis_deleted','PK: id\nFK: offering_id→\n  course_offering.id\nFK: student_id→\n  student.id\nUNIQUE: (offering_id,\n  student_id)','offering_id,\nstudent_id','score=NULL 未录入；\nis_deleted=1 表示退选'],
             ['teacher_course\n(讲授表)','id, teacher_id,\ncourse_id, create_time,\nupdate_time, is_deleted','PK: id\nFK: teacher_id→\n  teacher.id\nFK: course_id→\n  course.id\nUNIQUE: (teacher_id,\n  course_id)','teacher_id,\ncourse_id','M:N 中间表'],
             ['grade_point_rule\n(绩点对照表)','id, min_score,\nmax_score, point','PK: id','—','5条初始规则']],
            [1.0,1.8,1.3,1.0,1.5], caption='表 4-1  数据库表结构汇总')
        ap(i+2,'各表关系如图 4-1 所示。')
        aimg(i+3,'table_relationships.png',5.5,'图 4-1  数据库关系图')

    i=fh('4.2 规范化处理说明')
    if i:
        an(i,[
            '数据库设计整体遵循第三范式（3NF）。同时，根据任务书 2.3 节"考虑设计必要的冗余数据列"的要求，在高频查询路径上进行了适度的反规范化设计。',
            '（1）1NF：所有字段均为不可再分的原子值。特别注意，教师与课程的讲授关系因是多对多，抽取为独立的中间表，而非在教师表中嵌入多值字段。',
            '（2）2NF：每张表有单一主键（id），所有非主属性完全函数依赖于主键。对于 enrollment 和 teacher_course 等中间表，业务主键（组合）通过 UNIQUE 约束保证。',
            '（3）3NF：消除传递依赖。例如学生表只存储 class_id（外键），班级的名称、年级、专业通过 JOIN class 表获取，不冗余存储。',
            '（4）反规范化——course_offering.current_students：该字段可通过 COUNT(enrollment) 实时计算，但选课场景下此值查询非常频繁（每次展示可选课程时都需要显示剩余名额）。通过触发器在选课/退选时自动维护该冗余字段，避免了每次查询的 COUNT 聚合，以有限的空间代价换取了显著的读取性能提升。',
            '（5）反规范化——student.weighted_score 和 gpa：学籍分需 JOIN enrollment、course_offering 和 course 三张表计算，绩点分还需额外 JOIN grade_point_rule。成绩查询是高频操作，每次查询都做四表聚合计算不可接受。通过触发器在成绩变化时自动重算并冗余存储，使成绩查询退化为单表主键查询。',
        ])

    i=fh('4.3')
    if i:
        ap(i,'逻辑设计阶段，在 DDL 语法检查和字段完整性校验方面借助了 AI。表 4-2 记录了 AI 给出的建议与人工审核后的最终设计之间的关键差异。')
        atbl(i+1,
            ['表 / 字段','AI 的原始建议','人工审核后的最终设计','修改理由'],
            [['course_offering.\neffective_time','新增此字段','删除','enroll_start_time 已表示\n生效时间，冗余字段'],
             ['enrollment.\nenroll_time','新增此字段','删除，使用\ncreate_time 代替','语义重复，选课记录的\n创建时间即选课时间'],
             ['enrollment.\nscore 的处理','单独建成绩表','保留在\nenrollment 表中','一条选课记录只有\n一个成绩，单独建表\n增加不必要的 JOIN'],
             ['student.\nweighted_score','通过视图\n实时计算','冗余列存储 +\n触发器维护','高频查询，预计算方案\n显著优于实时聚合'],
             ['teacher_course\n的实现','在 teacher 表\n加可授课程字段','独立中间表','M:N 关系，必须用\n中间表。在教师表中\n嵌入多值违反 1NF']],
            [1.4,1.6,1.6,2.0], caption='表 4-2  AI 辅助逻辑设计审核对照表')

    i=fh('4.4 其它完整性约束说明')
    if i:
        an(i,[
            '数据库通过以下多层约束机制保障数据完整性：',
            '（1）实体完整性：所有表以 id 为 PRIMARY KEY，AUTO_INCREMENT 自动生成唯一标识。',
            '（2）参照完整性：通过 FOREIGN KEY 约束建立 7 对表间引用关系，依赖 InnoDB 的事务机制和行级锁保证操作一致性。',
            '（3）用户定义完整性：UNIQUE 约束（student.no、teacher.no、course.name、enrollment(offering_id,student_id)、teacher_course(teacher_id,course_id)）；NOT NULL 约束（所有业务关键字段）；DEFAULT 值（status 默认 1、current_students 默认 0、create_time/update_time 由 MySQL 自动填充）；触发器（选课前名额检查→选课后/退选后人数自动更新→成绩变化时自动计算学籍分和绩点分）；存储过程（选课业务的 4 项校验、成绩录入的 5 项校验）。',
            '（4）逻辑删除策略：所有删除操作均为 UPDATE is_deleted=1，不进行物理删除。所有视图均过滤 is_deleted=0，确保应用层不感知已删除数据。',
        ])

def ch5():
    print('Ch5...')
    i=fh('5.1 存储结构与索引设计')
    if i:
        an(i,[
            '数据库采用 MySQL 8.0，存储引擎选择 InnoDB。理由：支持事务（ACID），保证选课/退选/成绩录入等操作的原子性；支持行级锁和外键约束，适合选课场景下的并发控制；支持崩溃恢复，保障数据安全。不考虑分库分区分表等存储方式的设计。',
            '索引设计策略：根据任务书 2.4 节的要求，分析了系统中各查询的 WHERE 条件、JOIN 条件和排序分组需求，为高频查询路径设计索引如下。',
            '（1）主键索引：所有表 id 字段的聚簇索引，数据物理有序存储。',
            '（2）唯一索引：student.no、teacher.no、course.name——既保证数据唯一性，又加速按学号/工号/课程名的精确查询。enrollment(offering_id, student_id) 和 teacher_course(teacher_id, course_id) 的复合唯一索引同时服务于这两组字段的联合查询。',
            '（3）外键索引：class_id、course_id、teacher_id、offering_id、student_id 均建辅助索引，加速 JOIN 操作。例如按班级查学生（student JOIN class），idx_student_class 使查询从全表扫描变为索引查找。',
            '（4）业务查询索引：course_offering.semester——按学期筛选排课；course_offering(enroll_start_time, enroll_end_time) 复合索引——支持选课时间段范围查询（WHERE NOW() BETWEEN enroll_start_time AND enroll_end_time），这是选课列表页面最频繁的查询条件。',
            '（5）未建索引的考量：status 字段仅 0/1 二值，选择性极低（约 50%），B+Tree 索引对其无效，优化器倾向于全表扫描。仅在与其他高选择性字段组成复合索引时才有价值。',
        ])

    i=fh('5.2 AI辅助索引设计记录')
    if i:
        ap(i,'物理设计阶段，索引方案由自己根据系统查询模式确定。AI 在此阶段扮演了任务书附录一所述的"性能分析顾问"角色——分析 SQL 性能瓶颈并提出索引建议，最终是否采纳由自己决策。')

    i=fh('5.2.1')
    if i:
        ap(i,'在向 AI 咨询索引设计问题时，遵循了任务书附录一的要求——在提示词中提供了待优化的 SQL、业务场景说明和 EXPLAIN 输出结果。提示词示例如下：')
        aquote(i,'"以下 SQL 用于按班级和课程统计成绩（平均分、最高分、最低分、及格率），该查询在教务管理页面频繁调用。当前 EXPLAIN 输出如下：[EXPLAIN 结果]。请分析当前执行计划，建议是否需要增加或调整索引来优化查询性能。请说明每个建议的理由。"')
        ap(i+1,'AI 针对此查询建议在 enrollment 表上建立 (offering_id, student_id, score) 复合索引，认为这样可以实现覆盖索引避免回表。此外，AI 还给出了两个通用建议——为所有 status 字段建立索引，以及为 student.name 建立全文索引。')

    i=fh('5.2.2')
    if i:
        ap(i,'对 AI 的建议，按照任务书要求使用 EXPLAIN 逐条验证。对于"按班级课程统计成绩"的查询，EXPLAIN 输出如图 5-1 所示，访问类型 type 为 ref 和 eq_ref，Extra 列无 Using filesort 或 Using temporary 等性能警告，rows 估算值小，索引设计合理。')
        aimg(i+1,'explain_plan.png',5.0,'图 5-1  EXPLAIN 执行计划分析（班级成绩统计查询）')

    i=fh('5.2.3')
    if i:
        an(i,[
            '经 EXPLAIN 验证和人工分析，对 AI 建议的采纳/不采纳决定如下：',
            '（1）采纳——为 enrollment 增加 (offering_id, student_id) 索引加速 JOIN：EXPLAIN 验证显示 type 从 ALL 改善为 ref，rows 大幅减少。已在 enrollment 表中实现了该索引（即 uk_enrollment 唯一约束自带的索引）。',
            '（2）未采纳——为所有 status 字段单独建索引：status 值域仅 {0,1}，选择性约 50%，B+Tree 索引几乎无过滤效果。经 EXPLAIN 对比测试，添加索引前后优化器均选择全表扫描。未采纳。',
            '（3）未采纳——为 student.name 建全文索引：系统对姓名的查询模式为精确匹配（WHERE name = \'张三\'）或前缀模糊匹配（WHERE name LIKE \'张%\'），普通 B+Tree 索引已可利用 ICP 优化。全文索引维护开销大、占用空间多，且对于精确匹配的效率不如 B+Tree。未采纳。',
        ])

def ch6():
    print('Ch6...')
    i=fh('6.1 数据库及基本表创建')
    if i:
        ap(i,'数据库实施按以下步骤完成。第一步，创建数据库 db_sms（字符集 utf8mb4，排序规则 utf8mb4_unicode_ci）。第二步，执行 DDL 脚本创建 5 张实体表和 3 张关系表。第三步，创建 3 个视图封装多表联查。第四步，创建 5 个触发器实现自动化数据维护。第五步，创建 9 个存储过程和 2 个存储函数。第六步，通过 Python 脚本批量导入测试数据（12 班级、120 学生、16 教师、12 课程、100 排课、835 选课记录）。')

    i=fh('6.1.1')
    if i:
        ap(i,'以下为核心 DDL 建表语句示例（完整脚本存放于 sql/ 目录）。')
        acode(i+1,'-- 班级表\nCREATE TABLE class (\n    id          INT             NOT NULL AUTO_INCREMENT  COMMENT \'班级ID\',\n    name        VARCHAR(50)     NOT NULL                 COMMENT \'班级名\',\n    grade       VARCHAR(10)     NOT NULL                 COMMENT \'年级\',\n    major       VARCHAR(100)    NOT NULL                 COMMENT \'专业\',\n    status      TINYINT(1)      NOT NULL DEFAULT 1       COMMENT \'1=在读 0=毕业\',\n    create_time DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP,\n    update_time DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n    is_deleted  TINYINT(1)      NOT NULL DEFAULT 0       COMMENT \'逻辑删除\',\n    PRIMARY KEY (id),\n    INDEX idx_class_name (name),\n    INDEX idx_class_grade (grade),\n    INDEX idx_class_major (major)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;')
        acode(i+2,'-- 学生表\nCREATE TABLE student (\n    id              INT             NOT NULL AUTO_INCREMENT  COMMENT \'学生ID\',\n    name            VARCHAR(50)     NOT NULL                 COMMENT \'姓名\',\n    no              VARCHAR(20)     NOT NULL                 COMMENT \'学号\',\n    class_id        INT             NOT NULL                 COMMENT \'班级ID\',\n    weighted_score  DECIMAL(5,2)    NOT NULL DEFAULT 0.00    COMMENT \'学籍分\',\n    gpa             DECIMAL(5,2)    NOT NULL DEFAULT 0.00    COMMENT \'绩点分\',\n    status          TINYINT(1)      NOT NULL DEFAULT 1       COMMENT \'1=在读 0=离校\',\n    create_time     DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP,\n    update_time     DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n    is_deleted      TINYINT(1)      NOT NULL DEFAULT 0       COMMENT \'逻辑删除\',\n    PRIMARY KEY (id),\n    UNIQUE INDEX uk_student_no (no),\n    INDEX idx_student_name (name),\n    INDEX idx_student_class (class_id),\n    CONSTRAINT fk_student_class FOREIGN KEY (class_id) REFERENCES class(id)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;')

    i=fh('6.1.2')
    if i:
        ap(i,'测试数据通过 Python 脚本使用 pymysql 的 executemany() 方法批量导入。数据来源于预先准备的 CSV 文件。绩点对照表的 5 条规则数据（来源于任务书表 2-1）直接在 SQL 脚本中以 INSERT 语句初始化。')

    i=fh('6.2 视图、存储过程、触发器、函数创建')
    if i:
        an(i,[
            '根据任务书 2.5 节"设计必要的视图或存储过程实现外模式提高数据独立性""设计必要的触发器实现用户自定义完整性以及冗余数据和冗余表的维护""设计必要的存储过程实现全部统计报表的输出"这三项要求，数据库实施了以下程序对象。',
            '（1）视图（3 个）。v_student_message：封装学生基本信息查询，对外隐藏学籍分和绩点分字段（任务书第 132 条："学生基本信息的获取要求使用视图，因为学生表上的学籍分在常规增删改查时不需要显示和查看，应该使用自己的外模式"）。v_course_plan：封装排课列表查询，通过 JOIN 和 CASE 表达式一次性输出课程名、教师名、剩余名额和选课状态。v_enrollment：封装选课详情，四表联查输出完整的选课信息。所有视图均过滤 is_deleted=0。',
            '（2）存储过程（9 个）。sp_show_courses（查询可选课程，过滤选课中、有名额、未选过、未选同课不同教师的排课）；sp_enroll（选课，含 4 层校验+事务）；sp_unenroll（退选，含时效校验+事务）；sp_grade_input（成绩录入，含身份校验+归属校验+时效校验+范围校验+事务）；sp_student_roster（按班级输出学生名单含学籍分，对应任务书第 139 条要求）；sp_class_grade_report（按班级课程统计平均/最高/最低/及格率，对应任务书第 137 条要求）；sp_student_semester_avg（某学生某学期均分，对应任务书第 138 条要求）；sp_teacher_info 和 sp_teacher_list（教师授课统计）。所有存储过程均使用 START TRANSACTION + COMMIT/ROLLBACK 保证原子性。',
            '（3）触发器（5 个）。trg_enrollment_before_insert（选课前名额检查——当 current_students >= max_students 时 SIGNAL 报错，对应任务书第 133 条"到达限定人数再选抛出异常"）；trg_enrollment_after_insert（选课后 current_students+1——对应任务书第 133 条"利用触发器、冗余列实时动态更新选课人数"）；trg_enrollment_after_update（退选后 current_students-1）；trg_enrollment_after_insert_score 和 trg_enrollment_after_update_score（成绩变化时自动重算 weighted_score 和 gpa——对应任务书 7.2.2 节运维需求变更中"新增冗余列保存实时绩点分，新增或修改触发器代码"的要求）。',
            '（4）存储函数（2 个）。fn_get_student_id（学号→学生 ID）和 fn_get_teacher_id（工号→教师 ID），供各存储过程复用，避免重复编写相同的 ID 查询逻辑。',
        ])

    # 6.3 AI代码生成
    i=fh('6.3 AI辅助代码生成与人工审核记录')
    if i:
        ap(i,'数据库实施阶段是 SQL 代码量最大的环节。在此阶段，严格遵循了任务书附录一要求的 AI 使用前置决策原则——在向 AI 提问前，先自行判断该功能应使用哪种数据库对象，然后在提示词中明确指定。以下以选课存储过程为例，记录完整的 AI 辅助流程。')

    i=fh('6.3.1')
    if i:
        ap(i,'编写选课功能时，我先分析了该功能的需求特征：①需要校验多项业务规则（学生存在性→选课有效期→是否重复→是否同课不同教师）；②涉及多表操作（查询 student、查询 course_offering、插入 enrollment）；③需要事务保证原子性——任何一步校验失败都应回滚；④由用户主动调用。对照任务书附录一的"数据库对象选型参考"（Table 2），需要主动调用、逻辑复杂、需事务控制的场景应使用存储过程。因此，我先确定了方案——使用存储过程，然后在提示词中明确告知 AI 要创建的是存储过程。提示词如下：')
        aquote(i,'"请创建一个 MySQL 存储过程 sp_enroll，用于实现学生选课功能。该存储过程接收两个输入参数：p_student_no（学号，VARCHAR(20)）和 p_plan_id（排课 ID，INT）。需要依次完成以下校验：①通过 fn_get_student_id 函数查找学生 ID，若不存在则报错\'学生不存在\'；②检查该排课是否在选课期内（NOW() BETWEEN enroll_start_time AND enroll_end_time）；③检查学生是否已选过同一排课（enrollment 表中已存在且 is_deleted=0）；④检查学生是否已选过同一门课的其他教师。全部校验通过后向 enrollment 表插入数据。使用 START TRANSACTION 和 COMMIT/ROLLBACK 实现事务控制。校验失败时使用 SIGNAL SQLSTATE \'45000\' 抛出具体的中文错误信息。"')

    i=fh('6.3.2')
    if i:
        ap(i,'AI 返回的代码包含了存储过程的标准结构：参数声明、变量声明（DECLARE）、事务框架（START TRANSACTION + COMMIT/ROLLBACK）、四条校验的 SQL 骨架以及 SIGNAL 错误抛出语句。代码结构清晰，基本覆盖了我在提示词中列出的所有校验点。但代码中存在几处需要修正的问题，详见下一节。')

    i=fh('6.3.3')
    if i:
        an(i,[
            '在人工审核和调试过程中，发现并修复了以下问题：',
            '（1）数据库名错误：AI 生成的 CREATE DATABASE 语句中数据库名写成了 db_purchase_sales_inventory（上一项目的数据库名）。这说明 AI 在某种程度上受到了上下文的影响。人工修正为 db_sms。',
            '（2）字段冗余：AI 在 course_offering 表中添加了一个 effective_time 字段，但该表已有 enroll_start_time 表示生效时间，两者功能重复。人工删除了此字段。enrollment 表中 AI 同时设计了 enroll_time 和 create_time，语义重复。人工分析后保留了 create_time（选课记录的创建时间就是选课时间），删除了 enroll_time。',
            '（3）COLLATE 冲突：存储过程中字符串比较（WHERE xxx = \'某值\'）未指定 COLLATE，在特定配置下引发\"Illegal mix of collations\"错误。人工在所有字符串比较处添加了 COLLATE utf8mb4_unicode_ci 子句。',
            '（4）NULL 值比较错误：触发器中判断成绩是否发生变化时，AI 使用了普通等号 =。但 enrollment.score 允许 NULL（未录入成绩时），而 SQL 中 NULL = NULL 返回 NULL 而非 TRUE。这意味着"成绩从 85 改为 85"和"成绩从未录入改为未录入"这两种情况都无法被正确判断。人工将比较运算符改为 MySQL 的 NULL 安全等于运算符 <=>。',
            '（5）sed 批量替换副作用：开发过程中使用 sed 命令进行代码批量替换时，部分存储过程名中包含的替换关键字被意外修改，导致存储过程无法调用。人工逐条检查并手动恢复了被误改的名称。',
            '总体而言，AI 在此阶段生成的初始代码约有 30%~40% 在后续人工调试中被修改或重写。这印证了任务书的要求——AI 的角色是辅助者而非替代者，所有 AI 生成的内容必须经过人工审核和测试。',
        ])

def ch7():
    print('Ch7...')
    i=fh('第七章 系统实现')
    if i:
        ap(i,'在完成数据库设计的基础上，使用 Python 语言开发了命令行（CLI）和 Streamlit Web 两种用户界面，实现了学生、教师、教务管理员三种角色的全部功能。')

        ah2(i+1,'一、系统架构')
        ap(i+2,'系统采用分层架构：表示层（CLI 终端 + Streamlit Web 页面）、业务逻辑层（选课/退选/成绩录入/统计查询/权限控制）、数据访问层（PyMySQL 驱动 + 存储过程调用 + 游标管理）和数据存储层（MySQL 8.0 + 8 表 + 3 视图 + 5 触发器 + 9 存储过程 + 2 函数）。各层通过清晰接口通信——表示层调用业务逻辑层函数，业务逻辑层通过数据访问层执行 SQL 和存储过程，数据存储层通过视图、触发器和存储过程向上层屏蔽内部实现细节，体现了数据独立性的设计原则。')
        aimg(i+3,'system_architecture.png',5.0,'图 7-1  系统架构图')

        ah2(i+4,'二、用户界面')
        an(i+5,[
            '（1）命令行界面（CLI）：通过菜单驱动交互，使用 pymysql 调用存储过程。实现了 CJK 字符宽度感知的表格对齐算法（解决中英文混排时的列对齐问题）、Paginator 分页组件（支持翻页和跳转）、清屏（兼容 Windows/Linux）、暂停和确认提示等交互优化。支持学生、教师、教务和测试员（可切换任意角色进行功能验证）四种角色。',
            '（2）Web 界面（Streamlit）：基于 Streamlit 框架构建响应式单页应用，使用 st.navigation 实现原生多页导航，页面按角色和功能分组折叠。学生端 5 个页面（可选课程/选课/退选/成绩查询/学期均分）；教师端 3 个页面（成绩录入/CSV 批量导入/课程学生查看）；教务端分为三大模块——数据查看（数据概览仪表板 + 班级名单 + 成绩统计 + 成绩明细 + 教师信息 + 教师列表共 6 页）、数据管理（班级/课程/教师/学生/排课/选课共 6 个 CRUD 页面）、系统工具（备份 + 恢复共 2 页）。',
        ])
        j=ft('（2）Web 界面（Streamlit）')
        if j:
            aimg(j,'ui_login.png',4.2,'图 7-2  系统登录页面')
            aimg(j+1,'ui_admin_dashboard.png',5.0,'图 7-3  教务端数据概览页面')
            aimg(j+2,'ui_student_enroll.png',5.0,'图 7-4  学生选课页面')

        ah2(i+8,'三、关键实现技术')
        an(i+9,[
            '（1）权限控制：基于角色的访问控制（RBAC）。登录时根据输入的学号/工号/admin 在 student 表和 teacher 表中匹配，确定角色和身份后存入 st.session_state。各页面根据角色动态渲染对应的导航和功能组件。',
            '（2）消息跨 rerun 持久化：Streamlit 的 st.toast/success/error 在页面 rerun 时状态丢失。通过在 st.session_state 中维护 msg 元组实现跨 rerun 的消息传递。',
            '（3）数据备份恢复：备份通过 Python 调用 mysqldump 实现，恢复通过 st.file_uploader 上传文件 + session_state 持久化 + mysql 命令行执行。恢复前自动 SET FOREIGN_KEY_CHECKS=0 避免导入顺序问题。',
            '（4）CSV 批量导入成绩：解析 CSV（兼容 UTF-8 BOM），逐行调用 sp_grade_input 存储过程，汇总统计成功/失败数量并提供逐条错误反馈。',
        ])

def ch8():
    print('Ch8...')
    i=fh('8.1 测试方案与测试用例')
    if i:
        ap(i,'系统测试采用黑盒测试方法，覆盖三种角色的正常功能流程和异常/边界场景。测试流程如图 8-1 所示。')
        aimg(i+1,'test_flow.png',5.0,'图 8-1  系统测试流程图')
        ap(i+2,'各模块主要测试用例及结果如表 8-1 所示。')
        atbl(i+3,
            ['测试模块','测试用例','预期结果','实际'],
            [['登录认证','学生输入正确学号','进入学生端','✓ 通过'],
             ['登录认证','教师输入正确工号','进入教师端','✓ 通过'],
             ['登录认证','输入 admin','进入教务端','✓ 通过'],
             ['登录认证','输入不存在的用户','提示"用户不存在"','✓ 通过'],
             ['学生选课','正常选课（有名额、选课期内）','选课成功，名额-1','✓ 通过'],
             ['学生选课','选已满课程','提示"选课已满"','✓ 通过'],
             ['学生选课','重复选同一排课','提示"已选过这门课"','✓ 通过'],
             ['学生选课','选同门课不同教师','提示相应的错误信息','✓ 通过'],
             ['学生退选','正常退选（选课截止前）','退选成功，名额+1','✓ 通过'],
             ['学生退选','选课截止后退选','提示"已过退选截止时间"','✓ 通过'],
             ['成绩录入','正常录入成绩','录入成功，学籍分/GPA更新','✓ 通过'],
             ['成绩录入','录入非本人课程成绩','提示"这不是你的课"','✓ 通过'],
             ['成绩录入','录入超出0~100范围的成绩','提示"不在0~100之间"','✓ 通过'],
             ['成绩录入','超截止时间录入','提示"已过成绩录入截止时间"','✓ 通过'],
             ['教务管理','各实体CRUD操作','增删改查正确执行','✓ 通过'],
             ['成绩统计','按班级课程统计成绩','数据与原数据一致','✓ 通过'],
             ['备份恢复','备份→删数据→恢复','恢复后数据完整一致','✓ 通过']],
            [1.0,2.2,2.2,0.7], caption='表 8-1  系统测试用例及结果')

    j=ft('8.2 测试结果与分析')
    if j:
        an(j,[
            '全部 17 个测试用例均通过验证，通过率 100%，系统功能完整、运行稳定。',
            '测试过程中发现并修复的主要 Bug：',
            '（1）成绩录入未校验教师授课权限——在 grade_input 中增加了排课归属校验。',
            '（2）成绩更新后学籍分/GPA 未自动重算——增加两个 AFTER 触发器实现自动更新（对应任务书 7.2.2 节运维需求变更要求）。',
            '（3）Streamlit 的 file_uploader 在 rerun 后丢失文件——改用 session_state 持久化。',
            '（4）数据恢复时中文乱码——Windows 环境下 GBK/UTF-8 不匹配，指定 encoding=\'utf-8\' 解决。',
            '测试结论：系统满足任务书 7.2.1 节的全部基本需求，可以投入使用。',
        ])

    i=fh('8.3 AI辅助测试')
    if i:
        an(i,[
            '测试阶段 AI 仅参与了以下三项辅助工作：',
            '（1）帮助分析边界值场景——如选课名额刚好满员时的并发行为、成绩为 0 和 100 分时的处理逻辑。',
            '（2）调试时帮助分析报错信息——如存储过程的 SIGNAL 错误信息解读、collation 冲突的原因分析。',
            '（3）生成了部分 CSV 格式的测试数据模板——按照我指定的格式（plan_id, student_no, score）和数值范围生成，用于批量导入功能测试。',
            '测试用例的设计（覆盖哪些场景、如何验证结果）、测试的执行和 Bug 的修复均由自己独立完成。',
        ])

def ch9():
    print('Ch9...')
    i=fh('第九章 总结')
    if i:
        an(i,[
            '（1）课程设计完成的主要工作和成果：',
            '本课程设计按照数据库设计的标准流程，完成了学生成绩管理系统的从零到完整实现。数据库层面：完成了需求分析、概念设计（全局 E-R 图）、逻辑设计（8 张表，满足 3NF，含必要的反规范化设计）、物理设计（InnoDB + 索引策略）、数据库实施（DDL + 3 视图 + 5 触发器 + 9 存储过程 + 2 函数 + 批量测试数据）。应用层面：使用 Python 开发了 CLI 和 Streamlit Web 双模式界面，实现了任务书 7.2.1 节要求的全部基本功能。系统通过了 17 个测试用例的验证。',
            '（2）收获和体会：',
            '通过本次课程设计，系统性地实践了数据库设计的六个阶段。概念设计阶段体会到了"正确识别实体和关系"对整个后续设计的决定性影响——如果 ER 图画错了，后面的表结构一定会出问题。逻辑设计阶段最大的收获是学会了在范式和性能之间做权衡——不是越符合高范式越好，而是要根据实际的查询模式来决定哪些地方需要保留冗余。物理设计阶段通过亲手使用 EXPLAIN 分析执行计划，直观地感受到了索引对查询性能的影响，也学会了不盲目建索引。数据库实施阶段是工程量最大的环节，在这个过程中对 MySQL 的存储过程、触发器、视图和函数的语法和调试有了远超课堂的深入理解。另外，Python 应用开发的过程让我深刻体会到"数据库设计得好，应用代码就简单"这个道理。',
            '（3）存在的不足和改进设想：',
            '(a) 当前登录仅通过学号/工号识别身份，未实现密码认证。后续可增加密码哈希存储和登录验证。(b) 选课并发控制依赖 InnoDB 行级锁，在选课高峰（如开学选课日）可能出现锁竞争。可考虑引入 Redis 缓存选课名额实现无锁扣减。(c) Streamlit 界面为单体应用，后续可考虑前后端分离，将数据访问层独立为 RESTful API。(d) 成绩统计维度较为有限，可扩展年级排名、课程横向对比等更多分析维度。',
            '（4）课程设计的感想：',
            '这次课程设计让我从一个不同的角度理解了数据库。以前学数据库，更多是在课本上看到一个个独立的概念——范式、索引、事务、触发器——但不知道它们在实际项目中是怎么配合使用的。这次自己从头设计一个系统，才真正体会到这些概念之间的关联。比如，为了满足 3NF，应该把选课人数通过 COUNT 实时计算，但实际选课页面每秒钟要查询几十次"还剩多少个名额"，每次都 COUNT 显然不合理，于是故意在 course_offering 表里冗余了一个 current_students 字段，再用触发器来保证它和 enrollment 表的数据一致性。这个过程涉及了范式、索引、触发器和反规范化设计四个知识点，是课堂上很难体会到的。另外，这次课程设计也让我认识到"细节决定成败"——一个 NULL 值比较方式、一个 COLLATE 的设置，都可能成为调试数小时的根源。这些经验只有亲手做过才能真正理解。',
        ])

    # 9.1 AI
    i=fh('9.1 AI辅助效果评估与反思')
    if i:
        ap(i,'本次课程设计中，按照任务书附录一的要求选择了路径 B，将 AI 定位为辅助编码和查阅资料的工具。以下从效率、质量和反思三个角度进行总结。')

    i=fh('9.1.1')
    if i:
        an(i,[
            'AI 节省时间的环节：（1）语法层面的检查——SQL 的括号配对、分号位置、关键字拼写，AI 检查得比较快。（2）重复性代码框架——Streamlit 的 st.dataframe()、st.metric() 这类组件调用格式固定，AI 生成框架比较高效。DDL 语句的基本结构也是如此。（3）代码冗余检查——在代码量较大的 admin.py（约 1176 行）中，AI 帮助发现了重复的字典构建和未使用的 import 等问题。',
            'AI 反而更耗时的环节：（1）隐藏的逻辑错误——AI 在触发器中用普通等号 = 比较 NULL 值，这个错误在正常测试数据下不会暴露（因为测试数据中成绩都有值），直到我刻意构造了成绩为 NULL 的场景才发现。排查这类"看起来正确但实际有 Bug"的问题比从头手写更耗时。（2）AI 对项目上下文理解有限——生成的代码风格不一致，或者使用了和已有部分不兼容的写法，需要统一调整。（3）部分 API 过时——AI 给出的 Streamlit 代码偶尔使用旧版 API，编译通过但功能不符合预期。',
        ])

    i=fh('9.1.2')
    if i:
        an(i,[
            '使用过程中遇到的典型 AI 错误和经验：',
            '（1）数据库名被上下文"污染"——AI 在建库语句中使用了上一项目的数据库名。这说明 AI 并不理解当前项目的实际语境，它只是在统计层面上生成"看起来合理"的文本。',
            '（2）NULL 值比较是一类隐蔽且危害大的错误。AI 多次在 SQL 中使用 = 来比较可能为 NULL 的值。这个错误的隐蔽之处在于：如果测试数据恰好不包含 NULL，它可以通过测试，但一到生产环境就会出现逻辑错误。这让我意识到，AI 生成的代码需要经过"刻意构造边界条件"的测试，而不能只跑正常流程。',
            '（3）AI 有时会以很确定的语气给出不正确的建议。比如它很自信地建议"为所有 status 字段建索引"，但经过 EXPLAIN 实测发现优化器根本不使用这些索引。这提醒我，对 AI 的输出必须保持怀疑态度，用实际测试来验证。',
            '总结几点经验：（1）AI 适合做"辅助"——语法检查、框架生成、格式整理，这些机械重复的工作可以交给它。（2）核心决策必须自己做——范式策略、索引设计、数据库对象选型，这些需要综合判断的事情，AI 无法替代人的思考。（3）AI 的每一行输出都要审、要测——不能因为"看起来对"就相信它是对的。',
        ])

    i=fh('9.1.3')
    if i:
        an(i,[
            '回顾整个课程设计，我认为以下环节是必须由人来完成、不能交给 AI 的：',
            '（1）需求分析和概念设计——识别业务实体、判断关系基数（1:1、1:N、M:N），需要对业务场景有真实的理解。AI 缺乏对教务管理流程的实际感知，容易做出不符合实际情况的判断（如将选课安排与学生关系误判为 N:1）。',
            '（2）范式和反范式的权衡——在哪些地方严格遵守 3NF、在哪些地方做反规范化设计来换取性能，需要对系统的查询模式和数据量有整体把握。AI 只能提供一般原则，无法针对具体系统做出恰当的设计决策。',
            '（3）索引设计——哪些字段建索引、建什么类型的索引、是否需要复合索引，需要分析具体的 SQL 和查询频率。AI 的建议需要人用 EXPLAIN 逐条验证后才能决定是否采纳。',
            '（4）测试用例设计——尤其是异常和边界场景的覆盖，需要设计者对业务规则有全面理解。AI 生成的测试用例偏向"快乐路径"，对异常和边界的覆盖不够。',
            '总的来说，本次课程设计的核心工作——数据库设计的主要决策和关键代码的实现——都是由自己完成的。AI 在我手中是一个"更快的搜索引擎"加"代码框架生成器"。通过这次实践，我最大的收获之一是学会了如何合理地使用 AI——知道它擅长什么、不擅长什么，什么时候用它提效，什么时候关掉它自己沉下心来思考和编码。这也是任务书希望我们通过这次课程设计掌握的——在智能化开发环境中，保持独立思考和判断的能力。',
        ])

def refs_app():
    print('Refs...')
    i=fh('参考文献')
    if i:
        an(i,[
            '[1] 王珊, 萨师煊. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.',
            '[2] 吴臣. 数据库系统课程设计任务书[Z]. 2026.',
            '[3] MySQL 8.0 Reference Manual [EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/.',
            '[4] Streamlit Documentation [EB/OL]. https://docs.streamlit.io/.',
            '[5] Python 3.12 Documentation [EB/OL]. https://docs.python.org/3.12/.',
        ])
    i=fh('附录')
    if i:
        an(i,[
            '本课程设计的完整代码和 SQL 脚本存放于项目目录中，主要文件结构如下：',
            'sql/ 目录：01_数据库创建.sql、02_基础数据表.sql、03_中间表.sql、04_视图.sql、06_触发器.sql、07_存储过程.sql、08_存储函数.sql，共 7 个脚本文件。',
            'src/ 目录：app.py（Streamlit Web 主应用）、main.py（CLI 入口）、admin.py（教务管理）、student.py（学生端）、teacher.py（教师端）、tester.py（测试工具）、core/（配置/认证/工具模块）。',
            'data/ 目录：class.csv、student.csv、teacher.csv、course.csv、course_offering.csv、enrollment.csv、teacher_course.csv，共 7 个测试数据文件。',
            'docs/ 目录：开发日志.md、AI修正日志.md、日志.md 等开发过程记录文档。',
            '',
            '开发过程中 AI 辅助修正的部分典型记录（摘录自 AI修正日志.md）：',
            '• 07-07：AI 在 course_offering 表多加了 effective_time 字段（与 enroll_start_time 重复）→ 人工删除。',
            '• 07-07：AI 在 enrollment 表同时设计了 enroll_time 和 create_time（语义重复）→ 人工删除 enroll_time。',
            '• 07-07：AI 生成的数据库名错为上一项目名称 → 人工改为 db_sms。',
            '• 07-08：sed 批量替换误伤存储过程名 → 人工逐条手动修复。',
            '• 07-09：Streamlit 代码中复制粘贴导致 session key 混用 → 人工逐一定位修正。',
            '• 07-09：数据恢复未指定编码致 Windows 下中文乱码 → 人工增加 encoding=\'utf-8\'。',
        ])

# ====== 执行 ======
if __name__=='__main__':
    print('='*50)
    ch1(); ch2(); ch3(); ch4(); ch5(); ch6(); ch7(); ch8(); ch9(); refs_app()
    doc.save(OUT)
    import os
    sz=os.path.getsize(OUT)/1024
    im=len([r for r in doc.part.rels.values() if 'image' in r.reltype])
    print(f'\nDone: {len(doc.paragraphs)}p | {len(doc.tables)}tbl | {im}img | {sz:.0f}KB')
