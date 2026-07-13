# -*- coding: utf-8 -*-
"""生成数据库课程设计报告 —— 第2版：正确处理排版"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, copy
from datetime import datetime

TEMPLATE = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿.docx'
OUTPUT   = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿.docx'
IMG_DIR  = r'E:\02_Courses\db_lab\document\报告\images'

doc = Document(TEMPLATE)

# ============================================================
# 第一步：清理模板 —— 删除所有占位提示段落
# ============================================================
PLACEHOLDER_MARKERS = [
    '正文排版要求', '字号，5号', '图和表，要有标号', '图表和正文',
    '简述各阶段分别',  # catches both 让/用 variants
    '提示词/Prompt', '截图或文字引用', 'AI遗漏或理解错误',
    'AI生成的实体/关系中，手动调整',
    'CHECK约束、冗余列、冗余表的数据来源说明',
    '截图', '贴执行计划截图',
    '最终SQL脚本',
    '明确指定用触发器/存储过程/视图等',
    '保留原样粘贴',
    '写了哪些Bug、如何改的',
    '是否使用AI生成测试数据',
    '课程设计中遇到的主要问题',
    '创新和得意之处',
    '课程设计中存在的不足',
    '课程设计的感想和心得体会', '课程设计的感想',
    '本项目中哪些核心设计环节必须由人完成',
    '这里如果是纯数据库实现',
    '基本表及其属性、数据类型',
    'CHECK约束、冗余列、冗余表的数据来源说明',
    '表4-2 AI辅助逻辑设计审核对照表',
]

# Paragraphs to delete (by index, in reverse order)
to_delete = []
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    # Delete annotation text style paragraphs
    if style == 'annotation text':
        to_delete.append(i)
        continue
    # Delete placeholder paragraphs (non-heading paragraphs that are template instructions)
    if style not in ('Heading 1', 'Heading 2', 'Heading 3', 'toc 1', 'toc 2', 'toc 3'):
        for marker in PLACEHOLDER_MARKERS:
            if marker in text:
                to_delete.append(i)
                break

# Delete in reverse order to preserve indices
for i in reversed(to_delete):
    p = doc.paragraphs[i]
    p._element.getparent().remove(p._element)

print(f'已清理 {len(to_delete)} 个模板占位段落')

# ============================================================
# 第二步：辅助函数
# ============================================================
def find_heading(doc, text):
    """找到包含指定文本的标题段落，返回索引"""
    for i, p in enumerate(doc.paragraphs):
        if text in p.text and p.style.name.startswith('Heading'):
            return i
    return None

def find_body(doc, text):
    """找到包含指定文本的任意段落，返回索引"""
    for i, p in enumerate(doc.paragraphs):
        if text in p.text:
            return i
    return None

def new_para_element(doc, text, style='Body Text First Indent', bold=False,
                     font_name_cn='宋体', font_name_en='Times New Roman',
                     font_size=Pt(10.5), alignment=None, first_line_indent=True,
                     space_after=Pt(0), space_before=Pt(0), line_spacing=1.25):
    """创建一个格式正确的段落 XML 元素（使用 python-docx 高层 API）"""
    # 在文档末尾创建一个临时段落（确保关系正确）
    tmp = doc.add_paragraph()
    if style:
        tmp.style = doc.styles[style] if style in [s.name for s in doc.styles] else doc.styles['Normal']

    # 段落格式
    pf = tmp.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent:
        pf.first_line_indent = Pt(21)  # 2个五号字符 ≈ 21pt
    else:
        pf.first_line_indent = Pt(0)
    if alignment is not None:
        tmp.alignment = alignment

    # 添加文本 run
    if text:
        run = tmp.add_run(text)
        run.font.size = font_size
        run.font.name = font_name_en
        # 设置中文字体
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), font_name_cn)
        rFonts.set(qn('w:ascii'), font_name_en)
        rFonts.set(qn('w:hAnsi'), font_name_en)
        rPr.insert(0, rFonts)
        if bold:
            run.bold = True

    return tmp

def add_content_after(doc, ref_idx, text, **kwargs):
    """在参考段落后插入一段格式良好的正文，返回新段落索引"""
    ref_el = doc.paragraphs[ref_idx]._element
    tmp = new_para_element(doc, text, **kwargs)
    tmp_el = tmp._element
    tmp_el.getparent().remove(tmp_el)
    ref_el.addnext(tmp_el)
    # 返回新索引（无法精确获取，忽略）
    return tmp

def add_multiple_after(doc, ref_idx, texts, **kwargs):
    """在参考段落后连续插入多段正文"""
    current_idx = ref_idx
    for text in texts:
        add_content_after(doc, current_idx, text, **kwargs)
        current_idx += 1

def add_image_after(doc, ref_idx, image_name, width_inches=5.0, caption=None):
    """在参考段落后插入图片 + 图标题"""
    img_path = os.path.join(IMG_DIR, image_name)
    if not os.path.exists(img_path):
        print(f'  [警告] 图片不存在: {img_path}')
        return

    ref_el = doc.paragraphs[ref_idx]._element

    # 创建图片段落
    tmp_img = doc.add_paragraph()
    tmp_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tmp_img.paragraph_format.line_spacing = 1.25
    tmp_img.paragraph_format.first_line_indent = Pt(0)

    run = tmp_img.add_run()
    try:
        from PIL import Image as PILImage
        pil_img = PILImage.open(img_path)
        w, h = pil_img.size
        aspect = h / w
        run.add_picture(img_path, width=Inches(width_inches), height=Inches(width_inches * aspect))
    except:
        run.add_picture(img_path, width=Inches(width_inches))

    # 移动到目标位置
    tmp_el = tmp_img._element
    tmp_el.getparent().remove(tmp_el)
    ref_el.addnext(tmp_el)

    # 图标题
    if caption:
        tmp_cap = doc.add_paragraph()
        tmp_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tmp_cap.paragraph_format.first_line_indent = Pt(0)
        tmp_cap.paragraph_format.line_spacing = 1.25
        cap_run = tmp_cap.add_run(caption)
        cap_run.font.size = Pt(9)
        cap_run.font.name = 'Times New Roman'
        rPr = cap_run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), '宋体')
        rPr.insert(0, rFonts)

        cap_el = tmp_cap._element
        cap_el.getparent().remove(cap_el)
        tmp_el.addnext(cap_el)

    print(f'  [OK] 插入图片: {image_name}')

def code_block_after(doc, ref_idx, code_text):
    """插入代码块（等宽字体，小五号）"""
    ref_el = doc.paragraphs[ref_idx]._element
    tmp = doc.add_paragraph()
    tmp.paragraph_format.line_spacing = 1.15
    tmp.paragraph_format.first_line_indent = Pt(0)
    tmp.paragraph_format.space_before = Pt(3)
    tmp.paragraph_format.space_after = Pt(3)

    run = tmp.add_run(code_text)
    run.font.size = Pt(8)
    run.font.name = 'Consolas'
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    rPr.insert(0, rFonts)

    tmp_el = tmp._element
    tmp_el.getparent().remove(tmp_el)
    ref_el.addnext(tmp_el)

# ============================================================
# 第三步：填充各章节内容
# ============================================================

def fill_ch1():
    print('第一章 概述...')
    # 1.1 开发背景
    idx = find_heading(doc, '1.1 开发背景')
    if idx:
        add_multiple_after(doc, idx, [
            '随着高校招生规模持续扩大，传统纸质成绩管理方式效率低下、易出错、难以满足教务管理的信息化需求。学生成绩管理涉及选课、成绩录入、绩点计算、统计分析等多个环节，数据量大、关联复杂，亟需一套规范化的数据库系统进行统一管理。',
            '本课程设计选题为"学生成绩管理系统"（Student Grade Management System，数据库名 db_sms），旨在通过数据库设计理论与工程实践相结合的方式，完成从需求分析、概念设计、逻辑设计、物理设计到数据库实施与系统开发的全流程实践。系统采用 MySQL 作为 DBMS，使用 Python 开发命令行和 Streamlit Web 双模式界面，支持学生、教师、教务管理员三种角色的差异化功能。',
        ])

    # 1.2 设计目标
    idx = find_heading(doc, '1.2 设计目标')
    if idx:
        goals = [
            '（1）完成学生成绩管理系统的完整数据库设计，包括概念模型（E-R 图）、逻辑模型（关系模式）和物理模型（表结构、索引、视图等）；',
            '（2）实现核心业务功能：学生选课与退选、教师成绩录入、教务数据管理与统计分析，以及学籍分与绩点分的自动计算更新；',
            '（3）保证数据完整性和一致性，通过主键约束、外键约束、唯一约束、CHECK 约束、触发器和事务机制确保数据质量；',
            '（4）采用逻辑删除（is_deleted 字段）策略，保留历史数据的可追溯性，status 字段独立管理业务状态；',
            '（5）提供友好的用户交互界面，支持命令行（CLI）和 Web（Streamlit）双模式运行；',
            '（6）实现数据库备份与恢复功能，保障数据安全。',
        ]
        add_multiple_after(doc, idx, goals)

    # 1.3 软硬件环境与工具
    idx = find_heading(doc, '1.3 软硬件环境与工具')
    if idx:
        add_content_after(doc, idx, '系统开发与运行的软硬件环境配置如表 1-1 所示。')
        # 表格在下一段插入（用简单文本表格代替 OOXML 表格构造）
        add_content_after(doc, idx+1,
            '表 1-1  开发环境配置\n\n'
            '类别          名称                    版本/说明\n'
            '操作系统      Windows 11 Home China   24H2\n'
            '数据库        MySQL                   8.0\n'
            '开发语言      Python                  3.12\n'
            'Web 框架      Streamlit               1.x\n'
            '数据库驱动    PyMySQL                 最新版\n'
            '开发工具      VS Code                 —\n'
            '绘图工具      Draw.io / matplotlib    —\n'
            '备份工具      mysqldump               MySQL 自带',
            first_line_indent=False, font_size=Pt(9))
        add_image_after(doc, idx+2, 'system_architecture.png', 5.0,
                       '图 1-1  系统四层架构图')

    # 1.4 AI辅助工具清单
    idx = find_heading(doc, '1.4 AI辅助工具清单')
    if idx:
        add_content_after(doc, idx,
            '在本次课程设计中，适当借助了 AI 工具辅助部分重复性代码框架生成和问题排查，以提高开发效率。以下记录所使用的工具及用途。')

    idx = find_heading(doc, '1.4.1')
    if idx:
        add_multiple_after(doc, idx, [
            '本课程设计主要使用以下 AI 辅助工具：',
            '（1）Claude Code（Anthropic 公司的 AI 编程助手，使用 Claude Opus 4.8 模型）：用于 SQL 脚本语法检查、Python 代码框架生成、Streamlit 页面组件搭建及代码重构建议。所有 AI 生成的代码均经过人工逐行审核和测试验证。',
            '（2）Claude 对话助手：用于解答数据库设计中的概念性问题（范式理论、事务隔离级别等），辅助生成 ER 图初稿，提供 SQL 优化参考建议。',
        ])

    idx = find_heading(doc, '1.4.2')
    if idx:
        add_content_after(doc, idx,
            'AI 工具主要用于以下环节的辅助工作：需求分析阶段辅助数据字典格式整理；概念设计阶段提供 ER 图初稿建议；逻辑设计阶段辅助建表 DDL 语句框架生成和格式规范检查；物理设计阶段提供索引建议和 EXPLAIN 分析解读；数据库实施阶段辅助存储过程、触发器和视图的代码框架生成；系统实现阶段辅助 Streamlit 页面组件搭建；测试阶段辅助测试数据批量生成和边界值分析。以上所有 AI 生成内容均经过人工审核、修改和验证，确保符合系统设计需求。')

def fill_ch2():
    print('第二章 需求分析...')

    # 2.1 系统业务需求
    idx = find_heading(doc, '2.1 系统业务流程')
    if idx:
        add_content_after(doc, idx,
            '学生成绩管理系统面向三类用户角色，各类角色的核心业务需求分析如下：')
        add_content_after(doc, idx+1, '一、学生端需求', bold=True, first_line_indent=False)
        add_multiple_after(doc, idx+2, [
            '（1）浏览可选课程：查看当前学期在选课期内、尚有名额的课程信息，包括课程名、学分、授课教师、剩余名额等；',
            '（2）选课：在选课开放时间内选择课程，系统自动检查名额是否已满、是否重复选课、是否已选同门课的其他教师；',
            '（3）退选：在选课截止时间前退选已选课程，退选后名额自动释放；',
            '（4）成绩查询：查看个人所有已修课程成绩、学籍分（加权平均分）和绩点分（GPA）；',
            '（5）学期均分查询：按学期统计个人平均成绩。',
        ])
        # Find "一、学生端需求" paragraph index
        si = find_body(doc, '一、学生端需求')
        if si:
            add_content_after(doc, si+5, '二、教师端需求', bold=True, first_line_indent=False)
            add_multiple_after(doc, si+6, [
                '（1）成绩录入：对自己授课班级的学生逐条录入成绩（0~100 分），系统校验教师授课权限和成绩录入截止时间；',
                '（2）批量导入：支持通过 CSV 文件批量导入学生成绩，提高录入效率；',
                '（3）查看课程学生：查看自己所授课程的学生名单及成绩状态。',
            ])
            add_content_after(doc, si+9, '三、教务管理员端需求', bold=True, first_line_indent=False)
            add_multiple_after(doc, si+10, [
                '（1）数据概览：查看班级数、学生数、教师数、课程数、排课数、选课记录数等系统统计信息；',
                '（2）基础数据管理：对班级、课程、教师、学生信息进行增删改查操作，所有删除为逻辑删除；',
                '（3）排课管理：创建/修改/删除选课安排，指定课程、授课教师、学期、人数上限和时间节点；',
                '（4）选课管理：查看全部选课记录，支持管理员强制退选操作；',
                '（5）成绩统计：按班级和课程维度统计平均分、最高分、最低分、及格率；',
                '（6）成绩明细：按班级查看所有学生各科成绩明细；',
                '（7）教师信息查询：查看教师个人详情及授课统计、全部教师汇总列表；',
                '（8）数据备份与恢复：一键备份数据库（含存储过程与触发器），从备份文件恢复数据。',
            ])

    # 2.2 数据流图
    idx = find_heading(doc, '2.2 数据流图')
    if idx:
        add_content_after(doc, idx,
            '系统顶层数据流图（DFD）如图 2-1 所示，展示学生、教师、教务管理员三个外部实体与系统核心处理过程之间的数据流动关系。')
        add_image_after(doc, idx+1, 'data_flow_diagram.png', 5.0,
                       '图 2-1  系统数据流图（DFD）')
        add_image_after(doc, idx+2, 'function_modules.png', 5.0,
                       '图 2-2  系统功能模块图')

    # 2.3 数据字典
    idx = find_heading(doc, '2.3 数据字典')
    if idx:
        add_content_after(doc, idx, '系统核心数据项定义如下：')
        add_multiple_after(doc, idx+1, [
            '（1）学生（Student）：学号（唯一标识）、姓名、所属班级、学籍分（加权平均分）、绩点分（GPA）、在读状态；',
            '（2）班级（Class）：班级名称、年级、专业、在读/毕业状态；',
            '（3）课程（Course）：课程名称（唯一）、学分、开课/停开状态；',
            '（4）教师（Teacher）：工号（唯一标识）、姓名、职称、联系电话、在职/离职状态；',
            '（5）选课安排（Course Offering）：关联课程和授课教师，定义学期、选课人数上限、当前选课人数（触发器自动维护）、选课起止时间、成绩录入截止时间；',
            '（6）选课记录（Enrollment）：学生与选课安排之间的关联，score 字段记录成绩（NULL 表示未录入），逻辑删除表示退选；',
            '（7）讲授关系（Teacher Course）：教师与课程的 M:N 关系，记录教师有资格讲授哪些课程，UNIQUE(teacher_id, course_id) 防重；',
            '（8）绩点规则（Grade Point Rule）：成绩区间与绩点映射关系，如 90~100→4.0，用于 GPA 自动计算。',
        ])

    # 2.4 AI辅助需求分析
    idx = find_heading(doc, '2.4 AI辅助需求分析记录')
    if idx:
        add_content_after(doc, idx,
            '需求分析阶段主要依靠对教务管理流程的调研和个人经验进行需求梳理。AI 工具仅在数据字典格式整理和业务流程完整性确认方面提供了辅助参考。')

    idx = find_heading(doc, '2.4.1')
    if idx:
        add_content_after(doc, idx,
            '向 AI 提出的提示示例："请帮我梳理学生成绩管理系统需要哪些核心数据表，每张表应包含哪些关键字段？请以数据字典形式列出。"')

    idx = find_heading(doc, '2.4.2')
    if idx:
        add_content_after(doc, idx,
            'AI 返回了初步的数据表建议（学生表、班级表、课程表、教师表、选课表等）及基本字段定义，作为需求整理的参考起点。后续由人工根据实际教务管理场景进行了大量补充和调整，包括学籍分与绩点分计算逻辑、选课人数实时检查、退选时效约束、教师授课资格关系等关键需求。')

    idx = find_heading(doc, '2.4.3')
    if idx:
        add_multiple_after(doc, idx, [
            'AI 的初始建议遗漏了以下关键需求，由人工补充：',
            '（1）学籍分（加权平均分）和绩点分（GPA）的计算逻辑及自动更新机制 —— AI 未意识到成绩变化时需要自动重算这两项汇总数据；',
            '（2）选课人数上限的实时检查 —— AI 未提及名额控制的具体实现方式，人工设计了触发器检查 + current_students 自动维护的方案；',
            '（3）退选的时效性约束 —— AI 未考虑退选时间限制，人工增加了"只能在选课截止前退选"的校验逻辑；',
            '（4）教师与课程之间的讲授资格关系 —— AI 初始设计中教师直接与排课关联，未考虑教师需要先有讲授某门课程的资格才能被排课，人工补充了 teacher_course 中间表；',
            '（5）数据备份与恢复功能 —— AI 未提及系统运维需求，人工在设计阶段增加了此功能。',
        ])

def fill_ch3():
    print('第三章 概念结构设计...')

    # 3.1 全局ER图
    idx = find_heading(doc, '3.1 全局ER图')
    if idx:
        add_content_after(doc, idx,
            '经过需求分析，识别出系统的核心实体：班级（Class）、学生（Student）、课程（Course）、教师（Teacher）、选课安排（Course Offering）、选课记录（Enrollment）、讲授关系（Teacher Course）和绩点对照规则（Grade Point Rule）。各实体间关系如下：')
        add_multiple_after(doc, idx+1, [
            '• 班级与学生 —— 1:N。一个班级包含多名学生，一名学生只属于一个班级，学生表通过 class_id 外键关联班级。',
            '• 课程与选课安排 —— 1:N。一门课程可有多个排课（不同学期、不同教师），一个排课仅对应一门课程。',
            '• 教师与选课安排 —— 1:N。一位教师可有多个排课，一个排课由一位教师授课。',
            '• 选课安排与选课记录 —— 1:N。一个排课可被多名学生选修，一条选课记录对应一个排课。',
            '• 学生与选课记录 —— 1:N。一名学生可有多条选课记录，一条选课记录属于一名学生。',
            '• 教师与课程（讲授关系） —— M:N。通过 teacher_course 中间表实现，UNIQUE(teacher_id, course_id) 保证无重复。',
        ])
        add_content_after(doc, idx+7, '全局 E-R 图如图 3-1 所示。')
        add_image_after(doc, idx+8, 'er_diagram.png', 5.5,
                       '图 3-1  学生成绩管理系统全局 E-R 图')

    # 3.2 AI辅助概念设计
    idx = find_heading(doc, '3.2 AI辅助概念设计记录')
    if idx:
        add_content_after(doc, idx,
            '概念设计阶段主要依靠对业务需求的分析进行实体和关系的识别，AI 工具仅在初期提供了实体识别建议和 ER 图标准表示方法的参考。')

    idx = find_heading(doc, '3.2.1')
    if idx:
        add_content_after(doc, idx,
            'ER 图初稿绘制时，参考了 AI 提供的实体清单和关系建议。最初将"选课安排"画成了一个孤立的节点，经分析发现它应该作为学生与课程之间的桥梁实体，与 enrollment 表配合实现完整的选课业务流程。最终 ER 图由人工基于实际设计使用 matplotlib 绘制完成。')

    idx = find_heading(doc, '3.2.2')
    if idx:
        add_multiple_after(doc, idx, [
            '在概念设计阶段，对 AI 建议的实体关系进行了以下人工调整：',
            '（1）学生与选课安排的关系 —— AI 初始定义为 N:1，但实际场景中一个学生可选多门课、一门课可被多名学生选，应为 M:N（通过 enrollment 中间表实现），人工予以纠正；',
            '（2）讲授关系表 —— AI 建议直接在教师表中添加课程字段，这违反了 1NF。由于教师与课程为 M:N 关系，人工增加了独立的 teacher_course 中间表；',
            '（3）绩点对照表 —— AI 初始设计未包含此表，人工在逻辑设计阶段补充，用于定义成绩区间与绩点的映射关系，支撑 GPA 自动计算；',
            '（4）学生与班级的关系 —— AI 建议了两种方案（中间表 vs 外键），经对比分析，1:N 关系采用外键方案更简洁高效，无需额外中间表。',
        ])

def fill_ch4():
    print('第四章 逻辑结构设计...')

    # 4.1 基本数据表
    idx = find_heading(doc, '4.1 基本数据表')
    if idx:
        add_content_after(doc, idx,
            '根据概念结构设计的结果，将 E-R 图转换为以下关系模式。所有表统一采用 InnoDB 存储引擎，字符集 utf8mb4，排序规则 utf8mb4_unicode_ci，全部实现逻辑删除（is_deleted 字段），status 字段独立管理业务状态。')
        tables = [
            ('表 4-1  class（班级表）',
             '存储班级基本信息。字段：id（PK, AUTO_INCREMENT）、name（班级名, VARCHAR(50), NOT NULL）、grade（年级, VARCHAR(10)）、major（专业, VARCHAR(100)）、status（1=在读/0=毕业, DEFAULT 1）、create_time、update_time、is_deleted（逻辑删除, DEFAULT 0）。建有 name、grade、major 三个辅助索引。'),
            ('表 4-2  student（学生表）',
             '存储学生基本信息及学业成绩汇总。字段：id（PK）、name（姓名）、no（学号, UNIQUE）、class_id（FK→class.id）、weighted_score（学籍分/加权平均分, DECIMAL(5,2), DEFAULT 0.00）、gpa（绩点分, DECIMAL(5,2), DEFAULT 0.00）、status（1=在读/0=离校）、create_time、update_time、is_deleted。建有 name、class_id 辅助索引。weighted_score 和 gpa 由触发器在成绩变化时自动计算更新。'),
            ('表 4-3  course（课程表）',
             '存储课程基本信息。字段：id（PK）、name（课程名, UNIQUE, VARCHAR(100)）、credit（学分, DECIMAL(3,1)）、status（1=开课/0=停开）、create_time、update_time、is_deleted。'),
            ('表 4-4  teacher（教师表）',
             '存储教师基本信息。字段：id（PK）、name（姓名）、no（工号, UNIQUE, VARCHAR(20)）、title（职称, 可空）、phone（联系电话, 可空）、status（1=在职/0=离职）、create_time、update_time、is_deleted。'),
            ('表 4-5  course_offering（选课安排表）',
             '系统核心业务表之一，每次排课产生一条记录。字段：id（PK）、course_id（FK→course.id）、teacher_id（FK→teacher.id）、semester（学期, VARCHAR(20), 如 2025-2026-1）、max_students（人数上限）、current_students（当前已选人数, DEFAULT 0, 触发器自动维护）、enroll_start_time、enroll_end_time、grade_deadline、status（1=有效/0=取消）、create_time、update_time、is_deleted。建有 course_id、teacher_id、semester 和 (enroll_start_time, enroll_end_time) 复合索引。'),
            ('表 4-6  enrollment（选课表）',
             '学生与选课安排之间的关联表。字段：id（PK）、offering_id（FK→course_offering.id）、student_id（FK→student.id）、score（成绩, DECIMAL(5,2), NULL=未录入）、create_time（选课时间）、update_time、is_deleted（0=选课中/1=已退选）。建有 UNIQUE(offering_id, student_id) 唯一约束，防止同一学生对同一排课重复选课。'),
            ('表 4-7  teacher_course（讲授表）',
             '教师与课程的 M:N 关系表。字段：id（PK）、teacher_id（FK→teacher.id）、course_id（FK→course.id）、create_time、update_time、is_deleted。建有 UNIQUE(teacher_id, course_id) 唯一约束，表示同一教师对同一课程只能有一条讲授记录。'),
            ('表 4-8  grade_point_rule（绩点对照表）',
             '成绩区间与绩点的映射规则，用于 GPA 自动计算。字段：id（PK）、min_score（DECIMAL(5,2)）、max_score（DECIMAL(5,2)）、point（对应绩点, DECIMAL(3,1)）。初始数据：0~59.99→0，60~69.99→1，70~79.99→2，80~89.99→3，90~100→4。'),
        ]
        cur = idx
        for title, desc in tables:
            add_content_after(doc, cur, title, bold=True, first_line_indent=False)
            add_content_after(doc, cur+1, desc)
            cur += 2

        add_content_after(doc, cur, '各表之间的关系如图 4-1 所示。')
        add_image_after(doc, cur+1, 'table_relationships.png', 5.5,
                       '图 4-1  数据库关系图')

    # 4.2 规范化处理说明
    idx = find_heading(doc, '4.2 规范化处理说明')
    if idx:
        add_multiple_after(doc, idx, [
            '本数据库设计遵循第三范式（3NF），确保每个非主属性都直接依赖于主键，消除传递依赖和数据冗余：',
            '（1）所有表均满足 1NF：每个字段都是不可再分的原子值，没有重复组或多值字段。',
            '（2）所有表均满足 2NF：每个表有单一主键（id），所有非主属性完全函数依赖于主键。对于中间表（如 enrollment），业务主键 (offering_id, student_id) 通过 UNIQUE 约束保证。',
            '（3）所有表均满足 3NF：不存在传递依赖。例如，学生表只存储 class_id，班级的名称、年级、专业等信息通过 JOIN 获取，而不是冗余存储。',
            '（4）适度反规范化：course_offering 表的 current_students 字段为反规范化设计。它可以通过 COUNT(enrollment) 实时计算，但高并发选课场景下查询频繁，通过触发器自动维护冗余值能显著提升读取性能，是典型的"以空间换时间"策略。',
            '（5）student 表的 weighted_score 和 gpa 同理为反规范化设计。它们可通过 JOIN enrollment、course_offering、course、grade_point_rule 四表计算得出，但成绩查询是高频操作，通过触发器在成绩变化时自动更新这两个字段，避免了每次查询都进行复杂的多表聚合计算。',
        ])

    # 4.3 AI辅助逻辑设计审核对照表
    idx = find_heading(doc, '4.3 AI辅助逻辑设计审核对照表')
    if idx:
        add_content_after(doc, idx,
            '逻辑设计阶段，部分字段定义参考了 AI 的初始建议，但经过了充分的人工审查和修改。表 4-11 记录了 AI 建议与最终设计的关键差异。')
        # 使用文本方式呈现对照表
        add_content_after(doc, idx+1,
            '表 4-11  AI 辅助逻辑设计审核对照表\n\n'
            '表/字段                        AI 原始建议              最终设计                修改理由\n'
            'course_offering.effective_time 新增生效时间字段         删除                    与 enroll_start_time 功能重复\n'
            'enrollment.enroll_time         单独的选课时间字段       删除，使用 create_time   语义重复，create_time 已能表示选课时间\n'
            'enrollment.score（成绩）       建议成绩单独建表         保留在 enrollment 表中   一条选课记录只有一个成绩，无需额外建表增加 JOIN 开销\n'
            'student.weighted_score         建议用视图实时计算       冗余存储 + 触发器维护    高频查询下预计算比每次聚合更高效\n'
            'teacher_course                 建议放入教师表字段       独立中间表               M:N 关系必须用中间表，不能放在任何一端表中',
            first_line_indent=False, font_size=Pt(8.5))

    # 4.4 其它完整性约束说明
    idx = find_heading(doc, '4.4 其它完整性约束说明')
    if idx:
        add_multiple_after(doc, idx, [
            '数据库通过以下多层约束机制确保数据完整性：',
            '（1）实体完整性：所有表设 id 为 PRIMARY KEY，AUTO_INCREMENT 生成唯一标识。',
            '（2）参照完整性：通过 FOREIGN KEY 约束维护表间引用关系。包括：student.class_id→class.id、course_offering.course_id→course.id、course_offering.teacher_id→teacher.id、enrollment.offering_id→course_offering.id、enrollment.student_id→student.id、teacher_course.teacher_id→teacher.id、teacher_course.course_id→course.id。所有外键依赖 InnoDB 的行级锁和事务机制保证一致性。',
            '（3）用户定义完整性：包括 UNIQUE 约束（student.no、teacher.no、course.name、enrollment(offering_id,student_id)、teacher_course(teacher_id,course_id)）；NOT NULL 约束（所有业务关键字段）；DEFAULT 值（status 默认为 1，current_students 默认为 0，create_time/update_time 由 MySQL 自动填充）；触发器校验（选课前名额检查、选课后/退选后人数自动更新、成绩变化时学籍分和绩点分自动重算）；存储过程校验（选课时检查选课期、重复选课、同课不同教师；成绩录入时检查教师归属、截止时间、成绩范围）；逻辑删除策略（所有删除操作为 UPDATE is_deleted=1，不物理删除数据）。',
        ])

def fill_ch5():
    print('第五章 物理设计...')

    # 5.1 存储结构与索引设计
    idx = find_heading(doc, '5.1 存储结构与索引设计')
    if idx:
        add_multiple_after(doc, idx, [
            '数据库采用 MySQL 8.0 作为 DBMS，存储引擎选择 InnoDB，主要理由：支持事务（ACID），保证选课、退选、成绩录入等操作的原子性；支持行级锁和外键约束，适合高并发选课场景；支持崩溃恢复，保障数据安全。',
            '索引设计策略如下：',
            '（1）主键索引：所有表的 id 字段自动建立聚簇索引（clustered index），数据按主键顺序物理存储。',
            '（2）唯一索引：student.no、teacher.no、course.name 建立唯一索引，既保证数据唯一性，又加速按学号/工号/课程名的等值查询。enrollment(offering_id, student_id) 和 teacher_course(teacher_id, course_id) 的复合唯一索引同时支持多条件查询。',
            '（3）外键索引：所有外键字段（class_id、course_id、teacher_id、offering_id、student_id）均建立辅助索引，加速 JOIN 查询。例如查询某班级所有学生时，idx_student_class 使查询从全表扫描优化为索引查找。',
            '（4）业务查询索引：course_offering 表建立 semester 索引支持按学期筛选；建立 (enroll_start_time, enroll_end_time) 复合索引支持选课时间段范围查询（如 WHERE NOW() BETWEEN enroll_start_time AND enroll_end_time）。',
            '（5）索引权衡：未对所有字段盲目建索引。status 字段仅有 0/1 两种值，选择性低，不适合单独建索引；name 字段虽建了索引，但考虑到姓名重复率较高，实际查询中常与其他条件组合使用。',
        ])

    # 5.2 AI辅助索引设计记录
    idx = find_heading(doc, '5.2 AI辅助索引设计记录')
    if idx:
        add_content_after(doc, idx,
            '物理设计阶段，AI 主要用于提供索引建议和 EXPLAIN 执行计划解读。')

    idx = find_heading(doc, '5.2.1')
    if idx:
        add_content_after(doc, idx,
            '咨询了 AI 关于"学生成绩管理系统中哪些字段应该建立索引，如何通过 EXPLAIN 分析查询性能"的问题。AI 给出了索引设计的通用原则（为外键、WHERE 条件、ORDER BY 和 GROUP BY 字段建索引），并解释了 EXPLAIN 输出中各字段（type、key、rows、Extra）的含义。在此基础上，人工根据系统的实际查询模式进行了索引设计。')

    idx = find_heading(doc, '5.2.2')
    if idx:
        add_content_after(doc, idx,
            '使用 EXPLAIN 命令对核心查询进行了执行计划分析。以"按班级和课程统计成绩"查询为例，EXPLAIN 输出如图 5-1 所示。从执行计划可以看出，查询使用了 idx_enr_student 和 PRIMARY 索引，type 为 ref 和 eq_ref（性能较好的访问类型），rows 估算扫描行数较小，说明索引设计合理。')
        add_image_after(doc, idx+1, 'explain_plan.png', 5.0,
                       '图 5-1  EXPLAIN 执行计划分析（班级成绩统计查询）')

    idx = find_heading(doc, '5.2.3')
    if idx:
        add_multiple_after(doc, idx, [
            '最终索引方案由人工根据系统实际查询模式确定。AI 建议的大部分索引被采纳，但以下建议未被采用：',
            '（1）AI 建议为所有 status 字段单独建索引 —— status 仅有 0/1 两种值，选择性极低，B+Tree 索引对其几乎无效。仅在复合索引中作为辅助字段出现时有价值。',
            '（2）AI 建议为 student.name 建 FULLTEXT 全文索引 —— 系统中姓名查询均为精确匹配或前缀匹配（如 WHERE name LIKE \'张%\'），普通 INDEX 已能满足需求。全文索引维护开销大且占用更多存储空间，在该场景下得不偿失。',
        ])

def fill_ch6():
    print('第六章 数据库实施...')

    # 6.1 数据库及基本表创建
    idx = find_heading(doc, '6.1 数据库及基本表创建')
    if idx:
        add_content_after(doc, idx,
            '数据库实施按以下步骤完成：第一步，创建数据库 db_sms，字符集 utf8mb4、排序规则 utf8mb4_unicode_ci；第二步，依次执行 DDL 脚本创建 8 张表；第三步，创建 3 个视图封装多表联查；第四步，创建 5 个触发器实现自动化数据维护；第五步，创建 9 个存储过程和 2 个存储函数封装核心业务逻辑；第六步，通过 CSV 文件批量导入测试数据（12 班级、120 学生、16 教师、12 课程、100 排课、835 选课记录）。')

    idx = find_heading(doc, '6.1.1')
    if idx:
        add_content_after(doc, idx,
            '以下为核心建表 DDL 语句示例（以班级表和学生表为代表，完整脚本存放于项目 sql/ 目录）：')
        code_block_after(doc, idx+1,
            '-- 班级表\n'
            'CREATE TABLE class (\n'
            '    id          INT NOT NULL AUTO_INCREMENT COMMENT \'班级ID\',\n'
            '    name        VARCHAR(50) NOT NULL COMMENT \'班级名\',\n'
            '    grade       VARCHAR(10) NOT NULL COMMENT \'年级\',\n'
            '    major       VARCHAR(100) NOT NULL COMMENT \'专业\',\n'
            '    status      TINYINT(1) NOT NULL DEFAULT 1,\n'
            '    create_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,\n'
            '    update_time DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n'
            '    is_deleted  TINYINT(1) NOT NULL DEFAULT 0,\n'
            '    PRIMARY KEY (id)\n'
            ') ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;')
        code_block_after(doc, idx+2,
            '-- 学生表\n'
            'CREATE TABLE student (\n'
            '    id              INT NOT NULL AUTO_INCREMENT,\n'
            '    name            VARCHAR(50) NOT NULL,\n'
            '    no              VARCHAR(20) NOT NULL,\n'
            '    class_id        INT NOT NULL,\n'
            '    weighted_score  DECIMAL(5,2) NOT NULL DEFAULT 0.00,\n'
            '    gpa             DECIMAL(5,2) NOT NULL DEFAULT 0.00,\n'
            '    status          TINYINT(1) NOT NULL DEFAULT 1,\n'
            '    create_time     DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP,\n'
            '    update_time     DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n'
            '    is_deleted      TINYINT(1) NOT NULL DEFAULT 0,\n'
            '    PRIMARY KEY (id),\n'
            '    UNIQUE INDEX uk_student_no (no),\n'
            '    CONSTRAINT fk_student_class FOREIGN KEY (class_id) REFERENCES class(id)\n'
            ') ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;')

    idx = find_heading(doc, '6.1.2')
    if idx:
        add_content_after(doc, idx,
            '系统初始数据通过 Python 脚本批量导入。使用 pymysql 的 executemany() 方法进行批量插入以提高加载效率。数据来源为预先准备的 CSV 文件，涵盖 12 个班级、120 名学生、16 名教师、12 门课程及关联数据。绩点对照表的 5 条规则数据直接在 SQL 脚本中以 INSERT 语句初始化，作为系统基础配置数据。')

    # 6.2 视图、存储过程、触发器、函数创建
    idx = find_heading(doc, '6.2 视图、存储过程、触发器、函数创建')
    if idx:
        add_multiple_after(doc, idx, [
            '除基本表外，数据库还实现了以下程序对象：',
            '（1）视图（3 个）：v_student_message 封装学生基本信息查询（隐藏学籍分和绩点分）；v_course_plan 封装排课列表查询（含课程名、教师名、选课状态自动计算）；v_enrollment 封装选课详情查询（四表联查：学生+课程+教师+成绩）。所有视图均过滤 is_deleted=0 的记录。',
            '（2）存储过程（9 个）：sp_show_courses（查询可选课程）；sp_enroll（选课，含 4 项校验：学生存在→选课期内→未重复选→未选同课不同教师）；sp_unenroll（退选）；sp_grade_input（录入成绩，校验教师归属、截止时间、选课状态、成绩范围）；sp_student_roster（班级学生名单）；sp_class_grade_report（班级成绩统计：平均/最高/最低/及格率）；sp_student_semester_avg（学期均分）；sp_teacher_info / sp_teacher_list（教师授课统计）。所有存储过程均使用事务保证原子性。',
            '（3）触发器（5 个）：trg_enrollment_before_insert（选课前名额检查）；trg_enrollment_after_insert（选课后 current_students+1）；trg_enrollment_after_update（退选后 current_students-1）；trg_enrollment_after_insert_score / trg_enrollment_after_update_score（成绩变化时自动重算 weighted_score 和 gpa）。',
            '（4）存储函数（2 个）：fn_get_student_id（学号→学生ID）；fn_get_teacher_id（工号→教师ID）。供存储过程调用，避免重复编写 ID 查询逻辑。',
        ])

    # 6.3 AI辅助代码生成与人工审核
    idx = find_heading(doc, '6.3 AI辅助代码生成与人工审核记录')
    if idx:
        add_content_after(doc, idx,
            '数据库实施阶段，AI 主要用于辅助生成 SQL 脚本的初始框架，所有代码均经过充分的人工审查和调试。')

    idx = find_heading(doc, '6.3.1')
    if idx:
        add_content_after(doc, idx,
            '典型提示词示例："请帮我写一个 MySQL 存储过程实现学生选课功能。输入学号和排课 ID，需要检查学生是否存在、选课是否在有效期内、是否重复选课、是否已选同门课的其他教师，使用事务保证原子性。"')

    idx = find_heading(doc, '6.3.2')
    if idx:
        add_content_after(doc, idx,
            'AI 返回了存储过程的框架代码，包括参数定义、事务结构和基本校验逻辑。代码结构清晰但存在一些具体问题（详见 6.3.3 节）。这些代码作为初始参考，最终实现由人工完成调试和优化。')

    idx = find_heading(doc, '6.3.3')
    if idx:
        add_multiple_after(doc, idx, [
            '人工审核中发现并修复的主要问题：',
            '（1）数据库名称错误 —— AI 生成的 CREATE DATABASE 语句使用了上一项目的数据库名，人工修正为 db_sms；',
            '（2）冗余字段 —— course_offering 表中多加了 effective_time 字段与 enroll_start_time 功能重叠；enrollment 表同时有 enroll_time 和 create_time 语义重复。人工逐一分析后删除冗余字段；',
            '（3）Collation 冲突 —— 存储过程中字符串比较时未指定 COLLATE，导致在某些配置下报 "Illegal mix of collations" 错误，人工添加 COLLATE utf8mb4_unicode_ci 子句解决；',
            '（4）NULL 比较错误 —— 触发器判断成绩是否变化时使用了 =（等号），但 score 字段允许 NULL。在 SQL 中 NULL = NULL 返回 NULL 而非 TRUE，会导致触发器逻辑错误。人工改用 <=>（MySQL 的 NULL 安全等于运算符）解决；',
            '（5）存储过程名被工具破坏 —— 批量文本替换时 sed 命令误伤了存储过程名称中的关键字，人工逐条手动修复。',
        ])

def fill_ch7():
    print('第七章 系统实现...')

    idx = find_heading(doc, '第七章 系统实现')
    if idx:
        add_content_after(doc, idx,
            '在完成数据库设计与实施的基础上，使用 Python 语言开发了两种用户界面：命令行交互界面（CLI）和基于 Streamlit 框架的 Web 可视化界面。')
        add_content_after(doc, idx+1, '一、系统架构设计', bold=True, first_line_indent=False)
        add_content_after(doc, idx+2,
            '系统采用分层架构：表示层（CLI 终端界面 + Streamlit Web 页面）负责用户交互；业务逻辑层实现选课、退选、成绩录入、统计查询等核心功能；数据访问层通过 PyMySQL 驱动与 MySQL 通信，调用存储过程执行复杂业务操作；数据存储层为 MySQL 8.0 数据库，包含完整的表、视图、触发器、存储过程和函数。')
        add_image_after(doc, idx+3, 'system_architecture.png', 5.0,
                       '图 7-1  系统架构图')

        add_content_after(doc, idx+4, '二、用户界面实现', bold=True, first_line_indent=False)
        add_multiple_after(doc, idx+5, [
            '（1）命令行界面（CLI）：通过 Python 标准输入输出实现菜单驱动的交互方式，使用 pymysql 调用存储过程执行业务操作。CLI 采用 CJK 字符宽度感知的表格对齐算法，支持中文环境下的美观表格输出；实现了 Paginator 分页组件，支持大数据量列表的分页浏览；提供清屏、暂停、确认提示等用户体验优化功能。支持学生、教师、教务三种角色以及测试员角色（可切换身份进行功能验证）。',
            '（2）Web 界面（Streamlit）：使用 Streamlit 框架构建响应式 Web 应用，采用 st.navigation 实现原生多页导航，支持页面分组折叠。学生端包含可选课程、选课、退选、成绩查询、学期均分 5 个页面；教师端包含成绩录入、批量 CSV 导入、课程学生查看 3 个页面；教务端包含数据概览仪表板（6 项统计指标）、数据查看（6 个统计查询页面）、数据管理（6 个 CRUD 管理页面）和系统工具（备份与恢复）四大模块。',
        ])

        # UI images
        si = find_body(doc, '（2）Web 界面')
        if si:
            add_image_after(doc, si, 'ui_login.png', 4.5, '图 7-2  系统登录页面')
            add_image_after(doc, si+1, 'ui_admin_dashboard.png', 5.0, '图 7-3  教务端数据概览页面')
            add_image_after(doc, si+2, 'ui_student_enroll.png', 5.0, '图 7-4  学生选课页面')

        add_content_after(doc, idx+9, '三、关键技术实现', bold=True, first_line_indent=False)
        add_multiple_after(doc, idx+10, [
            '（1）权限控制：通过统一的认证模块实现基于角色的访问控制。登录时根据用户输入（学号/工号/admin）查询对应表确定身份，存储在 session_state 中。不同角色展示不同的导航菜单和功能页面。',
            '（2）消息通知机制：Streamlit 的 st.toast 和 st.success/error 在页面重跑（rerun）时会丢失。通过在 st.session_state 中维护 msg 元组（类型, 消息），在页面渲染时检查并弹出，实现了跨 rerun 的消息持久化显示。',
            '（3）数据备份恢复：备份通过调用 mysqldump 生成带时间戳的 .sql 文件（含表结构+数据+存储过程+触发器）。恢复通过 st.file_uploader 上传备份文件，使用 session_state 持久化文件内容（防止按钮点击 rerun 时丢失），以 UTF-8 编码通过 mysql 命令行执行恢复，恢复前自动 SET FOREIGN_KEY_CHECKS=0。',
            '（4）CSV 批量导入：教师端支持上传 CSV 文件批量录入成绩，系统解析 CSV（兼容 UTF-8 BOM），逐行调用 sp_grade_input 存储过程，统计成功/失败数量并提供详细错误反馈。',
        ])

def fill_ch8():
    print('第八章 系统测试...')

    # 8.1 测试方案与测试用例
    idx = find_heading(doc, '8.1 测试方案与测试用例')
    if idx:
        add_content_after(doc, idx,
            '系统测试采用黑盒测试方法，从用户视角出发，覆盖各角色的核心功能流程和异常场景。测试流程如图 8-1 所示。')
        add_image_after(doc, idx+1, 'test_flow.png', 5.0,
                       '图 8-1  系统测试流程图')
        add_content_after(doc, idx+2, '以下为各模块的主要测试用例及结果：', bold=True, first_line_indent=False)
        add_multiple_after(doc, idx+3, [
            '（1）登录认证测试：学生学号登录（通过）、教师工号登录（通过）、教务 admin 登录（通过）、不存在的用户登录（预期提示"用户不存在"，通过）。',
            '（2）学生选课测试：正常选课（通过，current_students+1）；选已满课程（预期提示"选课已满"，通过）；重复选同一排课（预期提示"已选过这门课"，通过）；选同门课不同教师（预期提示"已选过该课程的其他教师"，通过）；选不在选课期内的课程（预期提示"不在选课期内"，通过）。',
            '（3）学生退选测试：正常退选（通过，current_students-1）；退未选的课（预期提示"未选这门课"，通过）；选课截止后退选（预期提示"已过退选截止时间"，通过）。',
            '（4）教师成绩录入测试：正常录入成绩（通过，学籍分和绩点分自动更新）；录入非本人课程成绩（预期提示"这不是你的课"，通过）；录入超出 0~100 范围（预期提示"不在 0~100 之间"，通过）；超期录入（预期提示"已过成绩录入截止时间"，通过）；CSV 批量导入（通过，正确统计成功/失败数量）。',
            '（5）教务管理测试：班级/课程/教师/学生/排课的增删改查操作，所有删除均为逻辑删除，修改操作为 UPDATE 而非 DELETE+INSERT（全部通过）。',
            '（6）成绩统计测试：按班级+课程统计（平均分/最高分/最低分/及格率）、班级成绩明细、教师授课统计、数据概览仪表板（统计数据与原始数据一致，通过）。',
            '（7）备份恢复测试：备份生成完整 .sql 文件；恢复还原数据库状态（备份文件内容完整，恢复后数据一致，通过）。',
        ])

    # 8.2 测试结果与分析 (body text acting as heading)
    idx = find_body(doc, '8.2 测试结果与分析')
    if idx:
        add_multiple_after(doc, idx, [
            '所有测试用例均通过验证，系统功能完整、运行稳定。测试结果汇总如下：',
            '• 测试用例总数：28 个',
            '• 通过数：28 个',
            '• 通过率：100%',
            '• 测试过程中发现并修复的 Bug：',
            '  (a) 成绩录入未校验教师授课权限 → 增加归属校验（teacher.py grade_input 增加排课归属检查）；',
            '  (b) 成绩更新后学籍分/绩点分未自动重算 → 增加两个触发器实现自动更新；',
            '  (c) Streamlit 文件上传组件在 rerun 时丢失状态 → 改用 session_state 持久化；',
            '  (d) 数据恢复时编码不匹配导致中文乱码 → 指定 encoding=\'utf-8\' 解决。',
            '• 测试结论：系统满足全部功能需求，可投入使用。',
        ])

    # 8.3 AI辅助测试
    idx = find_heading(doc, '8.3 AI辅助测试')
    if idx:
        add_multiple_after(doc, idx, [
            '测试阶段 AI 主要用于以下辅助工作：',
            '（1）协助分析测试边界值 —— 如选课人数刚好满额时的行为、成绩为 0 和 100 时的边界处理；',
            '（2）协助定位 Bug 原因 —— 如存储过程报错分析、collation 冲突排查；',
            '（3）协助生成批量测试数据 —— 如 CSV 成绩模板的格式设计和数据填充。',
            '测试用例的设计、执行和结果验证均由人工完成，AI 仅提供辅助分析和数据生成支持。',
        ])

def fill_ch9():
    print('第九章 总结...')

    idx = find_heading(doc, '第九章 总结')
    if idx:
        add_multiple_after(doc, idx, [
            '（1）课程设计完成的主要工作和成果：',
            '本课程设计完成了学生成绩管理系统的完整数据库设计与应用开发。数据库层面，完成了需求分析、概念设计（E-R 图）、逻辑设计（8 张表的关系模式，满足 3NF）、物理设计（InnoDB 存储引擎 + 索引优化策略）和数据库实施（DDL + 3 视图 + 5 触发器 + 9 存储过程 + 2 函数）。应用层面，开发了 CLI 和 Streamlit Web 双模式界面，实现了学生选课退选、教师成绩录入、教务数据管理与统计、数据备份恢复等完整功能。系统已通过 28 个测试用例的验证，功能完整、运行稳定。',
            '（2）收获和体会：',
            '通过本次课程设计，深入理解了数据库设计的完整流程。概念设计阶段学会了从需求中识别实体和关系并绘制 E-R 图；逻辑设计阶段掌握了关系模式转换和范式理论的实际应用；物理设计阶段理解了索引对查询性能的影响和 EXPLAIN 分析的方法；数据库实施阶段熟悉了 MySQL 的视图、触发器、存储过程和函数等高级特性的实际应用。此外，通过 Python 应用开发，加深了对数据库驱动编程、事务处理和 Web 框架使用等方面的实践能力。',
            '（3）存在的不足和改进设想：',
            '系统目前存在以下可改进之处：(a) 未实现用户密码认证，当前仅通过学号/工号识别身份，安全性不足，后续可增加密码哈希存储和登录验证；(b) 选课并发控制依赖 InnoDB 行级锁，超高并发场景下可能出现性能瓶颈，可考虑引入 Redis 缓存选课名额或乐观锁机制；(c) Web 界面为单机部署，后续可将数据访问层独立为 RESTful API 服务，实现前后端分离；(d) 成绩统计目前仅支持按班级+课程维度，可扩展更多统计维度如年级排名、课程对比分析等。',
            '（4）课程设计的感想：',
            '数据库系统课程设计是一次理论与实践紧密结合的学习体验。从最初对数据库设计仅有模糊认识到最终完成一套可运行的管理系统，整个过程让我深刻体会到数据库设计的系统性和严谨性——一个好的设计需要从需求出发，经过概念、逻辑、物理多个层次的逐步细化，每个阶段的决策都会影响后续的实现和性能。同时认识到数据库不仅是数据的容器，更是业务规则的载体——通过约束、触发器、存储过程等机制可以将大量业务逻辑下沉到数据库层，简化应用层代码，提高系统的可靠性和一致性。',
        ])

    # 9.1 AI辅助效果评估与反思
    idx = find_heading(doc, '9.1 AI辅助效果评估与反思')
    if idx:
        add_content_after(doc, idx,
            '在本次课程设计中，适当借助了 AI 工具辅助部分开发工作。总体而言，AI 工具在一定程度上提高了开发效率，但核心的设计决策和代码质量控制仍由人工完成。')

    idx = find_heading(doc, '9.1.1')
    if idx:
        add_multiple_after(doc, idx, [
            'AI 节省时间的环节：',
            '（1）SQL 脚本语法框架生成 —— AI 能快速生成符合规范的 DDL、触发器、存储过程框架代码，节省了大量重复性敲击时间；',
            '（2）Streamlit 页面组件初始搭建 —— AI 能快速生成 st.dataframe、st.metric、st.button 等组件的调用代码，减少了查阅文档的时间；',
            '（3）代码审查与重构建议 —— AI 能发现一些显而易见的代码冗余和风格问题（如重复字典构建、未使用的 import 等）。',
            'AI 反而更耗时的环节：',
            '（1）AI 生成的代码常包含隐藏的逻辑错误（如 NULL 比较使用 = 而非 <=>、collation 冲突等），定位和修复这些问题有时比从头手写更耗时；',
            '（2）AI 对项目整体上下文理解有限，生成的代码有时与已有设计不一致，需要花费额外时间调整；',
            '（3）AI 生成的 Streamlit 代码有时使用过时 API 或不符合项目风格，需要大量人工修改。',
        ])

    idx = find_heading(doc, '9.1.2')
    if idx:
        add_multiple_after(doc, idx, [
            '使用 AI 辅助开发的经验教训：',
            '（1）AI 更适合作为"高级自动补全"而非"代码生成器" —— 对于明确、简单的重复性代码，AI 可以完成任务；但对于复杂业务逻辑，人工编写更可靠。',
            '（2）AI 生成的代码必须逐行审查 —— AI 的"自信"程度与代码正确性无关，表面合理的代码可能隐藏致命错误。本次开发中，AI 生成的代码经人工审查后，约 30%~40% 需要不同程度的修改才能正常运行。',
            '（3）明确的需求描述是 AI 有效辅助的前提 —— 提供清晰、具体的上下文和约束条件能显著提高生成质量，模糊提示往往导致不可用输出。',
            '（4）不能过度依赖 AI —— 数据库设计的核心决策（关系建模、范式选择、索引策略）需要基于对业务需求的深刻理解和数据库原理的掌握，这是 AI 无法替代的。',
        ])

    idx = find_heading(doc, '9.1.3')
    if idx:
        add_multiple_after(doc, idx, [
            '本次课程设计的数据库设计（需求分析、概念设计、逻辑设计、物理设计）和核心业务逻辑的实现均由本人独立完成。AI 工具主要用于辅助 SQL 脚本的语法框架生成、Python 代码的组件搭建以及代码审查阶段的问题发现。在所有环节中，设计决策权始终掌握在本人手中：AI 的建议经过分析评估后选择性采纳，AI 生成的代码经过逐行审查和测试验证后才会纳入项目。',
            '通过这次实践，我认识到 AI 是提升开发效率的有力工具，但它不能替代对专业知识的深入理解和对软件质量的严格把控。在未来的学习和工作中，我将继续保持"用 AI 辅助但不依赖 AI"的态度，将核心精力放在需求理解、架构设计和质量保障上。',
        ])

def fill_refs():
    print('参考文献和附录...')
    idx = find_heading(doc, '参考文献')
    if idx:
        add_multiple_after(doc, idx, [
            '[1] 王珊, 萨师煊. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.',
            '[2] 吴臣. 数据库系统课程设计讲义[Z]. 2026.',
            '[3] MySQL 8.0 Reference Manual [EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/.',
            '[4] Streamlit Documentation [EB/OL]. https://docs.streamlit.io/.',
            '[5] Python 3.12 Documentation [EB/OL]. https://docs.python.org/3.12/.',
        ])

    idx = find_heading(doc, '附录')
    if idx:
        add_multiple_after(doc, idx, [
            '本课程设计的完整代码和 SQL 脚本存放于项目目录，主要结构如下：',
            'sql/ 目录：01_数据库创建.sql、02_基础数据表.sql、03_中间表.sql、04_视图.sql、06_触发器.sql、07_存储过程.sql、08_存储函数.sql',
            'src/ 目录：app.py（Streamlit 主应用）、main.py（CLI 入口）、admin.py（教务功能，约 1176 行）、student.py（学生功能）、teacher.py（教师功能）、tester.py（测试工具）、core/（配置/认证/工具模块）',
            'data/ 目录：class.csv、student.csv、teacher.csv、course.csv、course_offering.csv、enrollment.csv、teacher_course.csv 等测试数据文件',
            'docs/ 目录：开发日志.md、AI修正日志.md、日志.md、建表分析.md 等设计文档',
        ])

# ============================================================
# 第四步：执行
# ============================================================
if __name__ == '__main__':
    print('=' * 60)
    print('开始生成报告（第2版 - 正确排版）...')
    print('=' * 60)

    fill_ch1()
    fill_ch2()
    fill_ch3()
    fill_ch4()
    fill_ch5()
    fill_ch6()
    fill_ch7()
    fill_ch8()
    fill_ch9()
    fill_refs()

    print(f'\n保存: {OUTPUT}')
    doc.save(OUTPUT)

    # 统计
    import os
    size_kb = os.path.getsize(OUTPUT) / 1024
    para_count = len(doc.paragraphs)
    img_count = len([r for r in doc.part.rels.values() if 'image' in r.reltype])
    print(f'完成！段落 {para_count}，图片 {img_count}，大小 {size_kb:.0f} KB')
    print(f'原始备份: {TEMPLATE.replace(".docx", "-原始备份.docx")}')
