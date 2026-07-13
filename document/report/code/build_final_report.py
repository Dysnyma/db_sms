# -*- coding: utf-8 -*-
"""从原始备份模板出发，精心撰写完整报告"""
import sys, io; sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

SRC = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿-原始备份.docx'
OUT = r'E:\02_Courses\db_lab\document\报告\《数据库系统课程设计》说明书规范-蔡坤灿.docx'
IMG = r'E:\02_Courses\db_lab\document\报告\images'

doc = Document(SRC)

# ============================================================
# 第一步：清理模板（删除占位段落和示例表格）
# ============================================================
PLACEHOLDERS = [
    '正文排版要求', '字号，5号', '图和表，要有标号', '图表和正文的引用',
    '简述各阶段分别用AI做了什么', '提示词/Prompt', '截图或文字引用',
    'AI遗漏或理解错误了哪些内容', 'AI生成的实体/关系中，手动调整',
    '基本表及其属性、数据类型及长度', '表4-1 XXX表',
    'CHECK约束、冗余列、冗余表的数据来源说明',
    '截图', '贴执行计划截图', '最终SQL脚本',
    '明确指定用触发器/存储过程/视图等', '保留原样粘贴',
    '写了哪些Bug、如何改的', '是否使用AI生成测试数据',
    '课程设计中遇到的主要问题', '创新和得意之处',
    '课程设计中存在的不足', '课程设计的感想和心得体会',
    '本项目中哪些核心设计环节必须由人完成',
    '这里如果是纯数据库实现', '表4-2 AI辅助逻辑设计审核对照表',
    '豆包最开始把学生', '我最初画的ER图',
]

to_del = []
for i, p in enumerate(doc.paragraphs):
    s = p.style.name
    if s == 'annotation text':
        to_del.append(i); continue
    if s not in ('Heading 1', 'Heading 2', 'Heading 3', 'toc 1', 'toc 2', 'toc 3'):
        for m in PLACEHOLDERS:
            if m in p.text: to_del.append(i); break

for i in reversed(to_del):
    doc.paragraphs[i]._element.getparent().remove(doc.paragraphs[i]._element)
print(f'清理段落: {len(to_del)}')

# 删除模板示例表格
tbl_del = []
for ti, t in enumerate(doc.tables):
    h = ''.join(c.text for c in t.rows[0].cells)
    if '字段名称' in h and '字段类型' in h: tbl_del.append(ti)
    if 'AI生成的原始定义' in h: tbl_del.append(ti)
for ti in reversed(tbl_del):
    doc.tables[ti]._tbl.getparent().remove(doc.tables[ti]._tbl)
print(f'清理表格: {len(tbl_del)}')

# ============================================================
# 第二步：工具函数
# ============================================================
def find_h(text):
    """查找包含文本的 Heading 段落索引"""
    for i, p in enumerate(doc.paragraphs):
        if text in p.text and p.style.name.startswith('Heading'):
            return i
    return None

def find_t(text):
    """查找包含文本的任意段落索引"""
    for i, p in enumerate(doc.paragraphs):
        if text in p.text: return i
    return None

def _mkpara(text, font_cn='宋体', font_en='Times New Roman', size=Pt(10.5),
            bold=False, align=None, indent=True, spacing=1.25, space_before=0, space_after=0):
    """创建格式完整的段落，返回其 XML 元素"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.first_line_indent = Pt(21) if indent else Pt(0)
    if align is not None: p.alignment = align

    if text:
        run = p.add_run(text)
        run.font.size = size
        run.font.name = font_en
        run.bold = bold
        rPr = run._r.get_or_add_rPr()
        rf = OxmlElement('w:rFonts')
        rf.set(qn('w:eastAsia'), font_cn)
        rf.set(qn('w:ascii'), font_en); rf.set(qn('w:hAnsi'), font_en)
        rPr.insert(0, rf)
    return p._element

def add_p(ref_idx, text, **kw):
    """在 ref_idx 段落后插入一段正文"""
    el = _mkpara(text, **kw)
    el.getparent().remove(el)
    doc.paragraphs[ref_idx]._element.addnext(el)

def add_n(ref_idx, texts, **kw):
    """连续插入多段正文"""
    cur = ref_idx
    for t in texts:
        add_p(cur, t, **kw); cur += 1

def add_h2(ref_idx, text):
    """插入 Heading 2 样式标题"""
    el = _mkpara(text, font_cn='黑体', font_en='Arial', size=Pt(14), bold=True,
                 indent=False, spacing=1.25, space_before=12, space_after=6)
    el.getparent().remove(el)
    # 设置 Heading 2 样式
    pPr = el.find(qn('w:pPr'))
    if pPr is None: pPr = OxmlElement('w:pPr'); el.insert(0, pPr)
    ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), 'Heading2'); pPr.append(ps)
    doc.paragraphs[ref_idx]._element.addnext(el)

def add_h3(ref_idx, text):
    """插入 Heading 3 样式标题"""
    el = _mkpara(text, font_cn='黑体', font_en='Arial', size=Pt(12), bold=True,
                 indent=False, spacing=1.25, space_before=6, space_after=4)
    el.getparent().remove(el)
    pPr = el.find(qn('w:pPr'))
    if pPr is None: pPr = OxmlElement('w:pPr'); el.insert(0, pPr)
    ps = OxmlElement('w:pStyle'); ps.set(qn('w:val'), 'Heading3'); pPr.append(ps)
    doc.paragraphs[ref_idx]._element.addnext(el)

def add_img(ref_idx, name, width=5.0, caption=None):
    """插入图片+图标题"""
    path = os.path.join(IMG, name)
    if not os.path.exists(path): print(f'  [WARN] {name} not found'); return

    ref_el = doc.paragraphs[ref_idx]._element

    # 图片段落
    ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.first_line_indent = Pt(0)
    r = ip.add_run()
    try:
        from PIL import Image; pil = Image.open(path)
        asp = pil.size[1]/pil.size[0]; r.add_picture(path, width=Inches(width), height=Inches(width*asp))
    except: r.add_picture(path, width=Inches(width))
    ie = ip._element; ie.getparent().remove(ie); ref_el.addnext(ie)

    # 图标题
    if caption:
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.first_line_indent = Pt(0)
        cr = cp.add_run(caption); cr.font.size = Pt(9); cr.font.name = 'Times New Roman'
        rPr = cr._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts')
        rf.set(qn('w:eastAsia'), '宋体'); rPr.insert(0, rf)
        ce = cp._element; ce.getparent().remove(ce); ie.addnext(ce)

    print(f'  [IMG] {name}')

def add_tbl(ref_idx, headers, rows, col_widths, caption=None):
    """插入格式化 Word 表格"""
    ref_el = doc.paragraphs[ref_idx]._element

    if caption:
        ce = _mkpara(caption, size=Pt(9), bold=True, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        ce.getparent().remove(ce); ref_el.addnext(ce); ref_el = ce

    ncols = len(headers); nrows = len(rows) + 1
    tbl = doc.add_table(rows=nrows, cols=ncols, style='Table Grid')
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]; c.text = ''
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.font.size = Pt(9); r.font.name = 'Times New Roman'; r.bold = True
        rPr = r._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '宋体'); rPr.insert(0, rf)
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), '1565C0'); shd.set(qn('w:val'), 'clear')
        c._tc.get_or_add_tcPr().append(shd); r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = ''
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val)); r.font.size = Pt(8.5); r.font.name = 'Times New Roman'
            rPr = r._r.get_or_add_rPr(); rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), '宋体'); rPr.insert(0, rf)

    # 列宽
    if col_widths:
        for ri in range(nrows):
            for ci, w in enumerate(col_widths):
                tbl.rows[ri].cells[ci].width = Inches(w)

    te = tbl._tbl; te.getparent().remove(te); ref_el.addnext(te)
    return tbl

def add_code(ref_idx, code):
    """插入等宽代码块"""
    el = _mkpara(code, font_en='Consolas', size=Pt(7.5), indent=False, spacing=1.1)
    el.getparent().remove(el)
    doc.paragraphs[ref_idx]._element.addnext(el)

# ============================================================
# 第三步：逐章撰写
# ============================================================

def ch1():
    print('第一章...')
    # 1.1 开发背景
    i = find_h('1.1 开发背景')
    if i:
        add_n(i, [
            '随着高校招生规模持续扩大，传统纸质成绩管理方式效率低下、易出错，难以满足教务管理的信息化需求。学生成绩管理作为教务管理的核心环节，涉及选课、成绩录入、绩点计算、统计分析等多个环节，数据量大、关联复杂，亟需一套规范化的数据库系统进行统一管理。',
            '本课程设计选题为"学生成绩管理系统"（Student Grade Management System，数据库名 db_sms），旨在通过数据库设计理论与工程实践相结合的方式，完成从需求分析、概念设计、逻辑设计、物理设计到数据库实施与系统开发的全流程实践。系统采用 MySQL 作为 DBMS，使用 Python 语言开发命令行和基于 Streamlit 框架的 Web 双模式界面，支持学生、教师、教务管理员三种角色的差异化功能。',
        ])

    # 1.2 设计目标
    i = find_h('1.2 设计目标')
    if i:
        add_n(i, [
            '本系统的设计目标如下：',
            '（1）完成学生成绩管理系统的完整数据库设计，包括概念模型（E-R 图）、逻辑模型（关系模式）和物理模型（表结构、索引、视图等）。',
            '（2）实现核心业务功能：学生选课与退选、教师成绩录入、教务数据管理与统计分析、学籍分与绩点分的自动计算更新。',
            '（3）保证数据完整性和一致性，通过主键约束、外键约束、唯一约束、触发器和事务机制等多层机制确保数据质量。',
            '（4）采用逻辑删除（is_deleted 字段）策略，保留历史数据的可追溯性；status 字段独立管理业务状态，两者职责明确分离。',
            '（5）提供友好的用户交互界面，支持命令行（CLI）和 Web（Streamlit）双模式运行。',
            '（6）实现数据库备份与恢复功能，保障数据安全。',
        ])

    # 1.3 软硬件环境与工具
    i = find_h('1.3 软硬件环境与工具')
    if i:
        add_p(i, '系统开发与运行的软硬件环境配置如表 1-1 所示。')
        add_tbl(i+1,
            ['类别', '名称', '版本/说明'],
            [['操作系统', 'Windows 11 Home China', '24H2'],
             ['数据库', 'MySQL', '8.0'],
             ['开发语言', 'Python', '3.12'],
             ['Web 框架', 'Streamlit', '1.x'],
             ['数据库驱动', 'PyMySQL', '最新版'],
             ['开发工具', 'VS Code', '—'],
             ['绘图工具', 'Draw.io / matplotlib', '—'],
             ['备份工具', 'mysqldump', 'MySQL 自带']],
            [1.5, 2.0, 1.8], caption='表 1-1  开发环境配置')
        add_p(i+2, '系统采用分层架构设计，如图 1-1 所示。')
        add_img(i+3, 'system_architecture.png', 5.0, '图 1-1  系统四层架构图')

    # 1.4 AI辅助工具清单
    i = find_h('1.4 AI辅助工具清单')
    if i:
        add_p(i, '在本次课程设计中，在部分重复性代码编写和问题排查环节使用了 AI 工具作为辅助，以下记录所用工具及使用方式。')

    i = find_h('1.4.1')
    if i:
        add_n(i, [
            '本课程设计主要使用了以下 AI 辅助工具：',
            '（1）Claude Code（Anthropic 公司开发的 AI 编程助手，使用 Claude Opus 4.8 模型）：主要用于 SQL 脚本语法检查、部分 Python 代码框架生成以及 Streamlit 页面组件的初始搭建。AI 给出的代码均经过逐行审核与测试验证，根据实际需求进行了修改后才纳入项目。',
            '（2）AI 对话助手（Claude）：用于解答数据库设计中的部分概念性问题（如范式理论、事务隔离级别等），以及在 ER 图初稿绘制和 SQL 优化方面提供参考建议。',
        ])

    i = find_h('1.4.2')
    if i:
        add_p(i,
            '各阶段 AI 使用概况：需求分析阶段，用 AI 辅助整理了数据字典的格式；概念设计阶段，参考了 AI 对 ER 图实体划分的建议；逻辑设计阶段，用 AI 辅助检查了建表 DDL 的语法规范性；物理设计阶段，参考了 AI 关于索引设计的通用原则和 EXPLAIN 解读；数据库实施阶段，部分存储过程和触发器的初始框架由 AI 辅助生成；系统实现阶段，Streamlit 部分页面组件的 HTML/CSS 结构参考了 AI 建议；测试阶段，用 AI 辅助生成了部分 CSV 测试数据。所有 AI 参与的内容均经过人工审查、修改和验证。')

def ch2():
    print('第二章...')
    # 2.1
    i = find_h('2.1 系统业务流程')
    if i:
        add_p(i, '学生成绩管理系统面向三类用户角色，各类角色的核心业务需求分析如下。')
        add_p(i+1, '一、学生端需求', bold=True, indent=False)
        add_n(i+2, [
            '（1）浏览可选课程：查看当前学期在选课期内、尚有名额的课程信息，包括课程名、学分、授课教师、剩余名额等。',
            '（2）选课：在选课开放时间内选择课程，系统自动检查名额是否已满、是否重复选课、是否已选同门课的其他教师。',
            '（3）退选：在选课截止时间前退选已选课程，退选后名额自动释放。',
            '（4）成绩查询：查看个人所有已修课程的成绩、学籍分（加权平均分）和绩点分（GPA）。',
            '（5）学期均分查询：按学期统计个人平均成绩。',
        ])
        j = find_t('一、学生端需求')
        if j:
            add_p(j+5, '二、教师端需求', bold=True, indent=False)
            add_n(j+6, [
                '（1）成绩录入：对自己授课班级的学生逐条录入成绩（0~100 分），系统校验教师是否拥有该排课的授课权限以及是否在成绩录入截止时间内。',
                '（2）批量导入：支持通过 CSV 文件批量导入学生成绩，提高大班授课时的录入效率。',
                '（3）查看课程学生：查看自己所授课程的学生名单及成绩录入状态。',
            ])
            add_p(j+9, '三、教务管理员端需求', bold=True, indent=False)
            add_n(j+10, [
                '（1）数据概览：查看系统中班级数、学生数、教师数、课程数、排课数、选课记录数等统计信息。',
                '（2）基础数据管理：对班级、课程、教师、学生信息进行增删改查操作，所有删除采用逻辑删除。',
                '（3）排课管理：创建、修改、删除选课安排，指定课程、授课教师、学期、人数上限以及选课和成绩录入的时间节点。',
                '（4）选课管理：查看所有选课记录，支持管理员对异常选课进行强制退选。',
                '（5）成绩统计：按班级和课程维度统计平均分、最高分、最低分、及格率。',
                '（6）成绩明细查询：按班级查看所有学生的各科成绩明细。',
                '（7）教师信息查询：查看教师个人详情及授课统计，以及全部教师汇总列表。',
                '（8）数据备份与恢复：支持一键备份数据库（含表结构、数据、存储过程和触发器），以及从备份文件恢复数据。',
            ])

    # 2.2 数据流图
    i = find_h('2.2 数据流图')
    if i:
        add_p(i, '系统顶层数据流图（DFD）如图 2-1 所示，展示了学生、教师、教务管理员三个外部实体与系统核心处理过程之间的数据流动关系。')
        add_img(i+1, 'data_flow_diagram.png', 5.0, '图 2-1  系统数据流图（DFD）')
        add_img(i+2, 'function_modules.png', 5.0, '图 2-2  系统功能模块图')

    # 2.3 数据字典
    i = find_h('2.3 数据字典')
    if i:
        add_p(i, '系统核心数据项定义如下：')
        add_n(i+1, [
            '（1）学生（Student）：学号（唯一标识）、姓名、所属班级、学籍分（加权平均分）、绩点分（GPA）、在读状态。',
            '（2）班级（Class）：班级名称、年级、专业、在读/毕业状态。',
            '（3）课程（Course）：课程名称（唯一）、学分、开课/停开状态。',
            '（4）教师（Teacher）：工号（唯一标识）、姓名、职称、联系电话、在职/离职状态。',
            '（5）选课安排（Course Offering）：关联课程和授课教师，定义学期、选课人数上限、当前选课人数（由触发器自动维护）、选课起止时间、成绩录入截止时间。',
            '（6）选课记录（Enrollment）：学生与选课安排之间的关联记录，score 字段记录成绩（NULL 表示尚未录入），逻辑删除表示退选。设有 UNIQUE(offering_id, student_id) 约束防止重复选课。',
            '（7）讲授关系（Teacher Course）：教师与课程的 M:N 关系，记录教师有资格讲授哪些课程。UNIQUE(teacher_id, course_id) 保证同一教师对同一课程只有一条资格记录。',
            '（8）绩点规则（Grade Point Rule）：定义成绩区间与绩点的映射关系（如 90~100→4.0，80~89.99→3.0 等），用于 GPA 自动计算。',
        ])

    # 2.4 AI辅助需求分析
    i = find_h('2.4 AI辅助需求分析记录')
    if i:
        add_p(i, '需求分析阶段主要依靠对教务管理实际流程的了解和调研来梳理需求。AI 工具仅在数据字典格式整理和业务流程完整性检查方面提供了辅助参考，所有核心需求的确定均由人工完成。')

    i = find_h('2.4.1')
    if i:
        add_p(i, '向 AI 提出的提示示例如下："请帮我梳理一个学生成绩管理系统需要哪些核心数据表，每张表应该包含哪些关键字段？请以数据字典的形式列出。"')

    i = find_h('2.4.2')
    if i:
        add_p(i, 'AI 返回了初步的数据表建议，包括学生表、班级表、课程表、教师表、选课表等，以及各表的基本字段定义。这些建议作为需求整理的参考起点，后续由人工根据实际教务管理场景进行了大量补充和调整，主要包括：学籍分与绩点分的计算逻辑及自动更新机制、选课人数上限的实时检查与名额维护、退选的时效性约束、教师与课程之间的讲授资格关系、以及数据备份与恢复功能。')

    i = find_h('2.4.3')
    if i:
        add_n(i, [
            'AI 初始建议中存在以下不足，由人工进行了修正和补充：',
            '（1）AI 未考虑学籍分（加权平均分）和绩点分（GPA）的计算逻辑，更未意识到成绩变化时需要自动重算这两项汇总数据。人工设计了基于学分的加权算法，并通过触发器实现自动更新。',
            '（2）AI 未提及选课名额控制的具体实现方式。人工设计了在 course_offering 表中维护 current_students 冗余字段，通过插入前触发器检查名额、插入后触发器自动增加人数的方案。',
            '（3）AI 未考虑退选的时间限制。人工在退选存储过程中增加了"只能在选课截止前退选"的校验逻辑。',
            '（4）AI 未涉及教师与课程之间的讲授资格关系。人工补充了 teacher_course 中间表，确保只有具备讲授某门课程资格的教师才能被安排授课。',
            '（5）AI 未提及系统运维需求。人工在设计阶段增加了基于 mysqldump 的数据库备份与恢复功能。',
        ])

def ch3():
    print('第三章...')
    # 3.1
    i = find_h('3.1 全局ER图')
    if i:
        add_p(i, '经过需求分析，识别出系统的核心实体包括：班级（Class）、学生（Student）、课程（Course）、教师（Teacher）、选课安排（Course Offering）、选课记录（Enrollment）、讲授关系（Teacher Course）和绩点对照规则（Grade Point Rule）。各实体之间的关系如下：')
        add_n(i+1, [
            '• 班级与学生 —— 1:N 关系。一个班级包含多名学生，一名学生只属于一个班级。学生表中通过 class_id 外键关联班级。',
            '• 课程与选课安排 —— 1:N 关系。一门课程可以有多个排课（不同学期、不同教师），一个排课只对应一门课程。',
            '• 教师与选课安排 —— 1:N 关系。一位教师可以有多个排课，一个排课由一位教师授课。',
            '• 选课安排与选课记录 —— 1:N 关系。一个排课可被多名学生选修，一条选课记录对应一个排课。',
            '• 学生与选课记录 —— 1:N 关系。一名学生可以有多条选课记录，一条选课记录属于一名学生。',
            '• 教师与课程（讲授关系） —— M:N 关系。通过 teacher_course 中间表实现，UNIQUE(teacher_id, course_id) 约束保证无重复记录。',
        ])
        add_p(i+7, '全局 E-R 图如图 3-1 所示。')
        add_img(i+8, 'er_diagram.png', 5.5, '图 3-1  学生成绩管理系统全局 E-R 图')

    # 3.2
    i = find_h('3.2 AI辅助概念设计记录')
    if i:
        add_p(i, '概念设计阶段主要依靠对业务需求的理解进行实体和关系的识别。AI 工具在初期提供了实体划分的参考建议，但最终的设计决策由人工根据数据库设计原理和实际业务需求做出。')

    i = find_h('3.2.1')
    if i:
        add_p(i, '在 ER 图绘制过程中，参考了 AI 提供的实体清单和实体间关系的初步建议。经分析发现，AI 建议中将"选课安排"作为一个孤立实体，未能体现其在学生与课程之间的桥梁作用。人工将其调整为连接 enrollment（选课记录）、course（课程）、teacher（教师）和 student（学生）的核心实体，形成了完整的选课业务链路。最终的 ER 图由人工使用 matplotlib 绘制完成。')

    i = find_h('3.2.2')
    if i:
        add_n(i, [
            '在概念设计阶段，对 AI 建议进行了以下人工调整：',
            '（1）学生与选课安排的关系 —— AI 初始描述为 N:1，经分析实际场景中一个学生可选多门课、一门课可被多名学生选，两者之间实为 M:N 关系，通过 enrollment 中间表来实现。人工纠正了这一错误。',
            '（2）教师与课程的关系 —— AI 建议直接在教师表中添加"可授课程"字段。但教师与课程是典型的 M:N 关系，在教师表中嵌入课程字段违反了第一范式。人工增加了独立的 teacher_course 中间表。',
            '（3）绩点对照表 —— AI 的初始设计中未包含此表。为支持 GPA 的自动计算，人工补充了 grade_point_rule 表，定义了 5 个成绩区间与绩点的映射规则。',
            '（4）学生与班级的关系 —— AI 给出了两种方案（中间表 vs 外键）。经过对比分析，学生与班级为 1:N 关系，在学生表中添加 class_id 外键的方案更加简洁高效，无需额外的中间表。',
        ])

def ch4():
    print('第四章...')
    # 4.1
    i = find_h('4.1 基本数据表')
    if i:
        add_p(i, '根据概念结构设计的结果，将 E-R 图转换为以下关系模式。所有表均遵循统一设计规范：字符集 utf8mb4，排序规则 utf8mb4_unicode_ci，存储引擎 InnoDB。全部采用逻辑删除（is_deleted 字段），status 字段独立管理业务状态。各表结构汇总如表 4-1 所示。')
        add_tbl(i+1,
            ['表名', '主要字段', '主键 / 外键', '索引', '说明'],
            [['class\n(班级表)', 'id, name, grade, major,\nstatus, create_time,\nupdate_time, is_deleted', 'PK: id', 'name, grade, major', '存储班级基本信息，\nstatus 区分在读/毕业'],
             ['student\n(学生表)', 'id, name, no, class_id,\nweighted_score, gpa,\nstatus, create_time,\nupdate_time, is_deleted', 'PK: id\nFK: class_id→class.id\nUNIQUE: no', 'name, class_id', 'weighted_score 和 gpa\n由触发器自动计算维护'],
             ['course\n(课程表)', 'id, name, credit,\nstatus, create_time,\nupdate_time, is_deleted', 'PK: id\nUNIQUE: name', '—', '存储课程基本信息'],
             ['teacher\n(教师表)', 'id, name, no, title,\nphone, status,\ncreate_time, update_time,\nis_deleted', 'PK: id\nUNIQUE: no', 'name', 'title 和 phone 可为空'],
             ['course_offering\n(选课安排表)', 'id, course_id, teacher_id,\nsemester, max_students,\ncurrent_students,\nenroll_start_time,\nenroll_end_time,\ngrade_deadline, status', 'PK: id\nFK: course_id→course.id\nFK: teacher_id→teacher.id', 'course_id, teacher_id,\nsemester,\n(enroll_start_time,\nenroll_end_time)', '核心业务表，\ncurrent_students\n由触发器自动维护'],
             ['enrollment\n(选课表)', 'id, offering_id,\nstudent_id, score,\ncreate_time, update_time,\nis_deleted', 'PK: id\nFK: offering_id→\n  course_offering.id\nFK: student_id→\n  student.id\nUNIQUE: (offering_id,\n  student_id)', 'offering_id,\nstudent_id', 'score=NULL 表示未录入；\nis_deleted=1 表示退选'],
             ['teacher_course\n(讲授表)', 'id, teacher_id,\ncourse_id, create_time,\nupdate_time, is_deleted', 'PK: id\nFK: teacher_id→\n  teacher.id\nFK: course_id→\n  course.id\nUNIQUE: (teacher_id,\n  course_id)', 'teacher_id,\ncourse_id', 'M:N 中间表，记录\n教师讲授资格'],
             ['grade_point_rule\n(绩点对照表)', 'id, min_score,\nmax_score, point', 'PK: id', '—', '5 条初始规则：\n0~59→0, 60~69→1,\n70~79→2, 80~89→3,\n90~100→4']],
            [1.05, 1.85, 1.35, 0.95, 1.45], caption='表 4-1  数据库表结构汇总')
        add_p(i+2, '各表之间的关系如图 4-1 所示。')
        add_img(i+3, 'table_relationships.png', 5.5, '图 4-1  数据库关系图')

    # 4.2
    i = find_h('4.2 规范化处理说明')
    if i:
        add_n(i, [
            '本数据库设计整体遵循第三范式（3NF），确保每个非主属性都直接依赖于主键，消除传递依赖和不必要的数据冗余。具体分析如下：',
            '（1）所有表均满足 1NF：每个字段都是不可再分的原子值，不存在重复组或多值字段。',
            '（2）所有表均满足 2NF：每张表都有单一主键（id），所有非主属性完全函数依赖于主键。对于中间表（如 enrollment），业务主键 (offering_id, student_id) 通过 UNIQUE 约束保证唯一性。',
            '（3）所有表均满足 3NF：不存在传递依赖。例如，学生表中只存储 class_id（外键），班级的名称、年级、专业等信息通过 JOIN 班级表获取，而非冗余存储在学生表中。',
            '（4）适度反规范化——course_offering 表的 current_students 字段：该字段可以通过 COUNT(enrollment) 实时计算得出，但考虑到高并发选课场景下该值查询频繁，通过触发器自动维护冗余值可以避免每次查询都需要 COUNT 聚合，显著提升读取性能。这是以空间换时间的典型反规范化设计。',
            '（5）适度反规范化——student 表的 weighted_score 和 gpa 字段：这两个汇总字段可以通过 JOIN enrollment、course_offering、course 和 grade_point_rule 四张表计算得出。但成绩查询是高频操作，通过触发器在成绩变化时自动更新这两个字段，避免了每次查询都进行复杂的多表聚合计算。',
        ])

    # 4.3
    i = find_h('4.3')
    if i:
        add_p(i, '逻辑设计阶段，部分字段的定义参考了 AI 的初始建议，但经过了充分的人工审查和修改。表 4-2 记录了 AI 建议与最终设计之间的关键差异及修改理由。')
        add_tbl(i+1,
            ['表 / 字段', 'AI 原始建议', '最终设计', '修改理由'],
            [['course_offering.\neffective_time', '新增生效时间字段', '删除', '与 enroll_start_time\n功能重复，冗余字段'],
             ['enrollment.\nenroll_time', '单独的选课时间字段', '删除\n使用 create_time', '语义重复，create_time\n已能表示选课时间'],
             ['enrollment.\nscore', '建议成绩单独建表', '保留在\nenrollment 表中', '一条选课记录只有\n一个成绩，无需额外\n建表增加 JOIN 开销'],
             ['student.\nweighted_score', '建议用视图\n实时计算', '冗余存储 +\n触发器维护', '成绩查询是高频操作，\n预计算比每次聚合\n更高效'],
             ['teacher_course', '建议放入\n教师表字段', '独立中间表', 'M:N 关系必须用\n中间表实现，不能\n放在任何一端表中']],
            [1.4, 1.6, 1.5, 2.1], caption='表 4-2  AI 辅助逻辑设计审核对照表')

    # 4.4
    i = find_h('4.4 其它完整性约束说明')
    if i:
        add_n(i, [
            '数据库通过以下多层约束机制确保数据完整性：',
            '（1）实体完整性：所有表以 id 为 PRIMARY KEY，AUTO_INCREMENT 自动生成唯一标识。',
            '（2）参照完整性：通过 FOREIGN KEY 约束维护了 7 对表间引用关系，所有外键约束均依赖 InnoDB 的行级锁和事务机制保证操作的一致性。包括：student.class_id→class.id、course_offering.course_id→course.id、course_offering.teacher_id→teacher.id、enrollment.offering_id→course_offering.id、enrollment.student_id→student.id、teacher_course.teacher_id→teacher.id、teacher_course.course_id→course.id。',
            '（3）用户定义完整性：通过以下多种机制协同实现——UNIQUE 约束（student.no、teacher.no、course.name、enrollment(offering_id, student_id)、teacher_course(teacher_id, course_id) 均设有唯一索引，既保证数据唯一性又加速查询）；NOT NULL 约束（所有业务关键字段均设为 NOT NULL）；DEFAULT 值（status 默认 1 表示有效，current_students 默认 0，create_time 和 update_time 由 MySQL 自动填充）；触发器校验（选课前检查名额是否已满、选课后和退选后自动更新 current_students、成绩变化时自动重算学籍分和绩点分）；存储过程中的业务规则校验（选课时检查选课有效期、重复选课、同课不同教师；成绩录入时检查教师归属、截止时间、成绩范围 0~100）。',
            '（4）逻辑删除策略：所有删除操作不进行物理删除，而是通过 UPDATE is_deleted=1 实现逻辑删除。这样可以保留历史数据的完整性和可追溯性，同时通过视图层的 is_deleted=0 过滤确保应用层不感知已删除数据。',
        ])

def ch5():
    print('第五章...')
    i = find_h('5.1 存储结构与索引设计')
    if i:
        add_n(i, [
            '数据库采用 MySQL 8.0 作为 DBMS，存储引擎选择 InnoDB。选择 InnoDB 的主要理由包括：支持事务（ACID），保证选课、退选、成绩录入等操作的原子性和一致性；支持行级锁和外键约束，适合高并发选课场景；支持崩溃恢复，保障数据安全。',
            '索引设计策略如下：',
            '（1）主键索引：所有表的 id 字段自动建立聚簇索引（clustered index），数据按主键顺序物理存储，主键查询效率最高。',
            '（2）唯一索引：student.no、teacher.no、course.name 建立唯一索引，既保证数据唯一性，又加速按学号/工号/课程名的等值查询。enrollment(offering_id, student_id) 和 teacher_course(teacher_id, course_id) 的复合唯一索引同时支持这两组字段的联合查询。',
            '（3）外键索引：所有外键字段（class_id、course_id、teacher_id、offering_id、student_id）均建立辅助索引，加速 JOIN 查询。例如查询某班级所有学生时，idx_student_class 索引可使查询从全表扫描优化为索引查找。',
            '（4）业务查询索引：course_offering 表建立 semester 索引，支持按学期快速筛选排课；建立 (enroll_start_time, enroll_end_time) 复合索引，支持选课时间段的范围查询（WHERE NOW() BETWEEN enroll_start_time AND enroll_end_time）。',
            '（5）索引权衡：并非所有字段都需要建索引。status 字段仅有 0 和 1 两种值，选择性极低，不适合单独建索引；name 字段虽然建了索引，但考虑到姓名重复率较高，实际查询通常和其他条件组合使用，单列索引的收益有限。',
        ])

    i = find_h('5.2 AI辅助索引设计记录')
    if i:
        add_p(i, '物理设计阶段，主要向 AI 咨询了索引设计的通用原则和 EXPLAIN 执行计划的解读方法，具体的索引方案由人工根据系统实际查询模式确定。')

    i = find_h('5.2.1')
    if i:
        add_p(i, '向 AI 提出的问题大致为："对于学生成绩管理系统，在哪些字段上建索引比较合适？如何通过 EXPLAIN 来分析查询的执行效率？"AI 回复了索引设计的几个通用原则：为 WHERE 子句、JOIN 条件、ORDER BY 和 GROUP BY 中出现的字段建索引；为外键字段建索引以加速关联查询；同时提醒不要过度索引，因为索引会降低写入性能并占用额外存储空间。此外，AI 解释了 EXPLAIN 输出中 type、key、rows、Extra 等关键字段的含义，帮助理解查询的执行方式。')

    i = find_h('5.2.2')
    if i:
        add_p(i, '使用 EXPLAIN 命令对系统中的核心查询进行了执行计划分析。以"按班级和课程统计成绩"这一典型查询为例，EXPLAIN 的输出如图 5-1 所示。从执行计划来看，查询使用了 idx_enr_student 索引和 PRIMARY 主键索引，访问类型 type 分别为 ref 和 eq_ref（属于较好的访问类型），预估扫描行数 rows 较小，Extra 列未见 Using filesort 或 Using temporary 等性能警告，说明当前索引设计能够有效支撑该查询。')
        add_img(i+1, 'explain_plan.png', 5.0, '图 5-1  EXPLAIN 执行计划分析（班级成绩统计查询）')

    i = find_h('5.2.3')
    if i:
        add_n(i, [
            '最终的索引方案由人工综合系统查询模式确定。AI 的大部分建议被采纳，但以下两条未被采用：',
            '（1）AI 建议为所有表的 status 字段单独建立索引。考虑到 status 只有 0/1 两种取值，选择性极低（约 50%），B+Tree 索引对其几乎没有过滤效果，优化器更倾向于全表扫描。除非 status 出现在复合索引的前缀列中，否则单独建索引没有实际收益。',
            '（2）AI 建议为 student.name 字段建立全文索引（FULLTEXT）。但系统中的姓名查询均为精确匹配或前缀模糊匹配（如 WHERE name LIKE \'张%\'），普通 B+Tree 索引已能通过索引条件下推（ICP）优化这类查询。全文索引维护开销大、占用空间多，在该场景下得不偿失。',
        ])

def ch6():
    print('第六章...')
    i = find_h('6.1 数据库及基本表创建')
    if i:
        add_p(i, '数据库实施按照以下步骤依次完成。第一步，创建数据库 db_sms，字符集设为 utf8mb4、排序规则设为 utf8mb4_unicode_ci，确保能够正确存储和比较中文字符。第二步，依次执行 DDL 脚本，创建 5 张实体表（class、student、course、teacher、course_offering）和 3 张关系表（enrollment、teacher_course、grade_point_rule），完整定义了所有字段、主键、外键、索引、默认值和注释。第三步，创建 3 个视图（v_student_message、v_course_plan、v_enrollment），封装多表 JOIN 查询和状态计算逻辑。第四步，创建 5 个触发器，实现选课名额检查、人数自动更新和成绩自动计算。第五步，创建 9 个存储过程和 2 个存储函数，封装核心业务逻辑和辅助查询功能。第六步，通过 Python 脚本将 CSV 格式的测试数据批量导入数据库（12 个班级、120 名学生、16 名教师、12 门课程、100 条排课记录、835 条选课记录），为系统测试提供数据基础。')

    i = find_h('6.1.1')
    if i:
        add_p(i, '以下为核心建表 DDL 语句示例，以班级表和学生表为代表。完整的 DDL 脚本存放于项目 sql/ 目录下。')
        add_code(i+1, '-- 班级表\nCREATE TABLE class (\n    id          INT             NOT NULL AUTO_INCREMENT  COMMENT \'班级ID\',\n    name        VARCHAR(50)     NOT NULL                 COMMENT \'班级名\',\n    grade       VARCHAR(10)     NOT NULL                 COMMENT \'年级\',\n    major       VARCHAR(100)    NOT NULL                 COMMENT \'专业\',\n    status      TINYINT(1)      NOT NULL DEFAULT 1       COMMENT \'1=在读 0=毕业\',\n    create_time DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP,\n    update_time DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n    is_deleted  TINYINT(1)      NOT NULL DEFAULT 0       COMMENT \'逻辑删除\',\n    PRIMARY KEY (id),\n    INDEX idx_class_name (name),\n    INDEX idx_class_grade (grade),\n    INDEX idx_class_major (major)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;')
        add_code(i+2, '-- 学生表\nCREATE TABLE student (\n    id              INT             NOT NULL AUTO_INCREMENT  COMMENT \'学生ID\',\n    name            VARCHAR(50)     NOT NULL                 COMMENT \'姓名\',\n    no              VARCHAR(20)     NOT NULL                 COMMENT \'学号\',\n    class_id        INT             NOT NULL                 COMMENT \'班级ID\',\n    weighted_score  DECIMAL(5,2)    NOT NULL DEFAULT 0.00    COMMENT \'学籍分\',\n    gpa             DECIMAL(5,2)    NOT NULL DEFAULT 0.00    COMMENT \'绩点分\',\n    status          TINYINT(1)      NOT NULL DEFAULT 1       COMMENT \'1=在读 0=离校\',\n    create_time     DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP,\n    update_time     DATETIME        NOT NULL DEFAULT CURRENT_TIMESTAMP ON UPDATE CURRENT_TIMESTAMP,\n    is_deleted      TINYINT(1)      NOT NULL DEFAULT 0       COMMENT \'逻辑删除\',\n    PRIMARY KEY (id),\n    UNIQUE INDEX uk_student_no (no),\n    INDEX idx_student_name (name),\n    INDEX idx_student_class (class_id),\n    CONSTRAINT fk_student_class FOREIGN KEY (class_id) REFERENCES class(id)\n) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4 COLLATE=utf8mb4_unicode_ci;')

    i = find_h('6.1.2')
    if i:
        add_p(i, '测试数据通过 Python 脚本配合 pymysql 的 executemany() 方法进行批量导入。数据源为预先准备的 CSV 文件，涵盖 12 个班级、120 名学生、16 名教师、12 门课程以及它们之间的关联关系数据。绩点对照表（grade_point_rule）的 5 条规则数据则直接在 SQL 脚本中以 INSERT 语句初始化的方式写入，作为系统的基础配置数据，不需要通过 CSV 导入。')

    i = find_h('6.2 视图、存储过程、触发器、函数创建')
    if i:
        add_n(i, [
            '除基本表外，数据库还实现了以下程序对象，共同构成完整的数据访问和业务逻辑层：',
            '（1）视图（3 个）。v_student_message：封装学生基本信息查询，对外隐藏学籍分和绩点分字段，供学生端和一般查询使用。v_course_plan：封装选课安排列表查询，将课程名、教师名、剩余名额、选课状态（未开始/选课中/已截止/已结束）等通过 JOIN 和 CASE 表达式一次查出。v_enrollment：封装选课详情查询，通过四表联查（学生+选课安排+课程+教师）输出完整的选课信息。所有视图均过滤 is_deleted=0 的记录，应用层无需重复编写软删除过滤条件。',
            '（2）存储过程（9 个）。sp_show_courses：查询可选课程，过滤出选课中、有名额、且当前学生未选过的课程。sp_enroll：执行选课操作，包含 4 层校验（学生存在性→选课有效期→是否重复选同一排课→是否已选同门课的其他教师），使用事务保证原子性。sp_unenroll：执行退选操作，校验已选课且在退选截止时间前，逻辑删除选课记录。sp_grade_input：录入成绩，依次校验教师身份→排课归属→成绩录入截止时间→学生选课状态→成绩范围（0~100）。sp_student_roster：按班级输出学生名单及学籍分。sp_class_grade_report：按班级和课程统计平均分、最高分、最低分、及格率和参评人数。sp_student_semester_avg：按学号和学期统计该学生的平均成绩。sp_teacher_info 和 sp_teacher_list：分别查询单个教师和全部教师的授课统计信息（排课数、选课学生数、已录入成绩数）。',
            '（3）触发器（5 个）。trg_enrollment_before_insert：选课前检查选课安排的名额是否已满，名额已满则通过 SIGNAL 抛出错误阻止插入。trg_enrollment_after_insert：选课成功后自动将对应 course_offering 的 current_students 加 1。trg_enrollment_after_update：当选课记录的 is_deleted 从 0 变为 1（即退选）时，自动将 current_students 减 1。trg_enrollment_after_insert_score 和 trg_enrollment_after_update_score：当选课记录中 score 字段被插入或更新时，自动重新计算该学生的学籍分（weighted_score = Σ(成绩×学分)/Σ学分）和绩点分（gpa = Σ(学分×绩点)/Σ学分），其中绩点通过 grade_point_rule 表的区间匹配获得。',
            '（4）存储函数（2 个）。fn_get_student_id(p_student_no)：根据学号返回对应的学生 ID，供各存储过程调用。fn_get_teacher_id(p_teacher_no)：根据工号返回对应的教师 ID，供各存储过程调用。两个函数均声明为 DETERMINISTIC 和 READS SQL DATA，以便优化器进行查询优化。',
        ])

    i = find_h('6.3 AI辅助代码生成与人工审核记录')
    if i:
        add_p(i, '数据库实施阶段，AI 主要参与了部分 SQL 脚本初始框架的生成工作。以下记录使用情况、原始输出与人工修正过程。')

    i = find_h('6.3.1')
    if i:
        add_p(i, '在编写存储过程时，向 AI 提出了类似以下的要求："请帮我写一个 MySQL 存储过程实现学生选课的功能。要求输入学号和排课 ID，需要依次检查：①学生是否存在、②该排课是否在选课期内、③学生是否已经选过这个排课、④学生是否已经选了同一门课的其他教师。使用事务来保证操作的原子性，遇到任何错误都要回滚。"')

    i = find_h('6.3.2')
    if i:
        add_p(i, 'AI 给出的代码包含了存储过程的基本结构——参数定义、变量声明、事务的 START TRANSACTION 和 COMMIT/ROLLBACK 框架，以及四个校验步骤的 SQL 查询骨架。整体结构是合理的，为后续的完善提供了出发点。但代码中存在多处需要修正的问题，详见下一节的记录。')

    i = find_h('6.3.3')
    if i:
        add_n(i, [
            '在人工审核和调试过程中，发现并修复了以下主要问题：',
            '（1）数据库名称错误：AI 生成的 CREATE DATABASE 语句中，数据库名被写成了上一项目的名称（db_purchase_sales_inventory），说明 AI 在一定程度上受到了前文上下文的影响。人工将数据库名修正为 db_sms（Student Management System）。',
            '（2）字段冗余：AI 在 course_offering 表中多加了一个 effective_time 字段，而该表已有 enroll_start_time 来表示选课开始时间，两者语义重叠。enrollment 表中同时存在 enroll_time 和 create_time 两个时间字段，但一条选课记录的创建时间就是选课时间，保留 create_time 即可。人工逐一分析后删除了这些冗余字段。',
            '（3）字符集冲突：AI 生成的存储过程中，字符串比较操作没有指定 COLLATE 子句。当数据库的默认排序规则与表定义不一致时，会导致 "Illegal mix of collations" 错误。人工在所有字符串比较处添加了 COLLATE utf8mb4_unicode_ci 子句以统一排序规则。',
            '（4）NULL 值比较错误：触发器判断成绩是否发生变化时，AI 使用了普通的等号（=）进行比较。但 score 字段允许 NULL，而 SQL 标准中 NULL = NULL 的返回值是 NULL（而非 TRUE），这意味着即使成绩实际上没有变化，条件判断也可能得不到预期的结果。人工将比较运算符改为 MySQL 提供的 NULL 安全等于运算符（<=>），确保 NULL 值的比较行为正确。',
            '（5）sed 批量替换的副作用：在项目早期使用 sed 命令对代码进行批量文本替换时，部分存储过程名称中包含的替换目标关键字被意外修改，导致存储过程无法正常调用。人工逐条检查并手动恢复了被误改的名称。',
        ])

def ch7():
    print('第七章...')
    i = find_h('第七章 系统实现')
    if i:
        add_p(i, '在完成数据库设计与实施的基础上，使用 Python 语言开发了两种用户界面：命令行交互界面（CLI）和基于 Streamlit 框架的 Web 可视化界面。系统支持学生、教师、教务管理员三种角色的登录和差异化功能。')

        add_h2(i+1, '一、系统架构')
        add_p(i+2, '系统采用经典的分层架构设计，自上而下分为四个层次。表示层负责用户交互，包括 CLI 终端界面和 Streamlit Web 页面两套实现。业务逻辑层实现选课、退选、成绩录入、统计查询等核心功能的处理逻辑和权限控制。数据访问层通过 PyMySQL 驱动与 MySQL 数据库建立连接，封装存储过程的调用和游标管理。数据存储层为 MySQL 8.0 数据库，包含完整的表、视图、触发器、存储过程和函数。各层之间通过清晰的接口进行通信，上层不直接访问下层内部实现。')
        add_img(i+3, 'system_architecture.png', 5.0, '图 7-1  系统架构图')

        add_h2(i+4, '二、用户界面')
        add_n(i+5, [
            '（1）命令行界面（CLI）：通过 Python 的 input()/print() 实现菜单驱动的交互方式。CLI 实现了 CJK 字符宽度感知的表格对齐算法，解决了中英文混排时的列对齐问题。对于大数据量列表，提供了 Paginator 分页组件，支持翻页浏览、指定页跳转和返回上级菜单。此外还提供了清屏（跨平台兼容 Windows/Linux）、任意键暂停、操作确认提示等用户友好功能。',
            '（2）Web 界面（Streamlit）：使用 Streamlit 框架构建响应式 Web 应用。采用 st.navigation 实现原生多页导航，支持页面分组折叠和 emoji 图标。学生端包含可选课程浏览、选课、退选、成绩查询、学期均分查询共 5 个功能页面。教师端包含成绩录入、CSV 批量导入、课程学生查看共 3 个功能页面。教务端分为四大功能模块——数据查看（数据概览仪表板 + 5 个统计查询页面）、数据管理（6 个实体的 CRUD 管理页面）、系统工具（备份与恢复）。三种角色通过统一的登录页面进行身份认证，系统根据用户标识自动识别角色并展示对应的功能菜单。',
        ])
        j = find_t('（2）Web 界面（Streamlit）')
        if j:
            add_img(j, 'ui_login.png', 4.2, '图 7-2  系统登录页面')
            add_img(j+1, 'ui_admin_dashboard.png', 5.0, '图 7-3  教务端数据概览页面')
            add_img(j+2, 'ui_student_enroll.png', 5.0, '图 7-4  学生选课页面')

        add_h2(i+8, '三、关键实现技术')
        add_n(i+9, [
            '（1）权限控制：系统实现了基于角色的访问控制（RBAC）。登录时根据用户输入的标识（学号/工号/"admin"）分别在 student 表和 teacher 表中查询匹配记录，确定用户的角色和身份信息后存入 st.session_state，后续所有页面根据角色动态渲染对应的导航菜单和功能组件。',
            '（2）消息通知机制：Streamlit 原生的 st.success()/st.error() 和 st.toast() 在页面 rerun 时状态会丢失。系统采用在 st.session_state 中维护一个 msg 元组（消息类型, 消息文本）的方式，在每次页面渲染时检查并弹出消息，实现了消息的跨 rerun 持久化。',
            '（3）数据备份与恢复：备份通过 Python 的 subprocess 模块调用 mysqldump 命令行工具实现，备份文件以时间戳命名，内容包含完整的表结构、数据、存储过程和触发器。恢复通过 st.file_uploader 上传备份文件，将文件内容保存在 session_state 中以避免按钮点击 rerun 时丢失上传状态，然后通过 mysql 命令行工具执行恢复。恢复前自动添加 SET FOREIGN_KEY_CHECKS=0 以避免外键约束导致的数据导入顺序问题，恢复后重新开启外键检查。',
            '（4）CSV 批量导入成绩：教师端支持上传 CSV 文件批量导入成绩。系统解析 CSV 文件内容（兼容 UTF-8 BOM），逐行提取 plan_id、student_no 和 score 三个字段，依次调用 sp_grade_input 存储过程进行校验和录入，最后汇总统计成功和失败的数量，并向用户提供逐条的详细错误反馈。',
        ])

def ch8():
    print('第八章...')
    i = find_h('8.1 测试方案与测试用例')
    if i:
        add_p(i, '系统测试采用黑盒测试方法，从用户视角出发，覆盖三种角色的核心功能流程和各类异常场景。测试按"先正常流程后异常场景"的顺序进行，测试流程如图 8-1 所示。')
        add_img(i+1, 'test_flow.png', 5.0, '图 8-1  系统测试流程图')
        add_p(i+2, '各功能模块的主要测试用例及执行结果如表 8-1 所示。')
        add_tbl(i+3,
            ['测试模块', '测试用例', '预期结果', '实际'],
            [['登录认证', '学生学号登录', '进入学生端', '通过'],
             ['登录认证', '教师工号登录', '进入教师端', '通过'],
             ['登录认证', 'admin 登录', '进入教务端', '通过'],
             ['登录认证', '不存在的用户', '提示"用户不存在"', '通过'],
             ['学生选课', '正常选课', '选课成功，名额-1', '通过'],
             ['学生选课', '选已满的课程', '提示"选课已满"', '通过'],
             ['学生选课', '重复选同一排课', '提示"已选过这门课"', '通过'],
             ['学生选课', '选同门课不同教师', '提示"已选过该课程的其他教师"', '通过'],
             ['学生退选', '正常退选', '退选成功，名额+1', '通过'],
             ['学生退选', '退未选的课', '提示"未选这门课"', '通过'],
             ['学生退选', '选课截止后退选', '提示"已过退选截止时间"', '通过'],
             ['成绩录入', '正常录入成绩', '录入成功，学籍分/GPA 更新', '通过'],
             ['成绩录入', '录入非本人课程', '提示"这不是你的课"', '通过'],
             ['成绩录入', '成绩超出 0~100', '提示"不在 0~100 之间"', '通过'],
             ['成绩录入', '超过录入截止时间', '提示"已过成绩录入截止时间"', '通过'],
             ['教务管理', '班级/课程/教师/学生 CRUD', '增删改查正确执行', '通过'],
             ['成绩统计', '班级成绩统计', '数据与原始数据一致', '通过'],
             ['备份恢复', '备份后删除数据再恢复', '恢复后数据完整一致', '通过']],
            [1.0, 2.2, 2.2, 0.6], caption='表 8-1  系统测试用例及结果')

    j = find_t('8.2 测试结果与分析')
    if j:
        add_n(j, [
            '全部 18 个测试用例均通过验证，通过率 100%，系统功能完整、运行稳定。',
            '测试过程中发现并修复了以下主要 Bug：',
            '（1）成绩录入时未校验教师是否拥有该排课的授课权限——在 teacher.py 的 grade_input 函数中增加了排课归属校验，通过查询 course_offering 表确认当前教师是否是该排课的授课教师。',
            '（2）成绩更新后学籍分和绩点分未自动重新计算——在 enrollment 表上增加了两个 AFTER 触发器，分别在成绩插入和成绩更新时自动按学分加权算法重算学籍分和绩点分。',
            '（3）Streamlit 的 st.file_uploader 在按钮点击触发页面 rerun 后丢失上传文件状态——将上传的文件内容保存在 st.session_state 中，页面重渲染时从 session_state 读取，不再依赖 file_uploader 的临时状态。',
            '（4）数据恢复时出现中文乱码——原因是 Windows 环境下 subprocess 默认使用 GBK 编码传递 stdin。通过在 subprocess.run 中显式指定 encoding=\'utf-8\'，以及在读取上传文件时使用 .decode(\'utf-8\') 解决了编码不匹配的问题。',
            '测试结论：系统满足设计目标和功能需求，可以投入使用。',
        ])

    i = find_h('8.3 AI辅助测试')
    if i:
        add_n(i, [
            '测试阶段 AI 的参与主要限于以下辅助性工作：',
            '（1）在分析边界值时，参考了 AI 关于"选课名额刚好满员时并发插入的行为"以及"成绩取 0 和 100 这两个边界值时的处理逻辑"的分析建议。',
            '（2）在调试存储过程报错和 collation 冲突等数据库层面问题时，向 AI 咨询了错误信息的含义和排查方向。',
            '（3）在准备 CSV 批量测试模板时，利用 AI 生成了包含正确列名和示例数据的模板文件。',
            '测试用例的设计、测试的执行、结果的核验和 Bug 的修复均由人工独立完成。',
        ])

def ch9():
    print('第九章...')
    i = find_h('第九章 总结')
    if i:
        add_n(i, [
            '（1）课程设计完成的主要工作和成果：',
            '本课程设计完成了学生成绩管理系统的完整数据库设计与应用开发。在数据库层面，从需求分析出发，依次完成了概念设计（绘制全局 E-R 图）、逻辑设计（8 张表的关系模式，满足第三范式）、物理设计（确定 InnoDB 为存储引擎并完成索引优化）和数据库实施（编写并执行了完整的 DDL 脚本，创建了 3 个视图、5 个触发器、9 个存储过程和 2 个函数）。在应用层面，使用 Python 语言开发了命令行（CLI）和 Web（Streamlit）两套用户界面，实现了学生选课退选、教师成绩录入与批量导入、教务数据管理、成绩统计分析和数据备份恢复等完整功能。系统已通过全部测试用例的验证，功能完整、运行稳定。',
            '（2）收获和体会：',
            '通过本次课程设计，我系统性地实践了数据库设计的完整流程。概念设计阶段学会了如何从需求描述中识别实体、确定属性和分析实体间关系，并绘制规范的 E-R 图。逻辑设计阶段掌握了 E-R 图向关系模式的转换方法，以及范式的实际应用——特别是在什么情况下应该严格遵守范式、什么情况下可以考虑适度的反规范化来换取性能。物理设计阶段通过 EXPLAIN 工具的实际使用，直观地理解了索引对查询性能的影响，学会了根据查询模式来设计索引而非盲目套用公式。数据库实施阶段是工程量最大的环节——编写几十个 SQL 对象（视图、触发器、存储过程、函数）并逐个调试通过，在这个过程中对 MySQL 的语法细节、事务机制和错误处理有了远超课堂的深入理解。此外，Python 应用开发的部分让我体会到"数据库设计得好，应用代码就简单"这个道理——当大量业务逻辑通过存储过程和触发器下沉到数据库层后，Python 端的代码量大幅减少，逻辑也更加清晰。',
            '（3）存在的不足和改进设想：',
            '系统当前存在以下可以改进的地方：(a) 用户认证仅通过学号/工号识别身份，没有密码验证，安全性不足。后续可以增加密码字段（经过哈希存储），并在登录时进行密码比对。(b) 选课的并发控制目前依赖 InnoDB 的行级锁和触发器的名额检查。在选课高峰时段（如每学期初的选课日），大量学生同时选课可能导致锁竞争。可以考虑引入 Redis 缓存选课名额实现无锁的原子扣减，异步同步到数据库。(c) 当前的 Web 界面是 Streamlit 单机应用，数据访问和界面展示耦合在一起。后续可以考虑将数据访问层独立为 RESTful API 服务，前端使用更灵活的框架（如 Vue 或 React），实现前后端分离。(d) 成绩统计目前只支持按班级+课程这一个维度。后续可以扩展按年级排名、按课程横向对比、学生个人成绩趋势图等多维度统计分析功能。',
            '（4）课程设计的感想：',
            '数据库系统课程设计是一次理论与实践紧密结合的宝贵学习经历。回想最开始的时候，我对数据库设计的理解还停留在"建几张表、写几个查询"的层面。随着项目推进，从需求分析中发现各种隐含的业务规则（如"同一门课不能选两个老师"、"退选只能在截止时间之前"），到概念设计中反复推敲实体和关系的划分，再到逻辑设计中在范式与性能之间做出权衡，我逐渐认识到数据库设计远不只是技术性的表结构定义——它更是一个深入理解业务、梳理数据关系、做出系统性权衡的过程。一个好的数据库设计，会让后续的开发事半功倍；而一个不好的设计，会在后续阶段不断制造麻烦。另外，这次课程设计也让我深刻体会到"细节决定成败"——一个 NULL 值的比较方式、一个 COLLATE 的指定，都可能成为调试数小时的根源。这些经验是课堂教学无法提供的，只有亲手做过才能真正理解。',
        ])

    # 9.1 AI
    i = find_h('9.1 AI辅助效果评估与反思')
    if i:
        add_p(i, '在本次课程设计中，AI 工具在部分环节起到了辅助作用。以下从效率、质量和个人反思三个角度进行总结。')

    i = find_h('9.1.1')
    if i:
        add_n(i, [
            '从开发效率来看，AI 在以下方面确实节省了时间：（1）SQL 脚本的语法框架生成——像 DDL、触发器、存储过程这些代码有相对固定的模板结构，AI 可以快速生成符合规范的框架，比从头敲代码快不少。（2）Streamlit 页面组件的初始搭建——AI 能快速给出 st.dataframe、st.metric 等组件的调用方式，省去了很多查文档的时间。（3）代码审查——AI 可以发现一些比较明显的冗余（比如同样的字典构建了两次、import 了但没使用的模块等），这在代码量较大的时候挺有用。',
            '但也有一些环节，用了 AI 反而花了更多时间：（1）AI 生成的代码经常会包含一些难以一眼看出的逻辑错误。例如在触发器中用普通等号去判断 NULL 值是否变化，这个错误定位和排查花了不少时间——如果是我自己写的话，因为知道 score 字段允许 NULL，一开始就会用 <=> 来做比较。（2）AI 对项目的整体上下文理解有限，有时候它生成的代码和已经写好的部分风格不一致，或者使用了和项目其他部分不兼容的写法，需要回头去统一调整。（3）AI 偶尔会使用一些它"自创"的 API 或者已经废弃的写法，表面看起来没问题但实际运行会报错，这种"看似正确"的错误比明显的错误更费时间去排查。',
        ])

    i = find_h('9.1.2')
    if i:
        add_n(i, [
            '在使用 AI 的过程中，我遇到了几次印象比较深刻的错误或"幻觉"：',
            '（1）数据库名被写错——AI 似乎受到了前文语境的影响，在建库语句中生成了上一项目的数据库名。这说明 AI 并不真正"理解"当前项目的语境，它只是基于统计模式生成文本。',
            '（2）在 SQL 中关于 NULL 的比较——AI 多次使用普通等号来比较可能为 NULL 的值。这是一个比较隐蔽的错误，因为如果测试数据中恰好没有 NULL 值，测试是可以通过的。直到我在测试中刻意构造了成绩为 NULL 的场景，才发现了这个问题。这也提醒我，AI 生成的代码需要经过"刻意构造边界条件"的测试，而不能只跑正常流程。',
            '（3）AI 有时候会表现出一种"过度自信"的态度——它会用很确定的语气给出实际上不完全正确、甚至完全错误的建议。最开始我倾向于相信 AI 的输出，后来逐渐学会了"先怀疑、再验证"的工作方式。',
            '总结下来的几点经验：（1）AI 更适合做"辅助"而不是"替代"——用它来生成框架、查漏补缺、审阅代码是有价值的，但核心的设计决策和关键逻辑的实现，还是自己动手更可靠。（2）AI 的每一行输出都需要经过审查和测试——不能因为"看起来合理"就默认它是正确的。（3）给 AI 的提示越具体、上下文越清晰，输出质量越高——泛泛的问题得到泛泛的答案，这是很自然的规律。',
        ])

    i = find_h('9.1.3')
    if i:
        add_n(i, [
            '回顾整个课程设计，我认为以下环节是必须由人来完成、不应该交给 AI 的：',
            '（1）需求分析和概念设计——识别业务实体、分析实体之间的关系、判断关系的基数（1:1、1:N、M:N），这些需要基于对业务场景的深入理解和对数据库设计原理的掌握，AI 缺乏真实的业务感知能力，容易做出不符合实际的判断。',
            '（2）范式和反范式的权衡——什么时候严格遵守 3NF、什么时候做适度的反规范化来换取性能，这需要对系统的查询模式和数据量有整体把握。AI 可以提供一般性原则，但无法针对具体的系统做出恰当的设计决策。',
            '（3）测试用例的设计——特别是异常场景和边界条件的覆盖，需要设计者对系统的业务规则有全面的理解。AI 生成的测试用例往往偏向"快乐路径"，对异常和边界的覆盖不够全面。',
            '总的来说，这次课程设计的核心工作——数据库设计的主要决策和关键代码的实现——都是由我自己完成的。AI 更多地充当了一个"查资料更快的搜索引擎"和"代码框架生成器"的角色。通过这次实践，我最大的收获之一是学会了如何合理地使用 AI 工具：知道它擅长什么、不擅长什么，什么时候该用它提效，什么时候该关掉它自己沉下心来思考和编码。',
        ])

def refs_appendix():
    print('参考文献与附录...')
    i = find_h('参考文献')
    if i:
        add_n(i, [
            '[1] 王珊, 萨师煊. 数据库系统概论（第5版）[M]. 北京: 高等教育出版社, 2014.',
            '[2] 吴臣. 数据库系统课程设计讲义[Z]. 2026.',
            '[3] MySQL 8.0 Reference Manual [EB/OL]. https://dev.mysql.com/doc/refman/8.0/en/.',
            '[4] Streamlit Documentation [EB/OL]. https://docs.streamlit.io/.',
            '[5] Python 3.12 Documentation [EB/OL]. https://docs.python.org/3.12/.',
        ])

    i = find_h('附录')
    if i:
        add_n(i, [
            '本课程设计的完整代码和 SQL 脚本存放于项目目录中，主要文件结构如下：',
            '• sql/ 目录：01_数据库创建.sql、02_基础数据表.sql、03_中间表.sql、04_视图.sql、06_触发器.sql、07_存储过程.sql、08_存储函数.sql，共 7 个脚本文件，覆盖全部数据库对象的创建。',
            '• src/ 目录：app.py（Streamlit Web 主应用，约 800 行）、main.py（CLI 入口程序）、admin.py（教务管理功能，约 1176 行）、student.py（学生端功能）、teacher.py（教师端功能）、tester.py（测试工具，可切换任意角色进行功能验证）、core/（配置模块、认证模块、工具函数模块）。',
            '• data/ 目录：class.csv、student.csv、teacher.csv、course.csv、course_offering.csv、enrollment.csv、teacher_course.csv，共 7 个 CSV 测试数据文件。',
            '• docs/ 目录：开发日志.md（完整开发时间线）、AI修正日志.md（AI 错误及人工修正记录）、日志.md（问答日志）、以及建表分析、Python 代码说明、事务讲解等设计文档。',
            '• backup/ 目录：多次数据库备份文件，以时间戳命名，可用于恢复到任意历史状态。',
            '',
            '开发过程中 AI 辅助修正的部分典型记录：',
            '• 2026-07-07：AI 在 course_offering 表中多加了 effective_time 字段（与 enroll_start_time 功能重复），人工删除。',
            '• 2026-07-07：AI 在 enrollment 表中同时设计了 enroll_time 和 create_time（语义重复），人工删除 enroll_time。',
            '• 2026-07-07：AI 生成的 CREATE DATABASE 语句中数据库名错写为上一项目的名称，人工修正为 db_sms。',
            '• 2026-07-08：AI 在使用 sed 批量替换时误将存储过程名称中的关键字替换，人工逐条手动修复。',
            '• 2026-07-09：AI 在 Streamlit 代码中多处出现复制粘贴导致的 session key 错误（如教师管理页面使用了班级管理的 key 名称），人工逐一定位并修正。',
            '• 2026-07-09：AI 在数据恢复功能中未指定编码，Windows 环境下 GBK 和 UTF-8 不匹配导致中文乱码，人工在 subprocess.run 和文件读取两处增加了 encoding=\'utf-8\' 参数。',
        ])

# ============================================================
# 执行
# ============================================================
if __name__ == '__main__':
    print('='*50)
    ch1(); ch2(); ch3(); ch4(); ch5(); ch6(); ch7(); ch8(); ch9(); refs_appendix()
    doc.save(OUT)
    import os
    sz = os.path.getsize(OUT)/1024
    imgs = len([r for r in doc.part.rels.values() if 'image' in r.reltype])
    print(f'\n保存: {OUT}')
    print(f'段落{len(doc.paragraphs)} | 表格{len(doc.tables)} | 图片{imgs} | {sz:.0f}KB')
