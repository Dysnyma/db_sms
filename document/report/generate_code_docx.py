"""
生成完整代码附录 .docx（全部 Python + SQL 源码）
用法：python document/report/generate_code_docx.py
"""

import os, sys
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.15
style.paragraph_format.space_after = Pt(2)

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cn_font(run, name):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)

def add_code_block(text, filename=""):
    """添加代码块（等宽字体+背景框）"""
    # 文件标题
    if filename:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"📄 {filename}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
        set_cn_font(run, '黑体')

    # 代码内容
    lines = text.split('\n')
    # 每50行分割一个段落，避免单个段落过大
    chunk_size = 80
    for chunk_start in range(0, len(lines), chunk_size):
        chunk = lines[chunk_start:chunk_start + chunk_size]
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

        # 添加底纹背景
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        shading.set(qn('w:val'), 'clear')
        pPr = p._element.get_or_add_pPr()
        pPr.append(shading)

        code_text = '\n'.join(chunk)
        run = p.add_run(code_text)
        run.font.name = 'Consolas'
        run.font.size = Pt(7.5)
        set_cn_font(run, 'Consolas')

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        set_cn_font(run, '黑体')
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    return h

# ════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('数据库系统课程设计\n代码附录')
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
set_cn_font(run, '黑体')

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('学生成绩管理系统 — 完整源代��')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
set_cn_font(run, '黑体')

for _ in range(8):
    doc.add_paragraph()

info = [('学生姓名：', '蔡坤灿'), ('日    期：', '2026年7月')]
for label, val in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + val)
    run.font.size = Pt(14)
    set_cn_font(run, '宋体')

doc.add_page_break()

# ════════════════════════════════════════════
# 目录
# ════════════════════════════════════════════
add_heading_styled('目  录', level=1)
toc_sections = [
    ('第一部分  SQL 代码', ['01_数据库创建.sql', '02_基础数据表.sql', '03_中间表.sql',
       '04_视图.sql', '06_触发器.sql', '07_存储过程.sql', '08_存储函数.sql']),
    ('第二部分  Python 核心代码', ['src/app.py', 'src/main.py', 'src/admin.py',
       'src/student_tui.py', 'src/teacher_tui.py']),
    ('第三部分  Streamlit 页面', ['src/pages/admin_page.py', 'src/pages/teacher.py',
       'src/pages/student.py']),
    ('第四部分  工具模块', ['src/core/config.py', 'src/core/auth.py',
       'src/core/models.py', 'src/core/utils.py', 'src/core/majors.py']),
    ('第五部分  数据脚本', ['src/import_data.py', 'src/reset_data.py', 'src/tester.py']),
]
for section_title, files in toc_sections:
    p = doc.add_paragraph()
    run = p.add_run(section_title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    for f in files:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f'  {f}')
        run.font.size = Pt(10)

doc.add_page_break()

# ════════════════════════════════════════════
# 辅助函数：读取文件
# ════════════════════════════════════════════
def read_file(path):
    full = os.path.join(BASE, path)
    try:
        with open(full, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        return f'// 文件读取失败: {full}'

# ════════════════════════════════════════════
# 第一部分：SQL 代码
# ════════════════════════════════════════════
add_heading_styled('第一部分  SQL 代码', level=1)
sql_files = [
    'sql/01_数据库创建.sql',
    'sql/02_基础数据表.sql',
    'sql/03_中间表.sql',
    'sql/04_视图.sql',
    'sql/06_触发器.sql',
    'sql/07_存储过程.sql',
    'sql/08_存储函数.sql',
]
for f in sql_files:
    content = read_file(f)
    add_code_block(content, f)

doc.add_page_break()

# ════════════════════════════════════════════
# 第二部分：Python 核心代码
# ════════════════════════════════════════════
add_heading_styled('第二部分  Python 核心代码', level=1)
py_core = [
    ('src/app.py', 'Streamlit 主入口（导航 + 页面路由）'),
    ('src/main.py', 'CLI 入口'),
    ('src/admin.py', '教务功能'),
    ('src/student_tui.py', '学生功能（CLI）'),
    ('src/teacher_tui.py', '教师功能（CLI）'),
]
for f, desc in py_core:
    content = read_file(f)
    add_code_block(content, f'  {desc}')

doc.add_page_break()

# ════════════════════════════════════════════
# 第三部分：Streamlit 页面
# ════════════════════════════════════════════
add_heading_styled('第三部分  Streamlit 页面', level=1)
pages = [
    ('src/pages/admin_page.py', '教务页面'),
    ('src/pages/teacher.py', '教师页面'),
    ('src/pages/student.py', '学生页面'),
]
for f, desc in pages:
    content = read_file(f)
    add_code_block(content, f'  {desc}')

doc.add_page_break()

# ════════════════════════════════════════════
# 第四部分：工具模块
# ════════════════════════════════════════════
add_heading_styled('第四部分  工具模块', level=1)
modules = [
    ('src/core/config.py', '数据库连接配置'),
    ('src/core/auth.py', '登录认证'),
    ('src/core/models.py', 'Pydantic 校验模型'),
    ('src/core/utils.py', '工具函数'),
    ('src/core/majors.py', '专业列表管理'),
]
for f, desc in modules:
    content = read_file(f)
    add_code_block(content, f'  {desc}')

doc.add_page_break()

# ════════════════════════════════════════════
# 第五部分：数据脚本
# ════════════════════════════════════════════
add_heading_styled('第五部分  数据脚本', level=1)
scripts = [
    ('src/import_data.py', 'CSV 数据导入'),
    ('src/reset_data.py', '重建数据库'),
    ('src/tester.py', '测试员跳板'),
]
for f, desc in scripts:
    content = read_file(f)
    add_code_block(content, f'  {desc}')

# ── 保存 ──
output = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                      '代码附录.docx')
doc.save(output)
print(f'OK: {os.path.getsize(output)//1024} KB -> {output}')
