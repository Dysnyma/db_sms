"""
精准修改说明书.docx——只替换文字内容，不动任何格式
用法：python document/report/patch_docx.py
"""
import os
from docx import Document

BASE = os.path.dirname(os.path.abspath(__file__))
src = os.path.join(BASE, '《数据库系统课程设计》说明书.docx')
out = os.path.join(BASE, '《数据库系统课程设计》说明书_更新版.docx')

doc = Document(src)

# ── 替换映射（旧文本 → 新文本） ──
REPLACE = {
    # === 1.1 项目背景 ===
    '随着高校招生规模的不断扩大，传统纸质成绩管理方式效率低下、易出错、难统计。':
    '本课程设计选题为题目二「学生成绩管理系统」。系统通过分析学校日常教学管理中的课程、选课、学生、班级、教师和成绩等核心数据，设计并实现一个集教师授课管理、学生选课管理、成绩管理与统计等功能于一体的小型数据库应用系统。',
    '学生成绩管理系统（Student Grade Management System)，数据库名 db_sms。':
    '',

    # === 1.3 技术栈 ===
    'pymysql': 'PyMySQL + SQLAlchemy连接池',
    '—': '—',
    'Streamlit 1.x': 'Streamlit 1.x + Plotly',
    '8.0': '8.0 (InnoDB + utf8mb4)',

    # === 2.1 业务需求 ===
    '成绩录入后自动实时计算学生的学籍分和绩点分（GPA）\n学籍分 = ∑(成绩×学分) / ∑(学分)\n绩点分 = ∑(学分×绩点) / ∑(学分)':
    '成绩录入后触发器自动实时计算学籍分和绩点分（GPA）\n学籍分 = ∑(成绩×学分) / ∑(学分)\n绩点分 = ∑(学分×绩点) / ∑(学分)\n支持逐条录入和批量CSV录入两种方式，CSV导入按行反馈中文错误提示。',

    # === 数据量更新 ===
    '12条（初始数据）': '24条（初始），329条（大规槰测试）',
    '120条（初始数据）': '1,000条（初始），20,000条（大规模测试）',
    '16名（初始数据）': '25名（初始），300名（大规模测试）',
    '约12条（初始数据）': '24条（初始），329条（大规槰测试）',
    '约120条（初始数据）': '1,000条（初始），20,000条（大规模测试）',
    '约16条（初始数据）': '25名（初始），300名（大规模测试）',
    '约12条（初始数据）': '24条（初始），329条（大规模测试）',

    # === 7.1 项目目录结构 ===
    'CLI/': 'src/',
    'Streamlit 入口': 'Streamlit Web 入口（22个页面导航）',
    'CLI 入口': 'CLI 入口',

    # === 7.2 功能模块（学生端） ===
    '选课（单一 / 批量 CSV 导入）': '选课 / 退选',
    '浏览可选课程、选课、退选、查自己成绩和学籍分、查学期均分':
    '查询可选课程（过滤选课期/满员/同课异师）、选课/退选、查看我的成绩（含Plotly各科柱状图+及格线）、学期均分查询（selectbox下拉）',

    # === 7.2 功能模块（教师端） ===
    '录入成绩（单个 / CSV 批量导入）、查看任课班级学生名单':
    '逐条录入成绩（Pydantic校验）、CSV批量导入（按行反馈中文错误）、查看课程学生（成绩分布饼图+排行柱状图+未录入折叠名单+统计指标）',

    # === 7.2 功能模块（教务端） ===
    '数据概览、班级/学生/教师/课程 CRUD、排课管理、数据库备份恢复':
    '数据概览（含Plotly交互式图表）、6大实体CRUD+专业管理（CSV存储）、排课管理（datetime_input原生组件）、选课管理（按学号+学期查询）、班级成绩统计/明细/教师排行、备份恢复',

    # === 7.3 测试数据 ===
    '12个班': '24个班（基础）/329个班（大规槰）',
    '120名学生': '1,000名学生（基础）/20,000名（大规模）',
    '16名教师': '25名教师（基础）/300名（大规模）',
    '40条排课安排': '54条（基础）/1,004条（大规模）',
    '约200条选课记录': '约3,000条（基础）/2,002,688条（大规模）',

    # === 6.1 DDL脚本 ===
    '06_存储过程.sql → 08_存储函数.sql': '06_触发器.sql → 07_存储过程.sql → 08_存储函数.sql',

    # === 9.1 成果 ===
    '8张表、3个视图、1个触发器、8个存储过程、2个函数':
    '8张表、3个视图、5个触发器、9个存储过程、2个函数',
}

# ── 执行替换 ──
count = 0
for p in doc.paragraphs:
    for old, new in REPLACE.items():
        if old in p.text:
            # 保留原段落的所有 run 格式，只替换文本
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
                    break  # 只改第一个匹配的 run
            else:
                # 如果 old 跨 run 分布，整段替换所有 text
                full = p.text
                for r in p.runs:
                    r.text = ''
                p.runs[0].text = full.replace(old, new)
                count += 1

# ── 表格替换 ──
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for old, new in REPLACE.items():
                    if old in p.text:
                        full = p.text
                        for r in p.runs:
                            r.text = ''
                        p.runs[0].text = full.replace(old, new)
                        count += 1

doc.save(out)
print(f'✅ 替换完成：{count} 处修改')
print(f'📄 已保存到：{out}')
