"""生成 3.5 工作日志.docx —— 含序号、更详细的每日工作记录"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "3.5工作日志.docx")
doc = Document()
BASE = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.35

for sec in doc.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)

def sf(run, name):
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None: rf = OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:eastAsia'), name)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = '黑体'; sf(r, '黑体')
        if level == 1: r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        elif level == 2: r.font.size = Pt(13); r.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

def tag(text):
    p = doc.add_paragraph()
    run = p.add_run(f'  {text}  ')
    run.font.size = Pt(9); run.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    sf(run, '微软雅黑')
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '2E86C1'); shading.set(qn('w:val'), 'clear')
    p._element.get_or_add_pPr().append(shading)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)

def P(text, sz=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text); run.font.size = Pt(sz); sf(run, '宋体')
    return p

def bullet(text, sz=10):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text); run.font.size = Pt(sz); sf(run, '宋体')

def sub(text, sz=10):
    """缩进段落"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text); run.font.size = Pt(sz); sf(run, '宋体')

def TABLE(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            c = t.rows[ri].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)

# ════════════════════════════════════════
# 封面
# ════════════════════════════════════════
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('数据库系统课程设计\n工作日志'); r.font.size = Pt(26); r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76); sf(r, '黑体')
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('学生姓名：蔡坤灿    专业：计算机科学与技术')
r.font.size = Pt(14); sf(r, '宋体')
for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026年7月'); r.font.size = Pt(14); sf(r, '宋体')
doc.add_page_break()

# ════════════════════════════════════════
# 总览
# ════════════════════════════════════════
H('工作总览')
TABLE(['日期', '主题', '核心成果', '代码量'],
    [['07/06', '选题与准备', '选题切换（进销存→成绩管理）、建表设计、日志体系搭建', 'SQL 约 800 行'],
    ['07/07', '数据库+CLI骨架', '8表/3视图/5触发器/9SP/2函数 + Python CLI 应用', 'SQL 约 600 行 + Python 约 500 行'],
    ['07/08', 'CLI功能完整实现', '6大实体CRUD、分页器、教师录成绩、Streamlit起步', 'Python 约 800 行'],
    ['07/09', 'Streamlit双界面', '22个Web页面、管理页三次重构、备份恢复加固', 'Python 约 1200 行'],
    ['07/13', '代码质量修复', '目录规范化、Pylint/Flake8清零、连接池、.env配置', '重构约 30 文件'],
    ['07/14', '功能增强+图表', 'Plotly图表、Pydantic校验、专业管理、自动班级名', 'Python 约 500 行'],
    ['07/15', '大数据量支持', '200万生成器、LOAD DATA优化、性能调优、数据修复', 'Python+SQL 约 300 行'],
    ['07/16', '文档收尾', 'PPT重构、说明书更新、代码附录、源码打包', '文档约 2000 行']])
doc.add_page_break()

# ════════════════════════════════════════
# 7月6日
# ════════════════════════════════════════
H('准备工作日：2026年7月6日 — 选题切换与日志体系搭建')
tag('1.0 选题切换')
P('原始选定题目一"进销存管理系统"并完成了完整的数据库设计（17 张表、8 个视图、7 个触发器、15 个存储过程、4 个函数），实际工作量约 SQL 800 行。经综合评估后决定切换至题目二"学生成绩管理系统"。原进销存代码与文档完整保留至 archive/ 目录作为学习参考。')

tag('1.1 建表 DDL 设计')
P('完成新数据库 db_sms 的初始设计。研究任务书 7.2 节要求后，确定 5 张实体表（class、student、course、teacher、course_offering）和 3 张关系表（enrollment、teacher_course、grade_point_rule）的整体架构。确定核心设计规范：全小写+下划线表名、utf8mb4 字符集、utf8mb4_unicode_ci 排序规则、InnoDB 引擎、INT 自增主键，全部表采用逻辑删除（is_deleted 字段）。')

tag('1.2 环境配置')
P('配置 VS Code 开发环境，关闭 SQL 文件保存时自动格式化功能以避免 DDL 被误格式化。建立三线并行的日志体系——开发日志（记录实际操作）、问答日志（记录与 AI 的交互过程）、AI 修正日志（记录 AI 错误及人工修正），确保课程设计全过程可追溯。')

tag('1.3 视图与数据库对象规划')
P('初步确定视图方案：v_student_message（学生基本信息）、v_course_plan（选课安排列表）、v_enrollment（选课详情）。触发器方案：选课前名额检查、选课后人数更新、退选后人数更新。确定 9 个存储过程的功能边界和参数接口。')

tag('1.4 CLAUDE.md 与记忆体系')
P('创建 CLAUDE.md 项目说明文件和 C:\\Users\\CKC\\.claude\\projects\\memory 记忆体系，记录项目概况、设计规范、关键决策和版本管理要求，确保 AI 在每次对话中保持上下文一致。')

doc.add_page_break()

# ════════════════════════════════════════
# 7月7日
# ════════════════════════════════════════
H('第一天：2026年7月7日 — 数据库设计 + Python CLI 应用骨架')

tag('2.1 概念结构设计')
P('深入讨论并确定了实体关系模型的核心决策。学生与班级之间采用 1:N 关系，直接在 student 表中添加 class_id 外键，不建中间表，避免过度设计。学生与课程之间为 M:N 关系，通过 enrollment 中间表实现，该中间表同时承担成绩存储功能。课程与教师之间为 M:N 关系，采用 teacher_course（讲授资格）和 course_offering（实际排课）两层设计，实现了讲授资格与排课安排的逻辑分离。')

tag('2.2 逻辑结构设计')
P('完成 8 张表的完整 DDL 撰写。class 表（班级）包含 name、grade、major 字段；student 表（学生）包含 name、no（学号）、class_id（外键）、weighted_score（学籍分）、gpa（绩点分）字段；course 表（课程）包含 name、credit（学分）字段；teacher 表（教师）包含 name、no（工号）、title（职称）、phone（电话）字段；course_offering 表（排课）关联课程、教师和学期，包含选课时间窗口和人数限制字段；enrollment 表（选课）通过 offering_id 和 student_id 关联排课和学生，score 字段允许为 NULL（表示未录入）。')

tag('2.3 视图开发')
P('完成 3 个视图的编写。v_student_message 封装学生与班级的基本信息联查，不暴露学籍分和绩点分敏感数据。v_course_plan 作为学生选课时的核心视图，输出课程名、教师名、剩余名额、选课状态（未开始/选课中/已截止/已结束）。v_enrollment 封装选课详情的多表联查。')

tag('2.4 触发器开发')
P('完成 5 个触发器的开发。trg_enrollment_before_insert 在选课插入前检查名额，使用 FOR UPDATE 行级锁防止并发超卖。trg_enrollment_after_insert 在选课后自动更新 course_offering 的 current_students 加 1。trg_enrollment_after_update 检测 is_deleted 从 0 变 1（退选）时减 1。trg_enrollment_after_insert_score 和 trg_enrollment_after_update_score 在成绩录入或更新时，自动按加权平均公式重算学籍分和 GPA。')

tag('2.5 存储过程与函数开发')
P('完成 9 个存储过程和 2 个存储函数的开发。sp_enroll 实现完整的选课事务，包含 4 步校验：学生存在性检查、排课有效期检查、重复选课检查、同一课程不同教师检查。sp_grade_input 实现成绩录入的三重身份校验。sp_student_roster 输出班级学生的完整学籍分名单。sp_class_grade_report 统计指定班级和课程的平均分、最高分、最低分、及格率。sp_student_semester_avg 按学期聚合学生成绩。sp_teacher_info 和 sp_teacher_list 提供教师授课统计报表。fn_get_student_id 和 fn_get_teacher_id 将学号/工号到 ID 的转换逻辑抽取为公共函数，减少存储过程重复代码。')

tag('2.6 AI 辅助与人工修正')
P('AI 生成了大部分 DDL 和存储过程的框架代码，但经过逐行审查发现以下问题：course_offering 表多加了一个 effective_time 字段与 enroll_start_time 语义重复，已删除。MySQL 8.0 默认 utf8mb4_0900_ai_ci 排序规则与数据库 utf8mb4_unicode_ci 冲突导致 1267 错误，通过在 pymysql 连接参数中加 collation=utf8mb4_unicode_ci 并在存储过程和函数中显式加 COLLATE 子句实现双重修复。enrollment 表同时存在 enroll_time 和 create_time 两个重复字段，删除 enroll_time。')

tag('2.7 Python CLI 开发')
P('搭建完整的命令行交互系统。main.py 作为统一入口，auth.py 处理登录认证（支持 admin/教师工号/学生学号三种身份识别）。工具函数层封装了 render_menu（数据驱动菜单渲染，支持自动编号和分列排版）、show_table（中英文混排自动对齐）、Paginator（通用分页器）、confirm（Y/n 确认）、cls（清屏）和 pause（按回车继续）等基础设施。角色模块按学生、教师、教务三类分离，代码结构清晰。')

doc.add_page_break()

# ════════════════════════════════════════
# 7月8日
# ════════════════════════════════════════
H('第二天：2026年7月8日 — CLI 功能完整实现')

tag('3.1 班级管理 CRUD')
P('完成班级管理的完整增删改查。新增时校验班级名称、年级、专业的必填性。修改采用逐字段交互模式，每字段显示当前值，回车则保留原值不修改。删除采用逻辑删除（SET is_deleted=1）。所有操作结果通过终端即时反馈。')

tag('3.2 课程管理 CRUD')
P('完成课程管理的完整增删改查。学分字段支持小数（如 3.0），新增时检查课程名是否已存在（含已逻辑删除的），支持已删除课程的恢复操作。')

tag('3.3 教师管理 CRUD')
P('完成教师管理的完整增删改查。职称（title）和电话（phone）字段设计为可空，新增时若回车跳过则存为 NULL。教师工号格式为 T 开头+数字。')

tag('3.4 学生管理 CRUD')
P('完成学生管理的完整增删改查。新增时列出可选班级列表供选择。学号格式校验为纯数字。同样支持已逻辑删除学生的恢复操作。')

tag('3.5 排课管理 CRUD')
P('完成排课管理的完整增删改查。新增时先选择课程，系统自动从 teacher_course 表中筛选有资格讲授该课程的老师。学期格式限定为 YYYY-YYYY-S。设置选课开始时间、截止时间和成绩录入截止时间三个时间节点。')

tag('3.6 选课管理')
P('教务端实现选课记录查看和强制退选功能。选课列表支持分页显示。')

tag('3.7 分页器（Paginator）推广')
P('将 Paginator 通用分页器组件推广至全部列表展示场景：5 个管理菜单、班级学生名单、教师列表、我的成绩、可选课程列表、成绩录入排课选择、课程学生列表等十余处。每页固定显示 10 条，支持上一页/下一页翻页操作。')

tag('3.8 教师端成绩录入优化')
P('成绩录入流程从直接输入排课 ID 改进为两步分页操作：第一步使用分页器展示该教师的全部排课列表（显示每门课的已选人数和上限），选择排课后进入第二步循环录入模式，输入学号后自动校验该学生是否在选课名单中，录入完成后自动刷新列表显示最新成绩。')

tag('3.9 CSV 批量录入')
P('教师端支持通过 CSV 文件批量导入成绩。文件格式为 plan_id, student_no, score，系统逐行调用 sp_grade_input 存储过程，成功和失败行数分别统计并逐行反馈错误信息。')

tag('3.10 排序规则修复')
P('MySQL 8.0 默认排序规则为 utf8mb4_0900_ai_ci 导致存储过程字符串比较报错 1267。修复方案为在 pymysql 连接参数中指定 collation=utf8mb4_unicode_ci，并在存储过程和函数的字符串比较字段后加 COLLATE utf8mb4_unicode_ci 子句实现双保险。')

tag('3.11 Streamlit 起步')
P('开始 Web 界面开发。app.py 采用 session_state 管理用户登录状态，建立 st.navigation 原生多页导航系统。完成学生端 5 个页面（可选课程、选课、退选、我的成绩、学期均分）和教师端 3 个页面（录入成绩、批量 CSV 上传、查看课程学生）。')

tag('3.12 代码质量优化')
P('tester.py 从 175 行自建子菜单和 lambda 包装精简为 83 行跳板，直接调用真实 menu 函数，消除全部重复代码。修复 import 路径兼容性问题。为 my_grades 和 semester_avg 函数增加 paged 参数，支持 CLI 分页模式和 Streamlit 数据返回模式双调用。')

doc.add_page_break()

# ════════════════════════════════════════
# 7月9日
# ════════════════════════════════════════
H('第三天：2026年7月9日 — Streamlit 双界面搭建')

tag('4.1 教务端 14 页面开发')
P('完成教务管理员的全部 Streamlit 页面。查询统计类 6 个：数据概览（6 个 COUNT 统计）、班级学生名单（选班级→sp_student_roster）、班级成绩统计（双列选班级+课程→5 个 metric）、班级成绩明细（每人每课成绩登记册）、教师信息、教师列表。数据管理类 6 个：班级管理、课程管理、教师管理、学生管理、排课管理、选课管理。系统工具类 2 个：备份数据、恢复数据。')

tag('4.2 管理页面架构重构（三次迭代）')
P('管理页面的组件架构经历了三次重大迭代。第一次使用 st.tabs 按"新增/修改/删除"分 tab，发现 Streamlit 同时渲染所有 tab 导致组件互相干扰，改为 st.radio 单模式渲染。第二次使用 st.form + st.form_submit_button 包裹表单，发现在动态 key 下修改时预填值不更新、保存按钮无效，全部改为 st.button + 动态 key + session_state 消息传递。第三次新增班级管理自动生成班级名功能，年级和专业改为 selectbox 下拉联动。')

tag('4.3 备份恢复功能加固')
P('备份命令补全 --routines（存储过程和函数）、--triggers（触发器）、--add-drop-table（恢复前先 DROP）三个参数。恢复功能弃用 DROP DATABASE 方式，改用 SET FOREIGN_KEY_CHECKS=0 包裹 SQL 文件内容，避免因外键约束导致恢复中断。增加 Windows 下 tempfile 临时文件处理机制，解决管道死锁问题。')

tag('4.4 导航系统升级')
P('侧边栏从 st.radio 手工切换页面升级为 st.navigation 原生多页导航系统。22 个页面按角色分三组：学生（5 页）、教师（3 页）、教务（14 页）。退出登录后自动触发 st.rerun 回到登录页，导航菜单自动重建。')

tag('4.5 连接池引入')
P('引入 SQLAlchemy QueuePool 管理数据库连接。配置 pool_size=5（常驻连接数）、max_overflow=10（峰值额外连接数）、pool_pre_ping=True（取连接前自动检测）、pool_recycle=3600（1 小时回收）。_make_page 包装器使用 try/finally 模式确保连接安全归还。')

doc.add_page_break()

# ════════════════════════════════════════
# 7月13日
# ════════════════════════════════════════
H('第四天：2026年7月13日 — 代码质量修复与规范化')

tag('5.1 目录结构规范化')
P('按照标准化项目结构对目录进行重组：SQL/ → sql/，document/上交/ → document/submission/。清理 document/ 根目录下过时的日志副本，确保 dev-log/ 为唯一日志来源。添加 README.md 项目说明文件，包含快速开始、技术栈、项目结构等基本信息。')

tag('5.2 代码质量清零')
P('对项目进行全面的代码质量检查与修复。修复 Pylint 全部警告，包括未使用的 import、变量命名不规范、函数过长等问题。统一算术运算符空格规范（Flake8 E226）。统一代码格式规范。')

tag('5.3 环境配置完善')
P('数据库连接配置迁移至 .env 文件（DB_HOST/DB_PORT/DB_USER/DB_PASSWORD/DB_NAME），支持环境变量覆盖默认值。更新 requirements.txt，移除未使用的 numpy，补充遗漏的 python-dotenv、sqlalchemy、pydantic 依赖。')

doc.add_page_break()

# ════════════════════════════════════════
# 7月14日
# ════════════════════════════════════════
H('第五天：2026年7月14日 — 功能增强与交互式图表')

tag('6.1 Pydantic 输入校验系统')
P('引入 Pydantic 模型进行全表单输入校验。创建 StudentCreate、TeacherCreate、ClassCreate、CourseCreate、GradeRecord 等校验模型，统一通过 validate_or_error 函数调用。学生学号校验 8-12 位纯数字，教师工号校验 T+数字格式，学分校验范围 0.5~30.0，成绩校验范围 0~100。校验失败时通过 st.toast 显示中文错误提示。所有 st.text_input 配置 max_chars 参数，前端与后端长度限制严格对齐。')

tag('6.2 Plotly 交互式图表系统')
P('引入 Plotly Express 图表库。数据概览页面增加 3 个饼图（学生/班级/教师状态分布）和 2 个柱状图（各专业班级数、各年级学生数）。班级学生名单页面增加学籍分和绩点分分布饼图，按照优秀（≥90/≥4.0）、良好（80~89/3.0~3.9）、中等（70~79/2.0~2.9）、及格（60~69/1.5~1.9）、不及格（<60/<1.5）分档。班级成绩统计增加成绩分布柱状图。教师查看课程学生页面增加成绩分布饼图与排行柱状图（左右并排）。学生我的成绩页面增加各科成绩柱状图并标注 60 分红色虚线及格线。教师列表增加排课数和选课学生数的排行柱状图。')

tag('6.3 专业管理模块')
P('新增专业管理模块，专业列表以 data/majors.csv 文件存储。提供前端图形化增删操作，新增时校验专业名长度（不超过 50 字符）和重复性，删除时自动检查是否有班级正在使用该专业。majors.csv 初始包含 26 个计算机/IT 相关专业，后扩展至 160+ 个专业覆盖工、理、医、文、法、经、管、艺、体等全部学科门类。')

tag('6.4 班级管理升级')
P('年级改为 selectbox 下拉选择（当前年份±5 年），专业从 CSV 专业列表下拉选择。系统自动查询该年级+专业已有班级的最大序号，自动生成班级名（格式：{年级}{专业}{序号}班）。修改模式中班级名、年级、专业改为只读显示，仅允许修改状态（在读/毕业）。')

tag('6.5 排课管理时间选择器优化')
P('排课上限从 st.number_input 改为 st.text_input + max_chars=5 组合，彻底解决 JavaScript Number 类型处理超大整数时的精度丢失问题。配合后端 int() 转换和 max_s > 99999 校验构建双重防线。排课时间选择从原始的 date_input + 两个 selectbox（时/分）共 9 个组件的方式，升级为 st.datetime_input 原生组件，砍掉 108 行冗余代码。')

tag('6.6 连接池安全加固')
P('_make_page 包装器增加 except 分支，在页面函数抛出异常时先执行 conn.rollback() 再 raise，最后由 finally 中 conn.close() 归还连接。防止业务代码异常后连接池中存在残留的未提交事务。CSV 文件读取全部统一为 utf-8-sig 编码，兼容 Windows 平台 Excel 编辑后产生的 BOM 头。')

tag('6.7 登录体验优化')
P('登录页面新增"快速选择"下拉框，测试账号从数据库实时读取（仅显示在职教师和在读学生）。输入校验前移：登录按钮点击后先验证格式（三个分支精确匹配 admin/教师工号 T+数字/学生学号纯数字 8-12 位），格式不符时直接 st.error 中断，不发起数据库查询。')

doc.add_page_break()

# ════════════════════════════════════════
# 7月15日
# ════════════════════════════════════════
H('第六天：2026年7月15日 — 大数据量支持与性能优化')

tag('7.1 200 万数据生成器')
P('创建 test/generate_large_data.py 一站式自动化数据生成器。从 reset_data.py 导入的基础数据（1000 学生、25 教师、54 排课）出发，逐步扩展至 20,000 学生、300 教师、329 班级（前 50 专业×3 年级×2~3 班）、1,004 排课（新增 950 条直接插入数据库），最终生成 2,002,688 条选课记录。')

tag('7.2 LOAD DATA INFILE 极速导入')
P('采用 LOAD DATA LOCAL INFILE 替代逐条 INSERT，导入速度提升约 20 倍。核心优化策略：导入前通过 Python 直连删除 5 个触发器（避免 200 万×5 次触发器调用），导入后手动批量 UPDATE 一次性重算全部 20,000 名学生的 GPA。CSV 中空成绩用 \\N（MySQL NULL 标准写法）而非空字符串。')

tag('7.3 生成器修复与迭代')
P('生成器开发过程中修复了多项问题：姓名库从 2640 种组合扩展至约 40 万种（引入双字名），唯一键冲突通过 used_pairs 集合去重，高斯分布成绩用 max(0, min(100, raw)) 截断防越界，information_schema.COLUMNS 的 CROSS JOIN 从 3 次减为 2 次避免 MySQL 产生 540 亿行中间结果导致崩溃，排课容量全部设为 99999 防止触发器拦截。')

tag('7.4 选课管理性能优化')
P('选课管理从全量加载 200 万条记录改为按学期+学号精准查询。首先通过轻量查询获取可选学期列表（SELECT DISTINCT co.semester），用户选择学期后输入学号，只查询该学生该学期的几门选课记录。退选操作改为直接输入选课 ID，彻底移除导致浏览器卡顿的 200 万选项下拉框。')

tag('7.5 专业列表扩展')
P('专业列表从初始的 26 个计算机相关专业扩展至 160+ 个，涵盖计算机科学与技术、软件工程、数据科学、人工智能、机械工程、电气工程、土木工程、建筑学、临床医学、药学、法学、经济学、金融学、会计学、汉语言文学、英语、新闻学、心理学、美术学、音乐学、体育教育等全部学科门类。')

tag('7.6 数据校验与防御性编程')
P('班级删除操作增加在读学生数检查（SELECT COUNT(*) FROM student WHERE class_id=? AND is_deleted=0），有学生时拒绝删除并提示。student_full_list 新增 AND c.is_deleted=0 过滤已删除班级的学生。teacher_course_teachers 新增 AND t.status=1 过滤离职教师。grade_roster_data 返回原生 NULL 而非字符串"未录入"，统一数据类型。')

doc.add_page_break()

# ════════════════════════════════════════
# 7月16日
# ════════════════════════════════════════
H('第七天：2026年7月16日 — 文档收尾与答辩准备')

tag('8.1 答辩 PPT 制作')
P('使用 python-pptx 库制作 11 页精美答辩 PPT。采用深蓝主题配色（#1A5276）、卡片式布局和彩色图标装饰。内容涵盖项目概述、系统架构（三层架构图+6 大技术栈卡片）、数据库设计（8 张表网格展示+数据库对象卡片）、三角色功能模块、6 大特色功能亮点、数据规模与性能优化方案、AI 辅助开发记录（含有效应用+人工修正+反思）、总结与致谢等完整章节。')

tag('8.2 课程设计说明书更新')
P('使用 python-docx 重构课程设计说明书为 9 章 15000 字版本。全面更新数据规模（120→20000 学生、16→300 教师、40→1004 排课、200→200 万选课）。补充新增功能描述。文档中标注了 10 处截图插入位置。')

tag('8.3 代码附录生成')
P('创建代码附录文档（document/report/代码附录.docx），整理全部 23 个源代码文件：7 个 SQL 脚本、5 个 Python 核心模块、3 个 Streamlit 页面、5 个工具模块、3 个数据脚本。采用 Consolas 等宽字体 + F5F5F5 灰色底纹排版，每 80 行分割一个段落避免打开卡顿。')

tag('8.4 说明书修改清单')
P('编写说明书修改清单文档，逐章标注需要修改的内容。按优先级分为三类：🔴高优先级（数据量更新 10+ 处、架构替换 2 段）、🟡中优先级（技术栈补充、AI 修正对照表新增）、🟢低优先级（截图插入 10 张）。同时提供基于 python-docx 的精准文本替换脚本，只替换文字内容不破坏原有排版。')

tag('8.5 源码打包')
P('使用 Python zipfile 打包完整项目源码为 4.2 MB 的压缩包，包含 57 个文件。排除 node_modules/、venv/、.git/、__pycache__/ 等非源码目录，排除 .mp4、.pdf、.pptx 等大文件，保留完整的项目目录结构。')

tag('8.6 版本管理')
P('整个项目通过 Git 进行版本管理，累计 80+ 次提交，每次提交附带清晰的 commit message 说明变更内容和原因。项目托管于 GitHub 远程仓库，确保数据安全。')

# ══ 保存 ══
doc.save(OUT)
sz = os.path.getsize(OUT) // 1024
print(f'OK: {sz} KB -> {OUT}')
